# -*- coding: utf-8 -*-
"""Generate the Parking Clarifications Response memo (reply to PIU request 16.05.2026). Concise version."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x55)
GREY = RGBColor(0x55, 0x55, 0x55)
HDRFILL = "1F3955"
ALTFILL = "EEF2F6"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.05

for lvl, sz in [("Heading 1", 13), ("Heading 2", 11.5)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = NAVY
    st.font.bold = True
    st.paragraph_format.space_before = Pt(8)
    st.paragraph_format.space_after = Pt(3)


def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexfill)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, white=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Calibri"
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)


def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, white=True)
        shade(t.rows[0].cells[i], HDRFILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, bold=(i == 0))
            if ri % 2 == 1:
                shade(cells[i], ALTFILL)
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ---- header ----
p = doc.add_paragraph()
r = p.add_run("RESPONSE TO CLARIFICATIONS REQUEST — PARKING")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(2)

meta = doc.add_table(rows=4, cols=2)
mrows = [
    ("In response to", "PIU — “Clarifications request on methodology and data sources”, 16 May 2026"),
    ("Prepared by / date", "Zurab Beradze — 25 June 2026"),
    ("Scope", "Parking-related items only (request Items 1, 5, 8, 9)"),
    ("Sources", "Output 8 and Output 10; project GeoJSON and yerevan-parking.vercel.app"),
]
for i, (k, v) in enumerate(mrows):
    set_cell(meta.rows[i].cells[0], k, bold=True, size=9.5)
    set_cell(meta.rows[i].cells[1], v, size=9.5)
    meta.rows[i].cells[0].width = Inches(1.6)
    meta.rows[i].cells[1].width = Inches(4.9)
for el in meta._tbl.iter():
    if el.tag == qn('w:tblBorders'):
        meta._tbl.remove(el)
doc.add_paragraph()

# ---- summary ----
doc.add_heading("In brief", level=1)
doc.add_paragraph(
    "Parking conclusions rest on two project-collected evidence bases — a desk-based supply inventory "
    "(15,897 spaces, 15% field-checked) and a field occupancy survey (six representative areas, hourly "
    "licence-plate sweeps 07:00–24:00, 29 May–3 June 2026). No prior model or ticketing data was used. "
    "Occupancy, duration, turnover and the displacement ratio are observed; supply totals and the scale of "
    "removal are inventory-based; the one figure carrying a real assumption — near-100% re-absorption of "
    "displaced demand — is conditional and flagged as such."
)

# ---- Item 9: data obtained ----
doc.add_heading("Parking data — obtained and used (Item 9)", level=1)
add_table(
    ["Data type", "Status", "What it supports"],
    [
        ["Occupancy", "Observed (field)", "Peak occupancy 93%; 69% of zones above 85%"],
        ["Duration of stay", "Observed", "Short-stay kerb: ~63% of stays ≤1 hour"],
        ["Turnover", "Observed", "5–11 vehicles per space per day"],
        ["Licence plate (manual)", "Observed", "Basis for duration/turnover; legality (~0.2% illegal)"],
        ["ANPR / payment", "Not available", "Not used. Would help payment-compliance & overnight demand if obtained"],
        ["Off-street", "Inventory only", "Capacity (7,035 spaces); occupancy not measured"],
    ],
    widths=[1.6, 1.4, 3.5],
)
doc.add_paragraph(
    "Note: occupancy is a kerbside (on-street) measurement. Off-street occupancy, overnight-residential demand "
    "and payment compliance were not observed."
)

# ---- Item 1: what each conclusion rests on ----
doc.add_heading("What each conclusion rests on (Item 1)", level=1)
add_table(
    ["Conclusion", "Basis", "Confidence"],
    [
        ["Supply 15,897 (8,862 on / 7,035 off)", "Inventory + 15% check", "High"],
        ["Paid-zone share 11.3%; signage; marking", "Inventory", "High"],
        ["Removal 5,953 (~85% of kerb)", "Inventory × corridor design", "High (retained ~1,052 preliminary)"],
        ["Occupancy / duration / turnover", "Observed", "High (weekday daytime)"],
        ["Displaced demand = 55% of removed", "Observed (local)", "High"],
        ["Re-absorption approaching 100%", "Conditional", "Depends on opening gated yards (92% of capacity)"],
    ],
    widths=[3.0, 1.9, 1.6],
)
doc.add_paragraph("No parking conclusion depends on a transport model — the displacement figures are empirical, not modelled.")

# ---- Item 5: methodology ----
doc.add_heading("Methodology — planned vs delivered (Item 5)", level=1)
add_table(
    ["Component", "Delivered as"],
    [
        ["Supply survey", "Desk inventory (360° video, satellite, Yandex) + 15% field cross-check, replacing the corridor-wide manual survey"],
        ["Occupancy survey", "Six representative areas (restored after initial descope), not the full 34 km"],
        ["Off-street occupancy", "Not surveyed — capacity only"],
    ],
    widths=[1.8, 4.7],
)
doc.add_paragraph(
    "The change was proportionate — the central-median design removes most kerb parking, so detailed temporal "
    "data on a baseline that will not survive has limited value — and was communicated to ADB and the PIU. "
    "The methodology-paper revisions and correspondence can be supplied as the decision record."
)

# ---- limitations / risk ----
doc.add_heading("Limitations for the risk matrix", level=1)
add_table(
    ["Limitation", "How to read it"],
    [
        ["Off-street capacity counted gross (92% gated yards)", "Re-absorption near 100% is a ceiling, conditional on opening those yards (mitigation in Output 10)"],
        ["Daytime window (07:00–24:00)", "Under-counts overnight-residential demand"],
        ["Six areas, not full corridor", "Representative sample, not a census"],
        ["Retained ~1,052 from concept design", "Preliminary until detailed design; removal of 5,953 is the firm figure"],
    ],
    widths=[2.9, 3.6],
)

# ---- reproducibility one-liner ----
doc.add_heading("Reproducibility (Item 8)", level=1)
doc.add_paragraph(
    "The parking figures are reproducible from project assets: the GeoJSON dataset (994 features), the "
    "metric-computation scripts (run live from the data), the KML layer and yerevan-parking.vercel.app. "
    "No prior parking model (WYG/ALG/Mott MacDonald) is an input; European BHLS precedent is corroboration only."
)

out = "Field Surveys/Field Surveys Report/Response to Clarifications Request - Parking - 25062026.docx"
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
