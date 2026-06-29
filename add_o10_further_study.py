#!/usr/bin/env python
"""Output 10 §7.2 'Further Studies Required': add a brief red mention of further
research into residential-yard administration into the EXISTING paragraph, and
remove the earlier verbose stand-alone paragraph if present. Edited IN PLACE on
the (rev) file so manual changes there are preserved. Idempotent.
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

REV = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Output 10 - Parking Analysis Report - 22062026 (rev).docx")
RED = RGBColor(0xC0, 0x00, 0x00)

OLD_MARKER = "Residential-yard opening — dedicated feasibility and implementation study"
NEW_SENTENCE = (
    " The administration of residential-yard opening in particular — who should lead it and "
    "under what organisational and legal arrangement — warrants a dedicated study before any "
    "roll-out."
)


def main():
    doc = Document(REV)
    paras = doc.paragraphs

    # 1) remove the earlier verbose stand-alone paragraph, if it was added
    for p in paras:
        if OLD_MARKER in p.text:
            p._p.getparent().remove(p._p)
            print("removed verbose paragraph")
            break

    # 2) locate the 7.2 body paragraph (first non-heading paragraph after the heading)
    paras = doc.paragraphs  # refresh after potential removal
    head = next((p for p in paras if p.text.strip().startswith("7.2") and "Further Studies" in p.text), None)
    if head is None:
        raise LookupError("7.2 Further Studies heading not found")
    children = list(doc.element.body.iterchildren())
    head_idx = children.index(head._p)
    body_para = None
    for el in children[head_idx + 1:]:
        if el.tag != qn("w:p"):
            break
        para = Paragraph(el, head._parent)
        if para.style.name and para.style.name.startswith("Heading"):
            break
        if para.text.strip():
            body_para = para
            break
    if body_para is None:
        raise LookupError("7.2 body paragraph not found")

    if NEW_SENTENCE.strip() in body_para.text:
        print("SKIP: sentence already present")
    else:
        r = body_para.add_run(NEW_SENTENCE)
        r.font.color.rgb = RED
        print("appended brief red sentence")

    doc.save(REV)
    print("WROTE:", REV)


if __name__ == "__main__":
    main()
