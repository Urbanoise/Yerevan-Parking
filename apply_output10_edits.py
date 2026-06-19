#!/usr/bin/env python
"""Apply the approved Output 10 redline edits to a working copy, in RED.

Source: "Parking Analysis Report - 20260417.docx" (= Output 10; its cover label
"Output 8: Parking Analysis Report" is a known mislabel). Edits land in
"...20260417 (rev).docx"; the source is left untouched. Additions are red;
superseded text is red strikethrough.

Follows the approved change-list (Output 10 - Proposed Changes (Redline Plan).docx)
as revised after client review:
  - truck-pocket %  -> show the calculation basis (comment 1)
  - Tbilisi          -> cautionary case, not a model (comment 2)
  - flat annual permit -> flagged IDEA pending PCS data, not a committed measure (comment 3)
  - Gai Avenue       -> plain wording + site-specific carve-out, not a new zone class (comment 0)
"""
import copy
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Parking Analysis Report - 20260417.docx")
REV = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Parking Analysis Report - 20260417 (rev).docx")

RED = RGBColor(0xC0, 0x00, 0x00)

shutil.copyfile(SRC, REV)
doc = Document(REV)
paras = doc.paragraphs


def find(sub, start=0):
    for i in range(start, len(paras)):
        if sub in paras[i].text:
            return paras[i]
    raise LookupError(f"not found: {sub!r}")


def _style_runs(p, color=RED, strike=False):
    for r in p.runs:
        r.font.color.rgb = color
        if strike:
            r.font.strike = True


def strike(sub, start=0):
    p = find(sub, start)
    _style_runs(p, RED, strike=True)
    return p


def _newpara_like(anchor):
    """New empty paragraph cloned from anchor's pPr (style/spacing)."""
    new_p = anchor._p.makeelement(qn("w:p"), {})
    pPr = anchor._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    return Paragraph(new_p, anchor._parent)


def _fill(p, text, bold_lead=None, size=None):
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = RED
        if size:
            r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.color.rgb = RED
    if size:
        r.font.size = Pt(size)
    return p


def ins_after(anchor_sub, text, bold_lead=None, start=0):
    anchor = find(anchor_sub, start)
    p = _newpara_like(anchor)
    _fill(p, text, bold_lead)
    anchor._p.addnext(p._p)
    return p


def ins_before(anchor_sub, text, bold_lead=None, start=0):
    anchor = find(anchor_sub, start)
    p = _newpara_like(anchor)
    _fill(p, text, bold_lead)
    anchor._p.addprevious(p._p)
    return p


def append_heading_title(anchor_sub, new_title):
    """Strike the old heading and append a red new title in the same para."""
    p = find(anchor_sub)
    _style_runs(p, RED, strike=True)
    r = p.add_run("  " + new_title)
    r.bold = True
    r.font.color.rgb = RED


def measure_note(anchor_sub, text):
    """Red 'Field-survey verdict:' note inserted after a measure's action line."""
    return ins_after(anchor_sub, text, bold_lead="Field-survey verdict.  ")


# ============================================================ 0. COVER LABEL (textbox)
# "Output 8: Parking Analysis Report" lives in two textboxes (drawing + VML
# fallback); doc.paragraphs does not reach them, so fix at the XML level.
OLD_COVER = "Output 8: Parking Analysis Report"
NEW_COVER = "Output 10: Parking Analysis Report"
cover_fixed = 0
for p_el in doc.element.iter(qn("w:p")):
    tnodes = p_el.findall(".//" + qn("w:t"))
    full = "".join(t.text or "" for t in tnodes)
    if full.strip() == OLD_COVER:
        # put the corrected string in the first run, blank the rest, colour red
        for idx, t in enumerate(tnodes):
            t.text = NEW_COVER if idx == 0 else ""
        for r_el in p_el.findall(".//" + qn("w:r")):
            rPr = r_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = r_el.makeelement(qn("w:rPr"), {})
                r_el.insert(0, rPr)
            for old in rPr.findall(qn("w:color")):
                rPr.remove(old)
            col = rPr.makeelement(qn("w:color"), {qn("w:val"): "C00000"})
            rPr.append(col)
        cover_fixed += 1
print("cover labels fixed:", cover_fixed)

# ============================================================ 1. PURPOSE & SCOPE
ins_after("Detailed survey methodologies, raw datasets, block-level inventories",
    "Revision note: a targeted field occupancy survey has since been completed on six "
    "representative areas, on ordinary mid-week working days. Its measured results — occupancy, "
    "duration, turnover and a locally-measured displacement test — now underpin the impact and "
    "mitigation analysis in this report and are documented in full in the companion Traffic and "
    "Parking Surveys and Analysis Report (Output 8).")

# ============================================================ 2. TARIFF / PAYMENT -> flat annual permit (flagged idea)
ins_after("This multi-channel payment ecosystem",
    "One of these subscription options is a flat annual zonal-parking permit, bought once for "
    "unlimited use of a parking zone for the year. Because a flat annual fee removes the marginal "
    "cost of each additional hour parked, it works against the curb turnover the corridor strategy "
    "depends on. A candidate idea — flagged here for decision, not adopted as a measure — is to "
    "phase this permit out in the medium-to-long term in favour of occupancy-based pricing. It "
    "would first need to be substantiated with Parking City Service permit-uptake and tariff data "
    "(see Mitigation Measures for Yerevan).")

# ============================================================ 3. SURVEY SECTION REFRAME (ANPR -> conducted)
ins_after("Should such data become available during the project",
    "Superseded: the demand-side evidence in this report is no longer dependent on the municipal "
    "ANPR transaction database. A proportionate field occupancy survey of six representative areas "
    "was carried out (late May–early June 2026), providing the occupancy, duration and turnover "
    "evidence directly. The municipal ANPR record remains useful to the city for ongoing monitoring "
    "and for the two dimensions the daytime survey cannot fully observe — payment compliance and "
    "true overnight-residential demand — but it is not relied upon for the demand analysis here.")

append_heading_title("Occupancy and Turnover Metrics from ANPR Data",
    "→ superseded by the conducted field occupancy survey")

append_heading_title("Overview of Previous Field Survey Plans",
    "→ Field Occupancy Survey (Conducted)")
ins_after("While field surveys were abandoned, it is important to document",
    "Revision: this position was subsequently revised. Rather than forgo occupancy evidence "
    "altogether, a proportionate field occupancy survey — six representative areas (Komitas, "
    "Shiraz/Hasratyan, Garegin Nzhdeh, Gai Avenue/Mega Mall, Kentron and Malatia-Sebastia) rather "
    "than the full corridor — was conducted on five separate ordinary working days, with hourly "
    "licence-plate sweeps from 07:00 to 24:00 recording time, zone, plate, vehicle type, kerb "
    "location, parking method and legality. The planning described below was therefore largely "
    "executed, not merely prepared; the measured results are summarised in this report and detailed "
    "in Output 8.")

# ============================================================ 4. DIFFERENTIATED IMPACTS BY LAND USE
ins_after("Along Arshakunyats Avenue, the wide right-of-way",
    "Measured land-use signature (field survey). The survey confirms and quantifies these "
    "differences. The Kentron commercial core runs at 138% peak occupancy and is short-stay "
    "dominated (about 63% of vehicles stay one hour or less). The Malatia-Sebastia residential "
    "segment shows the lowest daytime pressure (54% peak) but the strongest long-stay / overnight "
    "signature, consistent with residential storage rather than turnover. Turnover across the "
    "surveyed areas runs at 5–11 vehicles per space per day.")

# ============================================================ 5. IMPACTS ON SPECIFIC USER GROUPS
ins_after("Persons with disabilities who depend on proximate parking access",
    "Measured user groups (field survey). The survey lets these groups be anchored to observed "
    "behaviour rather than assumption. Visitors and customers are the dominant group (about 63% of "
    "vehicles stay one hour or less, 77% two hours or less). Commuters and workers are a small "
    "all-day tail (around 6%), better addressed as a BRT mode-shift opportunity than as a fining "
    "target. Residents appear as an overnight share (about 7%) that the daytime window under-counts "
    "and that should be protected, not weakened. Delivery and freight show as a measurable truck "
    "presence concentrated in pockets — truck share, calculated as trucks divided by all classified "
    "vehicles surveyed in the area, reaches 12.6% in Malatia-Sebastia (548 of 4,339 vehicles) and "
    "8.0% in Shiraz (198 of 2,469), against 4.9% citywide — so loading provision should be sized to "
    "those pockets rather than uniformly. Taxi and persons-with-disabilities provision is retained "
    "by design.")

# ============================================================ 6. DISPLACEMENT POTENTIAL BY CORRIDOR ZONE
ins_after("Lower-sensitivity zones include the outer sections of both corridors",
    "Measured recalibration (field survey). The survey tests these a-priori labels against "
    "behaviour. Kentron is confirmed high-sensitivity (138% peak) but for an off-street-led reason "
    "— only about 14% of its displaced demand can be re-absorbed on nearby streets. Komitas, "
    "previously medium, behaves as high-sensitivity (123% peak, 0% on-street absorption) and is "
    "reclassified accordingly. Shiraz/Hasratyan and Malatia-Sebastia behave as genuinely "
    "lower-sensitivity (75% and 54% peak). Gai Avenue (Mega Mall) is a special case: it is the one "
    "surveyed area where, even counting every nearby off-street yard at full capacity, there is "
    "still not enough room to re-absorb the cars displaced by kerb removal (123% peak, best-case "
    "absorption below 100%). It therefore does not fit any of the three standard packages; the "
    "simplest handling is a site-specific carve-out — defer kerb removal here and lean on the Mega "
    "Mall's own off-street structure — rather than creating a new zone category. Across the surveyed "
    "areas, peak displaced demand was 908 vehicles against 1,643 spaces removed (55%), confirming "
    "with local measurement that displaced demand falls well below the supply removed.")

# ============================================================ 7. YEREVAN-SPECIFIC FRAMEWORK (off-street dependency)
ins_after("Third, the Soviet-era urban form that characterises",
    "Fourth — and load-bearing — the re-absorption of displaced demand is conditional on off-street "
    "capacity. About 92% of the absorptive capacity that closes the aggregate gap to roughly 100% is "
    "gated residential yards, counted at gross capacity as if already open (Kentron about 86% "
    "yard-dependent, Komitas about 100%). Re-absorption is therefore not automatic: it depends on "
    "those yards being opened and actively managed, which is why Open Residential Yards is treated "
    "below as a committed precondition rather than an optional add-on.")

# ============================================================ 8. MITIGATION MEASURES — per-measure verdicts
measure_note("Roll out red/blue-line system onto side streets within 100 m",
    "Right lever (about 90% of the kerb is currently unpriced), but in the core it rations demand "
    "rather than absorbing it — surviving side-street capacity is only about 14% in Kentron and 0% "
    "in Komitas. Frame extension in the core as demand management, not as a place to re-house "
    "displaced cars.")

measure_note("Designate time-restricted loading bays where possible",
    "Confirmed by the short-access dominance in the survey. Size loading bays to the measured "
    "freight pockets (truck share = trucks ÷ all classified vehicles surveyed in the area: 12.6% in "
    "Malatia-Sebastia, 8.0% in Shiraz, 4.9% citywide) rather than distributing them uniformly.")

measure_note("Enforce maximum stay durations",
    "Re-scope, do not remove. The data contradicts the premise: about 63% of vehicles already stay "
    "one hour or less, all-day parking is only around 6%, and turnover is already 5–11 per space per "
    "day. Demote visitor caps to a targeted contingency for genuinely low-turnover spots, lead with "
    "pricing, and treat commuters as a BRT mode-shift opportunity rather than a fining target.")

measure_note("This permit arrangement should be understood as a transitional instrument",
    "Keep, but design it actively, using Tbilisi as a cautionary case rather than a model: Tbilisi "
    "currently grants up to two free vehicles per apartment with no paid resident permit — precisely "
    "the long-stay, low-turnover pattern to avoid — which is why its new strategy is tightening this "
    "to one per apartment. Yerevan's permit should be priced and/or quantity-limited from the outset, "
    "with yards and permits actively managed and clearly communicated to residents.")

measure_note("Delineate and sign currently informal parking",
    "Strongly confirmed (about 88% of corridor parking is unmarked; roughly 24.5% sits off the "
    "carriageway). Widen this beyond the lower-pressure outer zones — Komitas (about 35% informal) "
    "and Kentron (about 60% angled) need organising most.")

measure_note("Park & ride is a peripheral parking measure under which drivers leave",
    "Do not oversell. The daytime fleet is overwhelmingly short-stay, which weakens the 'many "
    "commuters to intercept' premise. Keep park-and-ride as a study item and frame BRT/P&R as serving "
    "the commuter, not as a primary displacement sink.")

measure_note("Encourage building associations to open gated courtyard parking",
    "Elevate to a committed precondition (load-bearing). Gated residential yards carry about 92% of "
    "the off-street absorptive capacity, and the roughly 100% aggregate re-absorption is conditional "
    "on opening them. Move this from an optional outer-zone pilot to a core, council-led precondition "
    "with a resident-engagement and management plan so it does not backfire.")

ins_before("Taken together, these measures form a layered mitigation package",
    "Phase out the flat annual permit — flagged idea, not yet a measure. The flat annual fee removes "
    "the marginal cost of dwell time and so suppresses the curb turnover the data shows is the binding "
    "lever. It is presented here as a candidate to cancel in the medium-to-long term and replace with "
    "occupancy-based pricing (the same logic as unwinding free resident parking), but it must first be "
    "substantiated with Parking City Service permit-uptake and tariff data before it is adopted as a "
    "measure.",
    bold_lead="")

# ============================================================ 9. PHASING (load-bearing sequencing)
ins_after("A three-phase approach is recommended",
    "Because re-absorption depends on the gated off-street yards (above), opening and actively "
    "managing them — together with the pricing — must precede kerb removal. The off-street "
    "dependency makes 'mitigation before removal' load-bearing rather than advisory: removing kerb "
    "capacity ahead of the yards would strand the displaced demand the survey measures.")

# ============================================================ 10. LIMITATIONS AND RESIDUAL RISKS
ins_after("It is important to acknowledge that not all displaced parking",
    "Survey-specific caveats. The field survey observed a daytime window (07:00–24:00) and so "
    "under-counts true overnight-residential demand; it covered six representative areas rather than "
    "the full corridor; and the absorption figures are gross best-case (yards counted as if already "
    "open). One residual risk is specific: Gai Avenue (Mega Mall) is the single surveyed area where "
    "even best-case capacity cannot fully absorb the displaced demand, and is handled as a "
    "site-specific carve-out (defer kerb removal pending the mall's off-street parking) rather than "
    "through the standard zone packages.")

# ============================================================ 11. RECOMMENDATIONS — packages by sensitivity zone
ins_after("This zoning approach is important because it turns the mitigation strategy",
    "Measured recalibration of the packages. Komitas moves into the high-sensitivity zone (123% "
    "peak, 0% on-street absorption). Open Residential Yards moves up into the high-sensitivity "
    "package as a committed precondition. Gai Avenue (Mega Mall) is handled as a site-specific "
    "carve-out — defer kerb removal pending the mall's off-street parking — rather than as a new "
    "zone class. The candidate idea to phase out the flat annual permit is noted against the "
    "cross-cutting package as a flagged option, pending Parking City Service data. Figure 10 "
    "(measures by zone) and the sensitivity map are updated to match.")

# ============================================================ 12. CONCLUSIONS (measured)
ins_after("Enforcement and modal shift underpin demand reduction",
    "The demand side is now measured, not inferred. A targeted field occupancy survey of six "
    "representative areas recorded capacity-weighted peak occupancy of 93% (69% of zones above the "
    "85% threshold), a short-stay, high-turnover kerb (about 63% of stays one hour or less; turnover "
    "5–11 per space per day), and peak displaced demand equal to only 55% of the spaces removed "
    "(908 against 1,643). Aggregate re-absorption approaching 100% is achievable but conditional on "
    "opening the gated residential yards, which carry about 92% of the absorptive capacity. The "
    "measured behaviour recalibrates the zoning — Komitas behaves as high-sensitivity (reclassify "
    "from medium) and Gai Avenue is handled as a site-specific carve-out — and reduces, but does not "
    "eliminate, the items requiring further study (overnight-residential demand, full-corridor "
    "extension, freight counts, and ANPR payment-compliance remain genuine gaps).")

doc.save(REV)
print("WROTE:", REV)
