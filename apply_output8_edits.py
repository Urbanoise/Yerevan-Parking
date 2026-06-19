#!/usr/bin/env python
"""Apply the approved Output 8 redline edits to a working copy, in RED.

Additions are red; deletions are red strikethrough. Source is left untouched;
edits land in "...08042026 (rev).docx". Edits follow the approved change-list
(Output 8 - Proposed Changes (Redline Plan).docx), incl. the two resolved
review comments: drop the 24,950 figure from the summary, and delete the
ANPR-as-future-demand thread.
"""
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Parking Surveys and Analysis Report 08042026.docx")
REV = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Parking Surveys and Analysis Report 08042026 (rev).docx")

RED = RGBColor(0xC0, 0x00, 0x00)

shutil.copyfile(SRC, REV)
doc = Document(REV)
paras = doc.paragraphs


def find(sub, start=0):
    for i in range(start, len(paras)):
        if sub in paras[i].text:
            return paras[i]
    raise LookupError(f"not found: {sub!r}")


def _style_runs(p, color=RED, strike=False):
    for r in p.runs:
        r.font.color.rgb = color
        if strike:
            r.font.strike = True


def strike(sub, start=0):
    p = find(sub, start)
    _style_runs(p, RED, strike=True)
    return p


def _newpara_like(anchor):
    """Create a new empty paragraph cloned from anchor's pPr (style/spacing)."""
    new_p = anchor._p.makeelement(qn("w:p"), {})
    pPr = anchor._p.find(qn("w:pPr"))
    if pPr is not None:
        import copy
        new_p.append(copy.deepcopy(pPr))
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, anchor._parent)


def _fill(p, text, bold_lead=None, size=None):
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = RED
        if size:
            r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.color.rgb = RED
    if size:
        r.font.size = Pt(size)
    return p


def ins_after(anchor_sub, text, bold_lead=None, start=0):
    anchor = find(anchor_sub, start)
    p = _newpara_like(anchor)
    _fill(p, text, bold_lead)
    anchor._p.addnext(p._p)
    return p


def ins_before(anchor_sub, text, bold_lead=None, start=0):
    anchor = find(anchor_sub, start)
    p = _newpara_like(anchor)
    _fill(p, text, bold_lead)
    anchor._p.addprevious(p._p)
    return p


def append_heading_title(anchor_sub, new_title):
    """Strike the old heading text and append a red new title in the same para."""
    p = find(anchor_sub)
    _style_runs(p, RED, strike=True)
    r = p.add_run("  " + new_title)
    r.bold = True
    r.font.color.rgb = RED


# ---------------------------------------------------------------- A. Exec Summary
ins_before("Introduction",
    "A targeted field occupancy survey has since been completed on six representative areas spanning the "
    "corridor's sensitivity range, surveyed on ordinary mid-week working days (29 May – 3 June 2026). It "
    "supplies the demand-side evidence previously unavailable. Across the surveyed areas, capacity-"
    "weighted peak occupancy reaches 93%, with 69% of zones exceeding the 85% efficiency threshold at the "
    "peak hour; parking is overwhelmingly short-stay (about 63% of vehicles stay one hour or less, 77% two "
    "hours or less) with turnover of 5–11 vehicles per space per day; and peak displaced demand equals "
    "only 55% of the spaces removed (908 cars against 1,643 spaces in the surveyed areas) — confirming, "
    "with local measurement, the international finding that displaced demand is materially lower than the "
    "supply removed.")

# ---------------------------------------------------------------- B. Purpose & Scope
ins_after("The report should be read in conjunction",
    "This revision incorporates the results of a targeted field occupancy survey carried out after the "
    "desk-based inventory, on six representative areas. The deliverable therefore now presents both "
    "supply-side evidence (the inventory) and measured demand-side evidence (occupancy, duration, turnover "
    "and locally-measured displacement).")

# ---------------------------------------------------------------- C. ToR
ins_after("the survey methodology was refined to prioritise",
    "The occupancy and duration component of the ToR (¶59–63), initially descoped on proportionality "
    "grounds, has subsequently been satisfied in representative form through this targeted field survey. "
    "The deliverable therefore meets the intent of the ToR rather than departing from it.")

# ---------------------------------------------------------------- D. Evolution
ins_after("The primary purpose of this project is to support",
    "This position was subsequently revised. Rather than forgo occupancy evidence entirely, a proportionate "
    "field occupancy survey — six representative areas rather than the full 34 km — was conducted in late "
    "May and early June 2026, restoring the occupancy, duration and turnover evidence that the corridor "
    "design and the displacement assessment require. The areas were surveyed on five separate ordinary "
    "working days (Komitas, Friday 29 May; Shiraz/Hasratyan, Monday 1 June; Garegin Nzhdeh and Gai Avenue, "
    "Tuesday 2 June; Kentron and Malatia-Sebastia, Wednesday 3 June), with no weekend or public holiday; "
    "because each area fell on a different day, a one-off disruption would show as a single-area outlier, "
    "whereas the behavioural patterns repeat across all six.")

# ---------------------------------------------------------------- E. Component 2 (ANPR -> field survey)
strike("Component 2: Municipal ANPR Data Integration")
ins_after("Component 2: Municipal ANPR Data Integration",
    "Component 2: Field Occupancy Survey (targeted). Hourly licence-plate sweeps from 07:00 to 24:00 "
    "across the six representative areas, recording for every parked vehicle the time, zone, plate, "
    "vehicle type, kerb location, parking method and legality. This provides the demand-side evidence — "
    "occupancy, duration and turnover — directly. (The municipal ANPR system maintained by Parking City "
    "Service CJSC remains available to the city for ongoing monitoring, but is not relied upon for the "
    "demand analysis in this report.)")

# ---------------------------------------------------------------- F. Justification chapter
append_heading_title("Justification for Excluding Full-Scale Occupancy and Duration Surveys",
    "→ Survey Scope: From Full-Corridor Census to Targeted Occupancy Survey")
ins_after("A formal justification letter was prepared",
    "Note (revision): The scope decision recorded in this section was subsequently revised. Rather than "
    "excluding occupancy and duration work altogether, a proportionate field occupancy survey was "
    "conducted on six representative areas (see ‘Occupancy and Demand: Field Survey Results’). The "
    "resource-proportionality argument below stands — a full 34-kilometre census remained disproportionate "
    "— but the conclusion that occupancy data was unnecessary, or had to await municipal ANPR data, no "
    "longer applies.")
strike("In lieu of field-based occupancy surveys, the ANPR transaction database")
ins_after("In lieu of field-based occupancy surveys, the ANPR transaction database",
    "Superseded: the demand-side evidence in this report is provided by the field occupancy survey "
    "described above, not by the ANPR database.")

# ---------------------------------------------------------------- G. Supply config/location bridge
ins_after("The predominance of parallel parking (75.3%)",
    "Field-survey behaviour (complementary). The inventory above records marked configuration; the field "
    "survey additionally recorded how vehicles actually parked. By vehicle, 55% parked parallel, 26% at "
    "45° and 19% perpendicular; measured by space, angled layouts concentrate in Kentron (about 60%) "
    "against 4% or less elsewhere. About 24.5% of observed parking sat off the carriageway (13.0% on "
    "footpaths and 11.5% in setbacks) — informal encroachment the supply map does not capture, highest in "
    "Komitas (35%). Legality was recorded for every vehicle: only about 0.2% was flagged as illegally "
    "parked, indicating that the kerbside problem is saturation and tolerated informality rather than "
    "overt illegality.")

# ---------------------------------------------------------------- H. Impact by Sensitivity Zone
ins_after("Impact by Sensitivity Zone",
    "Measured recalibration (field survey). The survey tests the a-priori sensitivity labels against "
    "behaviour. Kentron is confirmed high-sensitivity (138% peak occupancy) but for an off-street-led "
    "reason — only 14% of its displaced demand can be re-absorbed on nearby streets. Komitas, previously "
    "classed medium, behaves as high-sensitivity (123% peak, 0% on-street absorption) and should be "
    "reclassified accordingly. Garegin Nzhdeh sits between (92% peak, 30% absorption). Shiraz/Hasratyan "
    "and Malatia-Sebastia behave as genuinely lower-sensitivity (75% / 69% and 54% / 97% respectively). "
    "Gai Avenue (Mega Mall area) is an unclassified gap: 123% peak demand with total best-case absorption "
    "below 100% — the only surveyed area where even gross capacity cannot fully absorb the displaced "
    "demand, and it is not represented in the zone table below.")

# Table 9 — append measured findings (red) into each zone row's last cell
def append_cell_red(cell, text):
    p = cell.paragraphs[-1]
    r = p.add_run(text)
    r.font.color.rgb = RED

for tbl in doc.tables:
    head = [c.text.strip() for c in tbl.rows[0].cells]
    if head and head[0].startswith("Sensitivity") and any("Corridor" in h for h in head):
        for row in tbl.rows[1:]:
            label = row.cells[0].text.strip().lower()
            last = row.cells[-1]
            if label.startswith("high"):
                append_cell_red(last, "  Field survey (Kentron): 138% peak, only 14% absorbable on-street — confirmed high, off-street-led.")
            elif label.startswith("medium"):
                append_cell_red(last, "  Field survey: Komitas 123% peak / 0% absorption — behaves high, reclassify; Garegin 92% / 30%.")
            elif label.startswith("lower"):
                append_cell_red(last, "  Field survey: Shiraz 75% / 69%, Malatia 54% / 97% — confirmed lower.")
        break

# ---------------------------------------------------------------- I. Displacement Assessment
ins_after("International experience with European BHLS corridors",
    "Locally measured. Across the surveyed areas, peak displaced demand was 908 vehicles against 1,643 "
    "spaces removed — 55%. The principle that displaced demand falls well below the supply removed is "
    "therefore now measured in Yerevan, not merely inferred from European precedent.")
ins_after("1,857 on-street spaces on cross-streets within the influence area",
    "Measured absorption (surveyed areas). Of the peak displaced demand, surviving on-street capacity "
    "absorbs only about 43% (14% in the Kentron core, 0% in Komitas, up to 97% in Malatia-Sebastia). The "
    "balance is closed by off-street capacity — but about 92% of that off-street capacity is gated "
    "residential yards counted at gross capacity as if already open. The conclusion that displaced demand "
    "can be re-absorbed (approaching 100% in aggregate) is therefore conditional on those yards being "
    "opened and actively managed; this dependency, and the measures to address it, are carried in the "
    "companion Parking Analysis Report (Output 10).")

# ---------------------------------------------------------------- J. Occupancy: Current Status -> Results
append_heading_title("Occupancy and Demand Data: Current Status",
    "→ Occupancy and Demand: Field Survey Results")
ins_after("Occupancy and Demand Data: Current Status",
    "The demand-side analysis has been completed through the targeted field occupancy survey. Measured "
    "against the international 85% efficiency threshold, capacity-weighted peak occupancy across the "
    "surveyed areas reaches 93%, with 69% of surveyed zones exceeding 85% at the peak hour; it ranges from "
    "54% (Malatia-Sebastia) to 138% (Kentron). Duration of stay is dominated by short visits — about 63% "
    "of vehicles stay one hour or less and 77% two hours or less, while genuine all-day storage (eight "
    "hours or more) is about 7% — distinguishing high-turnover retail and errand parking from a small "
    "commuter and residential long-stay tail. Turnover runs at 5–11 vehicles per space per day. "
    "Payment-compliance and true overnight-residential demand remain the two dimensions the daytime survey "
    "cannot fully observe; these are the only items for which the municipal ANPR record, if released, "
    "would still add value.")
strike("The demand-side analysis depends on access to the ANPR")
strike("Occupancy rates benchmarked against the international 85% efficiency threshold")
strike("Duration of stay distributions distinguishing commuter parking from retail turnover")
strike("Turnover rates measuring commercial street vitality")
strike("Payment compliance ratios evaluating enforcement effectiveness")
strike("As of this report date, formal data-sharing approval")

# ---------------------------------------------------------------- K. Conclusions
ins_before("The full interactive results are available",
    "6. A targeted field occupancy survey on six representative areas (surveyed on ordinary working days) "
    "measured the demand side: capacity-weighted peak occupancy of 93%, with 69% of zones over the 85% "
    "threshold; a short-stay, high-turnover kerb (about 63% of stays one hour or less; turnover 5–11 per "
    "space per day); and peak displaced demand equal to only 55% of spaces removed (908 against 1,643 in "
    "the surveyed areas). 7. Aggregate re-absorption approaching 100% is achievable but conditional on "
    "opening gated residential off-street yards, which carry about 92% of the absorptive capacity. 8. The "
    "measured behaviour recalibrates the sensitivity zoning — Komitas behaves as high-sensitivity "
    "(reclassify from medium) and Gai Avenue is an unclassified local gap; demand-side mitigation measures "
    "are detailed in the companion Parking Analysis Report (Output 10).")

doc.save(REV)
print("WROTE:", REV)
