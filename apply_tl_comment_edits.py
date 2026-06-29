#!/usr/bin/env python
"""Apply the two Team-Leader comment edits to Output 8 and Output 10, in RED.

Single reproducible pass: copies each source fresh to "...22062026 (rev).docx" and
applies all red additions, so each rev carries every edit. Sources untouched.

COMMENT 1 — measured off-street occupancy (O8 + O10).
  The hourly off-street occupancy from the field survey (six representative
  facilities, one per area; field-survey-yards.geojson) was never folded into
  either report; both argued absorption from nominal inventory capacity alone.
  Added: O8 gets a narrative + a per-facility table; O10 gets a matching paragraph.
  Framing: avg ~52% shows real observed spare, but at peak three of six are
  at/over capacity. The six surveyed (accessible) facilities are excluded from the
  displacement absorber count because they already carry their own demand
  (compute_field_survey_metrics.mjs L380-382); the ~92% gated-yard dependency
  refers to the still-closed residential yards. So the data GROUNDS and REINFORCES
  the existing conclusion rather than softening it.

COMMENT 2 — residential-yard implementation, organizational options (O10 only).
  The "Open Residential Yards" measure stopped at "encourage building associations
  ... via EasyPark/Carsleep" without naming WHO drives it or HOW. Per direction,
  O10 should NOT prescribe an implementation mechanism (that is a separate future
  feasibility/implementation study); it should (a) cite a few international cases
  where opening under-used residential/institutional parking has worked, and
  (b) lay out the candidate organizational models (state-agency-led / PPP /
  association-led / combination) without committing. Backed by a precedent-research
  pass (China shared-parking, ParqEx/WhereiPark, Parking Benefit Districts; honest
  caveat that the concept is largely untested in a post-Soviet courtyard context
  where land title is often unresolved).
"""
import copy
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

RPT = r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
FILES = {
    "O8": (RPT + "Output 8 - Parking Surveys and Analysis Report 22062026.docx",
           RPT + "Output 8 - Parking Surveys and Analysis Report 22062026 (rev).docx"),
    "O10": (RPT + "Output 10 - Parking Analysis Report - 22062026.docx",
            RPT + "Output 10 - Parking Analysis Report - 22062026 (rev).docx"),
}
RED = RGBColor(0xC0, 0x00, 0x00)

OFFSTREET_HEAD = ["Area", "Surveyed off-street facility", "Formal spaces",
                  "Avg. occupancy", "Peak-hour occupancy"]
OFFSTREET_ROWS = [
    ["Kentron",                         "Nalbandyan yard",   "60",  "81%",  "142%"],
    ["Komitas",                         "Komitas City lot",  "123", "56%",  "101%"],
    ["Garegin Nzhdeh",                  "GN off-street",     "40",  "22%",  "45%"],
    ["Gai Avenue (Mega Mall)",          "Palace lot",        "32",  "115%", "194%"],
    ["Malatia-Sebastia",                "Sebastia yard",     "101", "32%",  "77%"],
    ["Shiraz / Hasratyan",              "Shiraz off-street", "71",  "37%",  "59%"],
    ["All six (total / weighted avg.)", "—",                 "427", "52%",  "—"],
]


# ---------------------------------------------------------------- helpers
def find(paras, sub, start=0):
    for i in range(start, len(paras)):
        if sub in paras[i].text:
            return paras[i]
    raise LookupError(f"not found: {sub!r}")


def _newpara_like(anchor):
    new_p = anchor._p.makeelement(qn("w:p"), {})
    pPr = anchor._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    return Paragraph(new_p, anchor._parent)


def _fill(p, text, bold_lead=None, size=None, italic=False):
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = RED
        if size:
            r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.color.rgb = RED
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def ins_after(anchor, text, bold_lead=None, size=None, italic=False):
    """Insert a red paragraph directly after `anchor` (a Paragraph); return it."""
    p = _newpara_like(anchor)
    _fill(p, text, bold_lead=bold_lead, size=size, italic=italic)
    anchor._p.addnext(p._p)
    return p


def ins_sequence_after(anchor, blocks):
    """Insert several red paragraphs in order after `anchor`. blocks: list of
    (bold_lead, text) tuples."""
    prev = anchor
    for bold_lead, text in blocks:
        prev = ins_after(prev, text, bold_lead=bold_lead)
    return prev


# ---- off-street table -------------------------------------------------------
def _set_cell_red(cell, text, bold=False, size=9):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.color.rgb = RED
    r.bold = bold
    r.font.size = Pt(size)


def _table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C00000")
        borders.append(el)
    tblPr.append(borders)


def build_offstreet_table(doc):
    tbl = doc.add_table(rows=1 + len(OFFSTREET_ROWS), cols=len(OFFSTREET_HEAD))
    tbl.alignment = 1
    for j, h in enumerate(OFFSTREET_HEAD):
        _set_cell_red(tbl.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(OFFSTREET_ROWS, start=1):
        bold = row[0].startswith("All six")
        for j, val in enumerate(row):
            _set_cell_red(tbl.rows[i].cells[j], val, bold=bold)
    _table_borders(tbl)
    return tbl


# ================================================================ COMMENT 1 text
O8_OCC_INTRO = (
    "Measured off-street occupancy (field survey). The off-street side of this "
    "balance is now partly observed, not only inventoried. In each of the six areas the "
    "survey recorded hourly occupancy in one representative off-street facility (427 "
    "surveyed spaces in total) using the same licence-plate method as on-street. Averaged "
    "across the day these accessible facilities run about 52% full — on the order of 205 "
    "spaces of observed spare on average — so real off-street headroom demonstrably exists. "
    "At the peak hour, however, three of the six are already at or beyond their marked "
    "capacity (Palace 194%, Nalbandyan 142%, Komitas City 101%) and only three keep "
    "meaningful spare (Garegin 45%, Shiraz 59%, Sebastia 77%). These surveyed facilities are "
    "accordingly excluded from the absorptive-capacity figure above — they already carry "
    "their own demand — so the ~92% gated-yard dependency refers to the still-closed "
    "residential yards that could not be surveyed. The measured occupancy thus grounds the "
    "absorption argument in field observation while reinforcing, not relaxing, its central "
    "condition: the spare capacity that must take displaced demand at the peak lies in the "
    "gated yards that remain shut, so opening them is a precondition rather than an assumption."
)
O8_OCC_CAPTION = (
    "Table — Measured occupancy of the six surveyed off-street facilities (one per area), "
    "field survey 29 May – 3 June 2026. Occupancy is shown as a share of formal (striped) "
    "capacity: average is the mean vehicles present across the 07:00–24:00 window; peak-hour "
    "is the busiest single hour. Values above 100% indicate vehicles parked beyond marked "
    "capacity (aisle stacking / double-parking). Peaks occur at different hours and are not "
    "additive across facilities."
)
O10_OCC_PARA = (
    "Measured off-street occupancy (field survey). This dependency is now supported by "
    "observation, not assumed. In each of the six areas the survey recorded hourly occupancy "
    "in one representative off-street facility (427 spaces). Averaged across the day these "
    "already-accessible facilities run about 52% full, so observed spare capacity does exist; "
    "but at the peak hour three of the six are already at or over their marked capacity "
    "(Palace 194%, Nalbandyan 142%, Komitas City 101%), with only Garegin (45%), Shiraz (59%) "
    "and Sebastia (77%) retaining real headroom. The facilities that are already in use are "
    "themselves busy at the peak — which is why they are set aside from the absorptive count "
    "— and the capacity that closes the displacement gap sits in the gated residential yards "
    "that remain closed. The measured occupancy therefore strengthens rather than relaxes the "
    "case for treating Open Residential Yards as a load-bearing precondition; the per-facility "
    "figures are tabulated in the companion Output 8."
)

# ================================================================ COMMENT 2 text (O10)
O10_YARD_BLOCKS = [
    ("International precedent (illustrative).  ",
     "Opening under-used residential and institutional parking to controlled daytime use is "
     "established practice in several countries, which is encouraging for feasibility even "
     "though the organisational model must be designed for Yerevan. A city-owned intermediary "
     "company, Stockholm Parkering (Sweden), brokers under-used private off-street stalls to "
     "motorists and helps developers meet parking requirements without new construction — a "
     "direct model for a city-backed entity bridging individual building associations and the "
     "parking authority. In Vancouver's West End, the city aligned on-street permit prices "
     "with market-rate off-street prices so that residents moved vehicles off the kerb into "
     "under-used building parkades, underlining that while courtyard parking remains free, "
     "kerbside demand cannot be managed. China's municipally-brokered 'shared parking' "
     "programmes open residential compounds and institutional lots on a time-staggered basis "
     "through a city app with automatic number-plate (ANPR) gate access; Hangzhou's scheme "
     "splits the revenue among the individual owner, the owners' committee (the equivalent of "
     "a Yerevan homeowners' association, hamatirutyun) and the building management, making the "
     "association a paid partner rather than only a gatekeeper. North American condominium and "
     "homeowner associations monetise surplus and visitor spaces through managed platforms "
     "(for example ParqEx and WhereiPark) under board control, and peer-to-peer platforms in "
     "the United Kingdom (JustPark), Japan (Akippa) and Australia (Share with Oscar) show the "
     "supporting machinery that builds resident trust: dedicated parking-sharing insurance, a "
     "small tax-free allowance on rental income, address privacy until a booking is confirmed, "
     "and legal rights that run with the land when a property is sold. A 'parking benefit "
     "district' approach — earmarking part of the revenue for visible local improvements on "
     "the same blocks — has repeatedly been the mechanism that turns residents from opponents "
     "into supporters."),
    ("Organisational options (to be selected, not fixed here).  ",
     "The same idea can be delivered through several organisational models, and choosing "
     "among them is a decision for the implementation stage: (i) public / state-agency-led, "
     "where a municipal parking authority or Parking City Service — or a dedicated city-owned "
     "company on the Stockholm Parkering model — brokers, leases and enforces the arrangement "
     "while the association consents and is paid; (ii) a "
     "public-private partnership, where a licensed operator or parking platform runs booking, "
     "access and payment under a municipal framework and shares revenue with the association; "
     "(iii) association-led and city-enabled, where the homeowners' association opens and "
     "manages its own courtyard under a standard municipal template and light regulation; or "
     "some combination of these. Each carries different requirements for the legal "
     "instrument, the revenue split, the tariff and enforcement."),
    ("A dedicated implementation study, not prescribed here.  ",
     "This report does not prescribe which model to adopt or how to implement it. Because the "
     "concept is, on the evidence reviewed, largely untested in a post-Soviet courtyard "
     "context — where the more common pattern has been to gate yards against outsiders, and "
     "where courtyard land ownership is often unresolved — the detailed design should be "
     "taken forward as a dedicated feasibility and implementation study before any roll-out. "
     "That study would resolve the load-bearing preconditions: who legally owns each "
     "courtyard and may therefore authorise the arrangement and retain revenue; the consent "
     "procedure and vote threshold under Armenian condominium law; the tax and insurance "
     "treatment of association parking income; the revenue-sharing and access-control model; "
     "and the pricing and enforcement approach. Sequencing should also follow the "
     "international experience, in which single-owner institutional and commercial lots open "
     "far more readily than residential courtyards and should be activated first."),
]


# ---------------------------------------------------------------- appliers
def edit_o8(src, rev):
    shutil.copyfile(src, rev)
    doc = Document(rev)
    # Comment 1: off-street occupancy narrative + table + caption
    anchor = find(doc.paragraphs, "Measured absorption (surveyed areas)")
    intro = ins_after(anchor, O8_OCC_INTRO)
    tbl = build_offstreet_table(doc)
    intro._p.addnext(tbl._tbl)
    caption = _newpara_like(anchor)
    _fill(caption, O8_OCC_CAPTION, size=8, italic=True)
    tbl._tbl.addnext(caption._p)
    doc.save(rev)
    print("WROTE:", rev)


def edit_o10(src, rev):
    shutil.copyfile(src, rev)
    doc = Document(rev)
    # Comment 1: off-street occupancy paragraph
    occ_anchor = find(doc.paragraphs, "re-absorption of displaced demand is conditional on off-street")
    ins_after(occ_anchor, O10_OCC_PARA)
    # Comment 2: residential-yard precedents + organisational options + defer-to-study
    yard_anchor = find(doc.paragraphs, "Elevate to a committed precondition")
    ins_sequence_after(yard_anchor, O10_YARD_BLOCKS)
    doc.save(rev)
    print("WROTE:", rev)


if __name__ == "__main__":
    edit_o8(*FILES["O8"])
    edit_o10(*FILES["O10"])
