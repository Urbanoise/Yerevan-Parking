#!/usr/bin/env python
"""Generate Post-Survey Debrief Interview Summary (.docx).

Documents the structured debrief call (30 Jun 2026) between the STS consulting
team and the Armenian field team, recording on-the-ground conditions, methodology
verification, and field observations for the Yerevan Parking Management Study.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Feeback/Post-Survey Debrief Interview - Summary Record.docx"

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED    = RGBColor(0xC0, 0x2A, 0x2A)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)
GREY   = RGBColor(0x55, 0x55, 0x55)
AMBER  = RGBColor(0xE6, 0x5C, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ── Base font ─────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def para(text="", bold=False, italic=False, size=10.5, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0,
         style_name="Normal"):
    p = doc.add_paragraph(style=style_name)
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color if color else RGBColor(0x1A, 0x1A, 0x1A)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size  = Pt(13)
    run.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),  "single")
    btm.set(qn("w:sz"),   "6")
    btm.set(qn("w:space"),"1")
    btm.set(qn("w:color"), "006D77")
    pBdr.append(btm)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = ACCENT
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25 + 0.2 * level)
    if bold_prefix:
        r0 = p.add_run(bold_prefix + " ")
        r0.bold = True
        r0.font.size = Pt(10.5)
        r0.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def callout(text, color_hex="142A4A", label="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p2 = cell.paragraphs[0]
    p2.paragraph_format.space_after  = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    r1 = p2.add_run(label + "  ")
    r1.bold = True
    r1.font.size  = Pt(10)
    r1.font.color.rgb = WHITE
    r2 = p2.add_run(text)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def two_col_row(tbl, label, value, label_bg="142A4A", shade_row=False):
    row  = tbl.add_row()
    lc, vc = row.cells[0], row.cells[1]
    set_cell_bg(lc, label_bg)
    if shade_row:
        set_cell_bg(vc, "F0F4F8")
    lp = lc.paragraphs[0]
    r = lp.add_run(label)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    vp = vc.paragraphs[0]
    r2 = vp.add_run(value)
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


# ══════════════════════════════════════════════════════════════════════════════
# COVER BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONCEPTUAL DESIGN AND DUE DILIGENCE FOR THE YEREVAN SUSTAINABLE URBAN TRANSPORT AND ELECTRIC MOBILITY")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Post-Survey Debrief Interview")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Summary Record of Field Team Interview")
r.bold = False; r.font.size = Pt(13); r.font.color.rgb = ACCENT
p.paragraph_format.space_after  = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Zurab Beradze  |  30 June 2026")
r.font.size = Pt(10); r.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Metadata table
meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
col_widths = [Inches(1.8), Inches(4.2)]
for i, w in enumerate(col_widths):
    for cell in meta.columns[i].cells:
        cell.width = w

rows_data = [
    ("Document Type",   "Post-Survey Debrief Interview — Summary Record"),
    ("Project",         "Conceptual Design and Due Diligence for the Yerevan Sustainable Urban Transport and Electric Mobility"),
    ("Client",          "Asian Development Bank (ADB)"),
    ("Prepared by",     "Zurab Beradze"),
    ("Interviewees",    "Hayk Sargsyan (Local Expert), Jorj Aslanyan (Field Survey Supervisor)"),
    ("Interview Date",  "30 June 2026"),
    ("Survey Period",   "June 2026 (dates noted per location in body text)"),
    ("Status",          "Final"),
]
for i, (lbl, val) in enumerate(rows_data):
    two_col_row(meta, lbl, val, shade_row=(i % 2 == 1))

doc.add_paragraph().paragraph_format.space_after = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# 1. BACKGROUND AND PURPOSE
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Background and Purpose")

para(
    "Following completion of the on-street parking occupancy and turnover surveys across "
    "six study areas in Yerevan, Zurab Beradze conducted a structured debrief interview "
    "with the Armenian field team on 30 June 2026. This is standard practice at the "
    "close of a field survey programme: the consulting team and the survey team review "
    "the data together, discuss on-the-ground conditions that affected individual zones, "
    "and verify that the observed patterns are consistent with what was experienced in "
    "the field.",
    space_after=8
)
para(
    "The interview was attended by Zurab Beradze and Nikoloz Archvadze (STS), "
    "Jorj Aslanyan (Field Survey Supervisor, Yerevan) and Hayk Sargsyan (Local Expert, "
    "Yerevan). The discussion covered data quality, site-specific conditions, "
    "interpretation of headline findings, and any anomalies or exceptions noted during "
    "the survey.",
    space_after=8
)
para(
    "This document summarises the interview, recording the field team's observations "
    "and confirmations alongside the relevant analytical findings.",
    space_after=8
)

callout(
    "All statements attributed to the field team in this document are drawn from "
    "a recorded and transcribed debrief call (30 Jun 2026). The original transcript "
    "is retained on file.",
    color_hex="006D77", label="SOURCE:"
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. SURVEY OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Survey Overview")

para(
    "Occupancy and turnover surveys were conducted across six study areas in Yerevan. "
    "Enumerators recorded vehicle licence plates and parking durations at regular intervals "
    "throughout the day, entering data into structured Excel workbooks. The zones surveyed "
    "correspond directly to the on-street parking segments mapped in the project GeoJSON "
    "database.",
    space_after=8
)

heading2("2.1  Study Areas")
areas = [
    ("Komitas",            "Major arterial corridor with high retail and service density"),
    ("Kentron",            "City centre — high-density commercial and civic uses"),
    ("Garegin Nzhdeh",     "Mixed residential-commercial; active pedestrian environment"),
    ("Gai Avenue (Mega Mall)", "Suburban retail corridor adjacent to Mega Mall"),
    ("Shiraz & Hasratyan",    "Residential-adjacent streets;"),
    ("Malatia-Sebastia",   "Outer district; mix of residential and local commerce"),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("Study Area").bold = True
hdr[1].paragraphs[0].add_run("Character").bold = True
set_cell_bg(hdr[0], "142A4A"); set_cell_bg(hdr[1], "142A4A")
for c in hdr:
    c.paragraphs[0].runs[0].font.color.rgb = WHITE
    c.paragraphs[0].runs[0].font.size = Pt(10)
for i, (area, desc) in enumerate(areas):
    row = tbl.add_row()
    if i % 2 == 1:
        set_cell_bg(row.cells[0], "F0F4F8")
        set_cell_bg(row.cells[1], "F0F4F8")
    for j, txt in enumerate([area, desc]):
        r = row.cells[j].paragraphs[0].add_run(txt)
        r.font.size = Pt(10)
        if j == 0:
            r.bold = True
doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# 3. DATA QUALITY AND SUPERVISION
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Data Quality and Supervision")

heading2("3.1  On-Site Supervisor Coverage")
para(
    "The field team was led by Jorj Aslanyan, who served as on-site supervisor across "
    "the majority of survey days. During the interview he described his approach to "
    "quality management:",
    space_after=4
)
bullet("He was personally present in the field for four of the six survey days for extended periods of eight to nine hours, and attended others in a supervisory capacity.")
bullet("After each survey day he reviewed all Excel workbooks submitted by enumerators, checking for duplicate entries, cross-zone double-counting, and data-entry errors.")
bullet("Errors identified during review — for example, a vehicle recorded twice in adjacent zones during a transition — were corrected before files were forwarded to Zurab Beradze.")
bullet("Any anomaly flagged by enumerators in the field was communicated to the supervisor in real time, who directed the appropriate corrective action.")

para(
    "In the supervisor's own words: \"I always checked all these results myself. I was always "
    "in contact with [the enumerators]. After that I reviewed all the Excel files and also "
    "checked all these results.\" (Transcript, 23:31-23:51)",
    italic=True, color=GREY, space_after=10
)

heading2("3.2  Real-Time Cross-Checking During Surveys")
para(
    "Beyond post-hoc review, the supervisor described active quality management during "
    "data collection:",
    space_after=4
)
bullet("Enumerators who recorded the same vehicle in overlapping zones were identified and directed to exclude the duplicate entry.")
bullet("The supervisor personally observed parking behaviour at first hand — watching individual vehicles move between zones — to validate enumerator records.")
bullet("Licence plate sequences were compared hour-by-hour within each workbook to detect implausible repetitions.")

callout(
    "Zone-crossover events (a vehicle leaving one zone and re-parking in an adjacent "
    "zone within the same survey period) are treated as two distinct parking events, "
    "each with its own start time. This is consistent with standard parking survey "
    "methodology — the vehicle's dwell in each zone is independent and short.",
    color_hex="142A4A", label="METHODOLOGY:"
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. ON-SITE CONDITIONS AND EXCEPTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. On-Site Conditions and Exceptions Recorded During the Interview")

para(
    "A central purpose of the debrief interview was to document site-specific conditions "
    "that affected individual zones during the survey period. The field team reported "
    "the following, each of which was handled in a consistent and documented manner.",
    space_after=8
)

heading2("4.1  Zones Affected by Road Works or Traffic-Rule Changes")
para(
    "In certain zones parking was unavailable at the time of the survey due to active "
    "road works or changes to traffic regulations:",
    space_after=4
)
bullet(
    "Moskovyan Street: the street's traffic direction had been changed to one-way, making "
    "parking on one side prohibited under revised traffic rules.",
    bold_prefix="Moskovyan Street -"
)
bullet(
    "Republic Square vicinity: a specific zone adjacent to the square was subject to a "
    "no-parking restriction in force at the time of the survey.",
    bold_prefix="Republic Square -"
)
bullet(
    "Malatia-Sebastia (two locations): zones where parking was not possible due to "
    "physical conditions on the survey day.",
    bold_prefix="Malatia-Sebastia -"
)
para(
    "The supervisor noted: \"There were no closed streets, but there were places where "
    "road works were being carried out and cars could not park. And in some zones, traffic "
    "rules in Armenia had changed -- the street became one-way and parking there is not "
    "allowed.\" (Transcript, 9:17-10:44)",
    italic=True, color=GREY, space_after=4
)
callout(
    "Zones where parking was prohibited at the time of survey were recorded as empty "
    "(zero occupancy). They are retained in the zone inventory but excluded from "
    "occupancy rate calculations.",
    color_hex="006D77", label="HANDLING:"
)

heading2("4.2  Zone Substitutions")
para(
    "Two zones could not be surveyed as originally designated and were substituted with "
    "adjacent zones of equivalent character, with the field coordinator's approval:",
    space_after=4
)
bullet(
    "Malatia-Sebastia: the designated zone was inaccessible on the day; the team "
    "contacted the coordinator (Hayk Sargsyan), who authorised a replacement zone "
    "approximately one block away.",
    bold_prefix="Malatia-Sebastia -"
)
bullet(
    "Kaipoghot (Gai Avenue area): a zone designated as part of the Mercedes-Benz "
    "dealership parking was closed to public access. It was excluded and the survey "
    "continued on adjacent street segments.",
    bold_prefix="Kaipoghot / Gai Avenue -"
)
para(
    "Both substitutions were approved in real time by the project coordinator and are "
    "reflected in the final zone inventory.",
    color=GREY, space_after=10
)

heading2("4.3  Time-Restricted Zones (Komitas, Zone 115)")
para(
    "Zone 115 in the Komitas area operated under a chain-barrier system "
    "(Armenian: 'shlangh-baun') that restricted entry after approximately 21:00-22:00. "
    "The supervisor confirmed this restriction with on-site parking attendants. "
    "Vehicles already parked before the closure were recorded as overnight stayers, "
    "with their licence plates noted in each subsequent hourly observation.",
    space_after=10
)

heading2("4.4  Location with Low Daytime Demand (Kerkurov Street)")
para(
    "Kerkurov Street (surveyed as part of the Komitas corridor) displayed near-zero "
    "occupancy during midday hours. The supervisor visited the location personally "
    "after receiving enumerator reports and confirmed the observation: the street lacks "
    "the adjacent attractors that generate daytime parking demand. Demand increased "
    "after 18:00 as residents returned home.",
    space_after=4
)
para(
    "The supervisor confirmed: \"First, when people said there were no cars, they sent "
    "photos. I went there myself, looked, and indeed there were no cars.\" "
    "(Transcript, 15:41-16:05)",
    italic=True, color=GREY, space_after=10
)

heading2("4.5  Long-Term / All-Day Parked Vehicles")
para(
    "A small number of vehicles remained stationary throughout the entire survey window. "
    "The field team's approach:",
    space_after=4
)
bullet("Recorded with their licence plates in every hourly observation (not excluded from the dataset).")
bullet("Counted once for occupancy purposes; their full-day dwell contributes correctly to the turnover distribution.")
bullet("Estimated at approximately 10 vehicles across the Komitas corridor (the area with the longest supervisor coverage of 8-9 hours).")
para(
    "These vehicles form a distinct long-dwell cohort and do not distort the short-stay "
    "statistics.",
    color=GREY, space_after=10
)

heading2("4.6  Taxi and Courier Vehicle Activity")
para(
    "Several zones -- notably in the Komitas and Garegin Nzhdeh areas -- had identifiable "
    "clusters of taxi and food-delivery courier vehicles. The supervisor described their "
    "behaviour:",
    space_after=4
)
bullet("Taxis typically occupy a fixed waiting point, leave when dispatched, and return after completing a job.")
bullet("Food-delivery couriers working fast-food and restaurant zones park briefly, collect orders, and cycle rapidly through the same spots.")
bullet("These are real demand patterns, correctly captured by the survey methodology, and contribute to the observed high short-stay rates in commercial areas.")


# ══════════════════════════════════════════════════════════════════════════════
# 5. FIELD TEAM VERIFICATION OF HEADLINE FINDINGS
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Field Team Verification of Headline Findings")

para(
    "During the interview, the consulting team presented the headline analytical findings "
    "to the field team and invited their assessment of whether the results were consistent "
    "with what they had observed in the field. The following summarises each finding and "
    "the field team's response.",
    space_after=8
)

heading2("5.1  62% of Parking Events Last Less Than One Hour")
para(
    "When presented with the finding that approximately 62% of observed parking events "
    "had a duration under one hour, the field supervisor confirmed this was consistent "
    "with his field observations and offered the following explanation:",
    space_after=4
)
bullet(
    "Commercial land use: all six survey areas are characterised by high concentrations "
    "of retail outlets, supermarkets, service businesses (barbershops, pharmacies), "
    "fast-food restaurants, and offices -- trip purposes that generate short-stay demand "
    "by nature."
)
bullet(
    "Zone-crossover events: when a driver parks in one zone, moves to an adjacent zone "
    "to be closer to a destination, and parks again, each event is recorded as a separate "
    "short stay. Both events are real and correctly observed."
)
bullet(
    "Taxi and courier cycling: as described in section 4.6, high-turnover vehicle "
    "categories contribute disproportionately to the short-stay count."
)
para(
    "The supervisor's response: \"I think this is logical. I can explain it because "
    "these streets have many shops, many supermarkets, many services -- barber shops "
    "or whatever. Cars will not stay very long. They came, did their business, and "
    "left.\" (Transcript, 24:08-24:45)",
    italic=True, color=GREY, space_after=10
)

heading2("5.2  Approximately 1% of Observed Vehicles Parked Illegally")
para(
    "The field supervisor was asked whether the low observed rate of illegal parking "
    "(approximately 1% across all six areas) matched his experience on the ground. "
    "He confirmed it did, explaining:",
    space_after=4
)
bullet(
    "Zone design: the survey zones covered designated legal parking spaces. Areas "
    "where parking is structurally prohibited (bus stops, no-stopping zones, "
    "pedestrian areas) were not within the survey zone boundaries. As the supervisor "
    "put it: '99% of the zones you sent us were legal parking zones.' (Transcript, 20:23)"
)
bullet(
    "Enforcement environment: all six study areas have dense camera coverage and "
    "active police presence. The supervisor noted: 'There are many cameras and many "
    "police officers. People cannot freely park on the pavement.' (Transcript, 21:58-22:05)"
)
bullet(
    "The supervisor further contextualised this: the areas surveyed are among the "
    "busiest and most visible streets in Yerevan. Higher illegal parking rates would "
    "be expected in residential back-streets and peripheral areas, which were outside "
    "the survey scope."
)

heading2("5.3  Morning Peak at 07:00-08:00 Higher Than at 09:00")
para(
    "The supervisor noted -- consistent with the analytical data -- that occupancy at "
    "07:00-08:00 tended to exceed that at 09:00 in several zones. He attributed this "
    "to commuter behaviour: workers park early and retrieve vehicles late, leaving a "
    "mid-morning gap as turnover occurs. This pattern is consistent with standard "
    "commuter parking demand profiles.",
    space_after=10
)

heading2("5.4  Hasratyan -- Asymmetric Demand Across Street Sides")
para(
    "On Krikor Hasratyan Street, the supervisor described an observed asymmetry: the "
    "side of the street fronting active retail had consistently higher occupancy, while "
    "the opposite side -- facing an unactivated frontage -- attracted overflow demand. "
    "Drivers unable to find space on the active side crossed to park opposite, producing "
    "a pattern where the less-attractive kerb appeared well-used due to spillover. This "
    "is correctly captured by zone-level analysis.",
    space_after=10
)

heading2("5.5  Vehicles Appearing in Non-Consecutive Hours")
para(
    "The interview also discussed instances where the same licence plate appears in hour 1, "
    "is absent in hours 2-3, and reappears in hour 4. The supervisor confirmed this is "
    "genuine behaviour:",
    space_after=4
)
bullet("Residents who park in the morning, leave for work, and return in the evening.")
bullet("Taxi drivers who occupy a waiting spot, respond to a call, and return to the same spot.")
bullet("Delivery vehicles cycling between a depot and the same delivery zone.")
para(
    "The supervisor: \"I only just told you that there were places where several cars "
    "were parked for several hours, then they left. And later in the evening they came "
    "back. I think these are people who live near this parking area.\" "
    "(Transcript, 27:37-27:52)",
    italic=True, color=GREY, space_after=10
)


# ══════════════════════════════════════════════════════════════════════════════
# 6. EXTERNAL CONDITIONS DURING THE SURVEY PERIOD
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. External Conditions During the Survey Period")

heading2("6.1  Weather")
para(
    "The supervisor confirmed that no rain occurred on any of the survey days in which "
    "he was personally present (four of the six survey days). Conditions were described "
    "as clear and normal throughout, and were not identified as a factor affecting "
    "parking behaviour or data quality.",
    space_after=10
)

heading2("6.2  Special Events")
para(
    "A public concert was held at Republic Square during the survey period. Both the "
    "field coordinator and supervisor assessed whether this could have affected the data:",
    space_after=4
)
bullet("Concert-goers travel to Republic Square primarily on foot or by public transport; no surface parking is available at the square itself.")
bullet("The survey zones near Republic Square cover the surrounding streets. Vehicles already parked in those zones would not vacate for a pedestrian event.")
bullet("Political rallies and public gatherings in central Yerevan are similarly pedestrian-dominated and do not generate unusual parking demand.")
para(
    "The field coordinator concluded: \"Even when these public events take place in "
    "central Yerevan, people go on foot. Those who are already parked stay parked -- "
    "they are doing their own business. So, can this seriously affect the results? "
    "No.\" (Transcript, 34:15-34:50)",
    italic=True, color=GREY, space_after=10
)


# ══════════════════════════════════════════════════════════════════════════════
# 7. SUMMARY OF KEY TOPICS DISCUSSED
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Summary of Key Topics Discussed")

topics = [
    (
        "Short-stay rate: 62% of events under one hour",
        "Confirmed as consistent with field observations. Driven by the commercial land "
        "use mix across all six areas, zone-crossover events, and high-frequency taxi "
        "and courier activity."
    ),
    (
        "Illegal parking rate: approximately 1%",
        "Confirmed as expected. Survey zones by design covered legal parking areas; "
        "the studied corridors operate under heavy enforcement with cameras and police "
        "presence."
    ),
    (
        "Zone substitutions",
        "Two zones were inaccessible on the survey day and replaced with equivalent "
        "adjacent zones following real-time coordination with the field coordinator. "
        "Both substitutions are documented."
    ),
    (
        "Zones affected by road works or traffic-rule changes",
        "Affected zones (Moskovyan, Republic Square vicinity, two in Malatia-Sebastia) "
        "were recorded as empty and excluded from occupancy rate calculations."
    ),
    (
        "Long-term parked vehicles",
        "Recorded with licence plates every hour; counted once for occupancy. "
        "Approximately 10 all-day vehicles identified across the Komitas corridor."
    ),
    (
        "Special events (Republic Square concert)",
        "Assessed by the field team as having no material impact on surrounding "
        "street zones, as the event was pedestrian-oriented."
    ),
    (
        "Data quality and supervision",
        "Supervisor reviewed all Excel files post-collection, corrected duplicates "
        "and cross-zone errors, and was present in the field for the majority of "
        "survey days."
    ),
    (
        "Weather conditions",
        "Confirmed clear and normal on all survey days with supervisor presence. "
        "No rain events recorded."
    ),
]

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
h0, h1 = tbl2.rows[0].cells
set_cell_bg(h0, "142A4A"); set_cell_bg(h1, "142A4A")
for cell, txt in [(h0, "Topic"), (h1, "Field Team Confirmation / Explanation")]:
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

for i, (q, a) in enumerate(topics):
    row = tbl2.add_row()
    if i % 2 == 1:
        set_cell_bg(row.cells[0], "EEF2F6")
        set_cell_bg(row.cells[1], "EEF2F6")
    r0 = row.cells[0].paragraphs[0].add_run(q)
    r0.bold = True; r0.font.size = Pt(10)
    r1 = row.cells[1].paragraphs[0].add_run(a)
    r1.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Conclusions")

para(
    "The post-survey debrief interview confirmed that the field survey data underpinning "
    "the Yerevan Parking Management Study was collected under professional supervision, "
    "with active quality-control procedures and transparent, documented handling of "
    "site-specific exceptions.",
    space_after=8
)
para(
    "The headline findings -- including the high short-stay rate, the low illegal parking "
    "rate, and the observed demand patterns across the six study areas -- were reviewed "
    "with the Armenian field team and confirmed as consistent with their direct "
    "observations in the field. The field team's explanations for each pattern align "
    "with the analytical interpretation presented in the project reports.",
    space_after=8
)
para(
    "This record, together with the original survey data and transcript on file, "
    "provides a complete account of the survey's conduct and findings.",
    space_after=12
)

callout(
    "Supporting materials on file: original debrief call transcript (30 Jun 2026); "
    "raw Excel workbooks per survey zone; zone-level photographs taken by enumerators; "
    "GIS zone boundary files.",
    color_hex="142A4A", label="ON FILE:"
)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Signature block
para("Prepared by:", bold=True, size=10, space_after=2)
sig_tbl = doc.add_table(rows=1, cols=1)
sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
p2 = sig_tbl.rows[0].cells[0].paragraphs[0]
r = p2.add_run("Zurab Beradze\n")
r.bold = True; r.font.size = Pt(10.5)
r2 = p2.add_run("Senior Parking Management Specialist")
r2.font.size = Pt(10); r2.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(4)
para("30 June 2026", color=GREY, size=10, space_after=0)
para("Conceptual Design and Due Diligence for the Yerevan Sustainable Urban Transport and Electric Mobility", color=GREY, size=9, italic=True)


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
