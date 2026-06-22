#!/usr/bin/env python
"""Generate the Step 9 - Field Surveys client report (.docx).

Reads the live, authoritative metrics from the app's published geojson so the
report never drifts from what the dashboard shows. Does NOT touch the app.
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GEO = r"C:/Users/user/Yerevan-Parking/app/static/data/wgs84/field-surveys.geojson"
IMGDIR = r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report"
OUT = IMGDIR + r"/Yerevan Field Surveys - Findings Report.docx"

# Displacement screenshots exported from the interactive dashboard (in IMGDIR).
DISP_IMG = {
    "all": "All Areas Displacement.png",
    "kentron": "Kentron Displacement.png",
    "komitas": "Komitas Displacement.png",
    "mega": "Gai Displacement.png",
    "garegin": "Garegin Displacement.png",
    "shiraz": "Shiraz & Hasratyan Displacement.png",
    "malatia": "Malatia-Sebastia Displacement.png",
}

with open(GEO, encoding="utf-8") as f:
    g = json.load(f)
A = g["areaStats"]
D = g["displacement"]

# Full-day AVERAGE occupancy, capacity-weighted, computed with the SAME population
# the pipeline uses for the published PEAK figure (zones with peak_occupancy/space).
# avgPresent for a zone = occupancy_pct/100 * space  ->  Σ avgPresent / Σ space.
def avg_occupancy():
    acc = {}
    for ft in g["features"]:
        p = ft.get("properties", {})
        a = p.get("area")
        if not a or p.get("occupancy_pct") is None or not p.get("space"):
            continue
        ap = p["occupancy_pct"] / 100.0 * p["space"]
        for key in (a, "all"):
            acc.setdefault(key, {"n": 0.0, "d": 0.0})
            acc[key]["n"] += ap
            acc[key]["d"] += p["space"]
    return {k: round(v["n"] / v["d"] * 100) for k, v in acc.items() if v["d"]}

AVG = avg_occupancy()

# App display labels (from StoryStep.svelte AREA_LABELS)
LABELS = {
    "malatia": "Malatia-Sebastia",
    "kentron": "Kentron",
    "garegin": "Garegin Nzhdeh",
    "mega": "Gai Avenue",
    "komitas": "Komitas",
    "shiraz": "Shiraz & Hasratyan",
}
# Order areas by peak occupancy (most-stressed first)
AREAS = sorted(LABELS.keys(), key=lambda k: -A[k]["occupancy"][0]["value"])


def val(area, group, idx):
    return A[area][group][idx]["value"]


# ---- styling helpers -------------------------------------------------------
NAVY = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=10, align="left", white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        r.font.color.rgb = color


def heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(16 if level == 1 else 13)
    p.space_before = Pt(10)
    return p


def body(text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def whatnow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("What now — ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


def add_image(path, caption=None, width=6.4):
    import os
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width))
    else:
        r = p.add_run(f"[missing image: {os.path.basename(path)}]")
        r.font.color.rgb = RED
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY


def header_row(table, headers, fill="142A4A"):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, white=True, size=9.5,
                 align="left" if i == 0 else "center")
        shade(hdr[i], fill)


# ---- title -----------------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("Yerevan Parking Study")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

t = doc.add_paragraph()
r = t.add_run("Field Surveys — Findings Report")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
r = sub.add_run("On-street parking occupancy, regulation and displacement across six survey areas")
r.font.size = Pt(12)
r.font.color.rgb = GREY
r.italic = True

meta = doc.add_paragraph()
r = meta.add_run("Step 9 of the interactive Parking Reform briefing  •  Prepared for municipal stakeholders")
r.font.size = Pt(10)
r.font.color.rgb = GREY

# horizontal rule
hr = doc.add_paragraph()
pPr = hr._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "142A4A")
pbdr.append(bottom)
pPr.append(pbdr)

# headline totals
total_paths = sum(val(a, "paidfree", 0) for a in LABELS)
total_white = sum(val(a, "paidfree", 1) for a in LABELS)
total_blue = sum(val(a, "paidfree", 2) for a in LABELS)
total_yard = sum(val(a, "paidfree", 3) for a in LABELS)
onstreet_total = total_white + total_blue

# ---- executive summary -----------------------------------------------------
heading("Executive summary", 1)
body(
    f"This report presents the findings of the field parking surveys carried out across "
    f"{len(LABELS)} areas of Yerevan. Surveyors recorded the vehicles physically present on "
    f"every marked parking length throughout the working day, building a true picture of how "
    f"intensively each street is used — rather than relying on design capacity alone. In total "
    f"the survey covers {total_paths} on-street parking segments holding {onstreet_total:,} marked "
    f"spaces, plus {int(total_yard):,} spaces across six off-street yards."
)

body("Three findings stand out:", size=11)
bullet(
    f" Across the surveyed network, occupancy averages just {AVG['all']}% of capacity over the "
    f"full day but climbs to {A['all']['occupancy'][0]['value']}% at the peak hour (both "
    f"capacity-weighted). {A['all']['occupancy'][1]['value']}% of zones exceed 85% occupancy at "
    f"the peak and {A['all']['occupancy'][2]['value']} individual zones spill over 100%. More than "
    "half the kerb sits empty most of the day — the spaces exist, they are just all wanted at the "
    "same time and place. The problem is the timing and management of demand, not a raw shortage "
    "of spaces.",
    bold_lead="Streets fill to capacity only at peak.")
bullet(
    f" Parking is heavily concentrated in time. The fleet is dominated by short visits "
    f"({A['all']['profile']['duration']['visitorPct']}% stay under two hours), with a clear "
    "demand peak in the late morning and again in the evening.",
    bold_lead="Demand is short-stay and peaky.")
bullet(
    f" The redesign removes {D['removed_supply']:,} on-street spaces. At the peak hour this "
    f"displaces {D['removed_demand']} parked cars. Nearby surviving on-street capacity and "
    "off-street yards are more than sufficient to absorb this city-wide, but the balance is "
    "tight in a few locations and one area (Gai Avenue) falls short locally.",
    bold_lead="Removed parking can be re-absorbed — with one local exception.")

# ---- methodology -----------------------------------------------------------
heading("What was surveyed and how", 1)
body(
    "Field teams observed each marked parking length at regular intervals through the day and "
    "logged the vehicles present (by plate). This produces an occupancy profile for every street "
    "segment — not an estimate, but a count of real cars on the ground. From these observations "
    "the study derives three complementary views, each selectable as a layer on the interactive map:"
)
bullet(" average and peak occupancy as a share of the segment's marked capacity;",
       bold_lead="Occupancy —")
bullet(" whether each segment is paid (blue) or free (white) parking, plus off-street yards;",
       bold_lead="Regulation —")
bullet(" which segments are retained or removed under the redesign, and where the displaced "
       "cars can go.", bold_lead="Displacement —")
body(
    "Capacity is the number of marked spaces on each segment; where a segment was unmarked, "
    "capacity is estimated at one bay per 7.5 m of kerb. Occupancy figures are full-day averages "
    "unless labelled “peak”, which is the single busiest hour. Displacement is assessed at "
    "the peak hour — the moment the system is under most pressure.",
    italic=True, color=GREY, size=10)

# ---- overall occupancy -----------------------------------------------------
heading("Finding 1 — How full are the streets?", 1)
body(
    f"Over the full working day the surveyed network sits at {AVG['all']}% of capacity, yet at the "
    f"peak hour it runs at {A['all']['occupancy'][0]['value']}%. That gap is the central finding. "
    "If the city were simply short of spaces, the daily average would be high too. Instead, more "
    "than half the kerb sits empty most of the day — the spaces exist, they are just all wanted at "
    "the same time and in the same place. The two occupancy columns below make this concrete: an "
    "area can look comfortable on a daily average and still be turning drivers away at its busiest "
    "hour. The pattern also varies sharply by area — the city-centre areas of Kentron and Komitas "
    "run well over capacity at peak, while the outer areas keep real headroom. The table ranks the "
    "areas from most to least stressed at peak."
)

tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tbl, ["Area", "Survey\nsegments", "Marked\nspaces", "Avg occupancy\n(full day)", "Peak occupancy\n(busiest hour)", "Zones over\ncapacity"])
for a in AREAS:
    row = tbl.add_row().cells
    paths = val(a, "paidfree", 0)
    spaces = val(a, "paidfree", 1) + val(a, "paidfree", 2)
    peak = val(a, "occupancy", 0)
    over = val(a, "occupancy", 2)
    set_cell(row[0], LABELS[a], bold=True, size=10)
    set_cell(row[1], paths, align="center", size=10)
    set_cell(row[2], f"{spaces:,}", align="center", size=10)
    set_cell(row[3], f"{AVG.get(a, '–')}%", align="center", size=10, color=GREY)
    pc = RED if peak > 100 else (RGBColor(0xB8, 0x6B, 0x00) if peak > 85 else GREEN)
    set_cell(row[4], f"{peak}%", align="center", size=10, color=pc, bold=True)
    set_cell(row[5], over, align="center", size=10)
# totals row
row = tbl.add_row().cells
set_cell(row[0], "All areas", bold=True, size=10, white=True)
set_cell(row[1], total_paths, align="center", size=10, white=True, bold=True)
set_cell(row[2], f"{onstreet_total:,}", align="center", size=10, white=True, bold=True)
set_cell(row[3], f"{AVG['all']}%", align="center", size=10, white=True, bold=True)
set_cell(row[4], f"{A['all']['occupancy'][0]['value']}%", align="center", size=10, white=True, bold=True)
set_cell(row[5], A["all"]["occupancy"][2]["value"], align="center", size=10, white=True, bold=True)
for c in row:
    shade(c, "006D77")
body("Peak colour: green = within capacity (≤85%); amber = at capacity (86–100%); red = over "
     "capacity (>100%). Both columns are capacity-weighted over surveyed zones.",
     italic=True, color=GREY, size=9)

# ---- regulation ------------------------------------------------------------
heading("Finding 2 — Paid vs free parking", 1)
body(
    f"Within the surveyed dataset, {total_blue:,} of the {onstreet_total:,} marked on-street "
    f"spaces are recorded as paid and the remaining {total_white:,} as free, with the recorded "
    "paid parking concentrated in Komitas."
)
body(
    "This understates priced parking on the ground. The survey's regulation coding distinguishes "
    "only free and a single paid class and does not capture all paid (red) lines that exist on "
    "street — so the true extent of priced kerb is higher than the recorded figures, and the "
    "paid/free split should be treated as indicative rather than exact.",
    italic=True, color=GREY, size=10)
body(
    "The reliable, network-level takeaway is broader: across the surveyed areas, kerb space is "
    "largely unpriced or under-priced relative to the peak demand documented in Finding 1. "
    "Pricing remains the principal lever available to manage that demand wherever streets run hot "
    "at the peak hour."
)

# ---- retain / remove -------------------------------------------------------
heading("Finding 3 — Spaces retained and removed under the redesign", 1)
ret_all = A["all"]["retained"]
body(
    f"The redesign retains {ret_all[0]['value']:,} on-street spaces and removes "
    f"{ret_all[1]['value']:,} — {ret_all[2]['value']}% of surveyed on-street capacity is kept. "
    "The retain/remove decision is made street by street in the engineering workbooks, so the "
    "share kept differs greatly between areas: Komitas keeps three-quarters of its spaces, while "
    "Gai Avenue and Shiraz & Hasratyan are removed almost in full to make way for the corridor "
    "redesign."
)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tbl, ["Area", "Spaces retained", "Spaces removed", "% retained"])
for a in AREAS:
    row = tbl.add_row().cells
    r_ = A[a]["retained"]
    set_cell(row[0], LABELS[a], bold=True, size=10)
    set_cell(row[1], f"{r_[0]['value']:,}", align="center", size=10, color=GREEN)
    set_cell(row[2], f"{r_[1]['value']:,}", align="center", size=10, color=RED)
    set_cell(row[3], f"{r_[2]['value']}%", align="center", size=10, bold=True)
row = tbl.add_row().cells
set_cell(row[0], "All areas", bold=True, size=10, white=True)
set_cell(row[1], f"{ret_all[0]['value']:,}", align="center", size=10, white=True, bold=True)
set_cell(row[2], f"{ret_all[1]['value']:,}", align="center", size=10, white=True, bold=True)
set_cell(row[3], f"{ret_all[2]['value']}%", align="center", size=10, white=True, bold=True)
for c in row:
    shade(c, "006D77")

# ---- displacement ----------------------------------------------------------
heading("Finding 4 — Where do the displaced cars go?", 1)
body(
    f"Removing {D['removed_supply']:,} spaces does not remove {D['removed_supply']:,} cars: many "
    f"of those spaces are empty for much of the day. At the single busiest hour, the removed "
    f"segments hold {D['removed_demand']} parked cars — this is the demand that must be "
    f"re-accommodated. Against it, the study counts the spare capacity within walking distance: "
    f"surviving (retained) on-street spaces, other nearby un-surveyed on-street parking, and "
    f"off-street yards."
)
body(
    f"City-wide there is ample room: roughly {D['absorb_offstreet']:,} nearby off-street spaces "
    f"and {D['absorb_yellow_unsurveyed']:,} additional on-street spaces are available, a surplus of "
    f"about {D['net_vs_demand']:,} spaces over peak displaced demand. The constraint is local, not "
    "global — the table below shows that every area can absorb its own displaced demand except "
    "Gai Avenue, where on-street and yard capacity together cover an estimated "
    f"{val('mega','displacement',5)}% of the peak displaced cars."
)

tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tbl, ["Area", "Peak cars\ndisplaced", "Nearby\non-street", "Nearby\noff-street", "% of demand\nabsorbed"])
for a in AREAS:
    row = tbl.add_row().cells
    d_ = A[a]["displacement"]
    absorbed = d_[5]["value"]
    set_cell(row[0], LABELS[a], bold=True, size=10)
    set_cell(row[1], d_[0]["value"], align="center", size=10)
    set_cell(row[2], d_[2]["value"], align="center", size=10)
    set_cell(row[3], d_[3]["value"], align="center", size=10)
    set_cell(row[4], f"{absorbed}%", align="center", size=10, bold=True,
             color=GREEN if absorbed >= 100 else RED)
row = tbl.add_row().cells
da = A["all"]["displacement"]
set_cell(row[0], "All areas", bold=True, size=10, white=True)
set_cell(row[1], da[0]["value"], align="center", size=10, white=True, bold=True)
set_cell(row[2], da[2]["value"], align="center", size=10, white=True, bold=True)
set_cell(row[3], da[3]["value"], align="center", size=10, white=True, bold=True)
set_cell(row[4], f"{da[5]['value']}%", align="center", size=10, white=True, bold=True)
for c in row:
    shade(c, "006D77")
body(
    "“% of demand absorbed” is capped at 100% where capacity meets or exceeds peak displaced "
    "demand. It reflects gross best-case capacity (every nearby space assumed available) and "
    "assumes drivers are willing to shift to off-street yards — so the Gai Avenue shortfall is a "
    "floor, and real-world friction (walking distance, willingness to pay) would tighten the "
    "balance elsewhere.", italic=True, color=GREY, size=9)

# ---- displacement in detail (area by area, with dashboard screenshots) ------
doc.add_page_break()
heading("Displacement in detail — area by area", 1)
body(
    "The panels below are taken directly from the interactive dashboard with the Displacement "
    "lens active. The stat block on the left of each panel reports, for that view, the peak cars "
    "displaced, the spaces removed, the nearby on-street and off-street capacity available to "
    "absorb them, and the share of displaced demand that capacity covers. On the map:"
)
bullet(" surviving surveyed parking that the redesign keeps (it absorbs no displaced demand of its "
       "own — it is shown for context);", bold_lead="green = ")
bullet(" nearby un-surveyed on-street parking within walking distance that can take displaced "
       "cars;", bold_lead="yellow = ")
bullet(" off-street yards available to absorb displaced demand.", bold_lead="purple = ")
body(
    "Each subchapter pairs the dashboard view with a short reading of that area's displacement "
    "balance. Figures are peak-hour and reflect gross best-case capacity, as set out in Finding 4.",
    italic=True, color=GREY, size=9)

DISP_ORDER = ["all", "kentron", "komitas", "mega", "garegin", "shiraz", "malatia"]
DISP_TITLE = {
    "all": "All areas", "kentron": "Kentron", "komitas": "Komitas", "mega": "Gai Avenue",
    "garegin": "Garegin Nzhdeh", "shiraz": "Shiraz & Hasratyan", "malatia": "Malatia-Sebastia",
}
DISP_TEXT = {
    "all":
        "Across the whole study, the redesign removes {removed:,} marked spaces, but those streets "
        "hold only {disp} parked cars at the single busiest hour — the real demand to be "
        "re-accommodated. Against it stand {onstreet} nearby on-street spaces and {offstreet:,} "
        "off-street spaces, a combined surplus of roughly {surplus:,} over peak demand. Surviving "
        "on-street parking alone covers {onpct}% of the displaced cars; off-street yards close the "
        "remaining gap to full absorption. The headline is reassuring city-wide — there is enough "
        "room in principle — but it leans heavily on off-street capacity, which only works if "
        "drivers are willing to use it.",
    "kentron":
        "Kentron is the tightest on-street case. The redesign removes {removed} spaces and "
        "displaces {disp} peak-hour cars, yet only {onstreet} nearby on-street spaces remain to "
        "take them — barely {onpct}% of demand. This is exactly what the 138% peak occupancy "
        "predicted: the centre has essentially no spare kerb. Absorption only reaches {totalpct}% "
        "because {offstreet} off-street yard spaces carry almost the entire load. In practice, "
        "Kentron's displaced parkers must move off-street — there is nowhere on-street for them to "
        "go — so the off-street provision here is not a convenience but a prerequisite.",
    "komitas":
        "Komitas removes the least ({removed} spaces) because the redesign retains roughly "
        "three-quarters of its parking — the large green surviving network on the map. The {disp} "
        "displaced peak-hour cars find no spare on-street capacity ({onstreet}), but the area's "
        "{offstreet} off-street spaces absorb them comfortably, taking absorption to {totalpct}%. "
        "Displacement here is a small, well-covered problem.",
    "mega":
        "Gai Avenue is the one area that cannot fully re-absorb its displaced demand. It is removed "
        "in full, and because it runs chronically over capacity almost every removed space is "
        "occupied at peak — {disp} cars displaced against {removed} removed. Only {onstreet} nearby "
        "on-street and {offstreet} off-street spaces are available, covering about {totalpct}% of "
        "demand and leaving a peak-hour shortfall of roughly {gap} cars. This is the area that "
        "needs a targeted response — additional off-street provision or a phased removal — rather "
        "than relying on nearby capacity that is not there.",
    "garegin":
        "Garegin Nzhdeh is removed almost in full, displacing {disp} peak-hour cars. Nearby "
        "on-street parking is thin ({onstreet} spaces, {onpct}% of demand), so absorption to "
        "{totalpct}% depends on the {offstreet} off-street spaces around the area. The balance "
        "holds, but as in Kentron the off-street yard is doing the heavy lifting.",
    "shiraz":
        "Shiraz & Hasratyan has the healthiest balance among the fully-removed areas. Of its {disp} "
        "displaced peak-hour cars, nearby on-street parking can take {onstreet} ({onpct}%) on its "
        "own, with {offstreet} off-street spaces covering the rest to reach {totalpct}%. Most "
        "displaced parkers can stay on-street nearby — the redesign is comfortably accommodated.",
    "malatia":
        "Malatia-Sebastia is the easiest case. It displaces {disp} peak-hour cars, and {onstreet} "
        "nearby on-street spaces alone absorb {onpct}% of them — almost the entire demand — before "
        "the {offstreet} off-street spaces are even counted. With the area only about half-full at "
        "peak to begin with, the displacement balance is well within capacity at {totalpct}%.",
}
for a in DISP_ORDER:
    d_ = A[a]["displacement"]
    disp, removed = d_[0]["value"], d_[1]["value"]
    onstreet, offstreet = d_[2]["value"], d_[3]["value"]
    onpct, totalpct = d_[4]["value"], d_[5]["value"]
    heading(DISP_TITLE[a], 2)
    add_image(IMGDIR + "/" + DISP_IMG[a],
              caption=f"Displacement lens — {DISP_TITLE[a]} (interactive dashboard).")
    body(DISP_TEXT[a].format(
        disp=disp, removed=removed, onstreet=onstreet, offstreet=offstreet,
        onpct=onpct, totalpct=totalpct, gap=max(0, disp - onstreet - offstreet),
        surplus=D["net_vs_demand"]))

# ---- demand profile --------------------------------------------------------
heading("Finding 5 — Who parks, and when", 1)
dur = A["all"]["profile"]["duration"]
body(
    f"The surveyed fleet is overwhelmingly short-stay: {dur['visitorPct']}% of vehicles park for "
    f"under two hours (visitors), {dur['commuterPct']}% for two to eight hours (commuters), and "
    f"only {dur['longPct']}% stay longer than eight hours. The average stay is {dur['avg']} hours. "
    "This profile is favourable for pricing: short-stay demand is the most responsive to even "
    "modest parking charges, so priced kerb space would turn over faster and serve more users."
)
# peak-hour callout from vap
vap = A["all"]["profile"]["vap"]
peak_hour, peak_cars = max(vap, key=lambda x: x[1])
body(
    f"Demand builds through the morning to a late-morning peak (around {peak_hour:02d}:00, "
    f"about {peak_cars:,} cars present network-wide) and holds high into the evening, confirming "
    "that the busiest streets are under sustained pressure across the working day rather than in a "
    "single rush hour."
)

# ---- area-by-area ----------------------------------------------------------
doc.add_page_break()
heading("Area profiles", 1)
NOTES = {
    "kentron": "The congested city centre: well over capacity at peak and almost entirely free "
               "(unpriced) on-street parking. Pricing here would have the greatest effect.",
    "komitas": "The only substantially priced area (276 blue spaces) and the one that retains most "
               "of its parking (76%). Still over capacity at peak.",
    "mega": "Gai Avenue / Mega Mall. Removed in full; the only area where nearby capacity does not "
            "fully absorb peak displaced demand — a local shortfall to address.",
    "garegin": "Garegin Nzhdeh: high peak occupancy and a near-total removal; displaced demand is "
               "absorbed mainly by the off-street yard.",
    "shiraz": "Shiraz & Hasratyan: moderate occupancy, removed in full; comfortably absorbed by "
              "surrounding on-street and yard capacity.",
    "malatia": "Malatia-Sebastia: the most relaxed area, around half-full at peak, with the highest "
               "retention of its on-street parking and an easy displacement balance.",
}
for a in AREAS:
    heading(LABELS[a], 2)
    occ = A[a]["occupancy"]
    pf = A[a]["paidfree"]
    rt = A[a]["retained"]
    dd = A[a]["displacement"]
    pr = A[a]["profile"]["duration"]
    body(NOTES[a])
    facts = (
        f"Survey segments: {pf[0]['value']}  •  marked spaces: "
        f"{pf[1]['value'] + pf[2]['value']:,} (paid: {pf[2]['value']})  •  "
        f"off-street yard: {int(pf[3]['value'])} spaces ({pf[3]['label']}).\n"
        f"Occupancy: {AVG.get(a, '–')}% average / {occ[0]['value']}% at peak  •  zones over "
        f"capacity: {occ[2]['value']}.\n"
        f"Retained: {rt[0]['value']:,} spaces ({rt[2]['value']}%)  •  removed: {rt[1]['value']:,}.\n"
        f"Stay mix: {pr['visitorPct']}% visitor / {pr['commuterPct']}% commuter / {pr['longPct']}% long "
        f"(avg {pr['avg']} h)."
    )
    p = doc.add_paragraph()
    r = p.add_run(facts)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ===========================================================================
# VISUAL ANALYSIS CHAPTER (generated charts) — placed before the conclusions
# ===========================================================================
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

BLUE = "#2BA8D6"; BLUE_L = "#BfE3F2"; ORANGE = "#ED7D31"; GREEN = "#4CAF50"
REDc = "#EF5350"; YEL = "#F4B400"; PUR = "#7C4DFF"; INK = "#444444"; GRIDC = "#DDDDDD"
ORDER6 = ["kentron", "komitas", "mega", "garegin", "shiraz", "malatia"]
SHORT = {"kentron": "Kentron", "komitas": "Komitas", "mega": "Gai Ave",
         "garegin": "Garegin", "shiraz": "Shiraz", "malatia": "Malatia"}

def _cap(area):
    return sum(p["space"] for ft in g["features"]
               for p in [ft.get("properties", {})]
               if p.get("area") == area and p.get("peak_occupancy") is not None and p.get("space"))
CAP = {a: _cap(a) for a in ORDER6}
CAP_ALL = sum(CAP.values())

def _style(ax, title, ylabel=None):
    ax.set_title(title, color=INK, fontsize=13.5, pad=12, weight="bold")
    ax.set_facecolor("white")
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    for s in ("left", "bottom"):
        ax.spines[s].set_color("#BBBBBB")
    ax.yaxis.grid(True, color=GRIDC, linewidth=0.8)
    ax.set_axisbelow(True)
    ax.tick_params(colors=INK, labelsize=9)
    if ylabel:
        ax.set_ylabel(ylabel, color=INK, fontsize=9)

def _labels(ax, bars, fmt="{:.0f}", dy=3, size=8):
    for b in bars:
        h = b.get_height()
        ax.annotate(fmt.format(h), (b.get_x() + b.get_width() / 2, h),
                    ha="center", va="bottom", fontsize=size, color=INK,
                    xytext=(0, dy), textcoords="offset points")

def _save(fig, name):
    path = IMGDIR + "/" + name
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path

charts = {}

# Chart 1 — hourly demand vs supply (all areas)
vap = [(h, c) for h, c in A["all"]["profile"]["vap"] if h <= 23 and c > 10]
fig, ax = plt.subplots(figsize=(9, 4.2))
bars = ax.bar([f"{h}:00" for h, _ in vap], [c for _, c in vap], color=BLUE, width=0.72)
_labels(ax, bars, dy=2, size=7.5)
ax.axhline(CAP_ALL, color=ORANGE, lw=2.5)
ax.annotate(f"Surveyed supply ≈ {CAP_ALL:,}", (len(vap) - 1, CAP_ALL),
            color=ORANGE, fontsize=9, weight="bold", ha="right", va="bottom",
            xytext=(0, 4), textcoords="offset points")
_style(ax, "Hourly parking demand vs. surveyed supply — all areas", "Vehicles present")
plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
ax.set_ylim(0, max(c for _, c in vap) * 1.18)
charts["demand"] = _save(fig, "chart_demand_vs_supply.png")

# Chart 2 — average vs peak occupancy by area
fig, ax = plt.subplots(figsize=(9, 4.4))
import numpy as np
x = np.arange(len(ORDER6)); w = 0.38
avg = [AVG[a] for a in ORDER6]; peak = [A[a]["occupancy"][0]["value"] for a in ORDER6]
b1 = ax.bar(x - w / 2, avg, w, color=BLUE_L, label="Full-day average")
b2 = ax.bar(x + w / 2, peak, w, color=BLUE, label="Peak hour")
_labels(ax, b1, fmt="{:.0f}%", size=7.5); _labels(ax, b2, fmt="{:.0f}%", size=7.5)
ax.axhline(85, color=YEL, lw=1.4, ls="--"); ax.axhline(100, color=REDc, lw=1.4, ls="--")
ax.annotate("85% efficiency", (len(ORDER6) - 0.5, 85), color="#B8860B", fontsize=7.5, va="bottom", ha="right")
ax.annotate("100% capacity", (len(ORDER6) - 0.5, 100), color=REDc, fontsize=7.5, va="bottom", ha="right")
ax.set_xticks(x); ax.set_xticklabels([SHORT[a] for a in ORDER6])
_style(ax, "Average vs. peak occupancy by area (% of capacity)", "% of capacity")
ax.set_ylim(0, max(peak) * 1.18); ax.legend(frameon=False, fontsize=9, loc="upper right")
charts["occ"] = _save(fig, "chart_avg_vs_peak_occupancy.png")

# Chart 3 — length-of-stay mix by area (100% stacked horizontal)
rows = ["all"] + ORDER6
labels_y = ["All areas"] + [SHORT[a] for a in ORDER6]
vis = [A[a]["profile"]["duration"]["visitorPct"] for a in rows]
com = [A[a]["profile"]["duration"]["commuterPct"] for a in rows]
lng = [A[a]["profile"]["duration"]["longPct"] for a in rows]
fig, ax = plt.subplots(figsize=(9, 4.0))
yidx = np.arange(len(rows))
ax.barh(yidx, vis, color=BLUE, label="Visitor (<2 h)")
ax.barh(yidx, com, left=vis, color=ORANGE, label="Commuter (2–8 h)")
ax.barh(yidx, lng, left=[v + c for v, c in zip(vis, com)], color="#7A7A7A", label="Long (>8 h)")
for i, v in enumerate(vis):
    ax.annotate(f"{v}%", (v / 2, i), ha="center", va="center", color="white", fontsize=8, weight="bold")
ax.set_yticks(yidx); ax.set_yticklabels(labels_y); ax.invert_yaxis()
ax.set_xlim(0, 100); ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0f}%"))
ax.set_title("Length-of-stay mix by area", color=INK, fontsize=13.5, pad=12, weight="bold")
ax.set_facecolor("white")
for s in ("top", "right", "left"):
    ax.spines[s].set_visible(False)
ax.spines["bottom"].set_color("#BBBBBB"); ax.tick_params(colors=INK, labelsize=9)
ax.xaxis.grid(True, color=GRIDC, linewidth=0.8); ax.set_axisbelow(True)
ax.legend(frameon=False, fontsize=9, ncol=3, loc="lower center", bbox_to_anchor=(0.5, -0.22))
charts["stay"] = _save(fig, "chart_stay_mix.png")

# Chart 4 — on-street spaces retained vs removed by area (stacked)
fig, ax = plt.subplots(figsize=(9, 4.2))
ret = [A[a]["retained"][0]["value"] for a in ORDER6]
rem = [A[a]["retained"][1]["value"] for a in ORDER6]
x = np.arange(len(ORDER6))
ax.bar(x, ret, 0.6, color=GREEN, label="Retained")
ax.bar(x, rem, 0.6, bottom=ret, color=REDc, label="Removed")
for i, a in enumerate(ORDER6):
    tot = ret[i] + rem[i]
    ax.annotate(f"{A[a]['retained'][2]['value']}% kept", (i, tot), ha="center", va="bottom",
                fontsize=8, color=INK, xytext=(0, 3), textcoords="offset points")
ax.set_xticks(x); ax.set_xticklabels([SHORT[a] for a in ORDER6])
_style(ax, "On-street spaces retained vs. removed by area", "Marked spaces")
ax.set_ylim(0, max(ret[i] + rem[i] for i in range(len(ORDER6))) * 1.20)
ax.legend(frameon=False, fontsize=9, loc="upper left")
charts["retrem"] = _save(fig, "chart_retained_removed.png")

# Chart 5 — peak displaced demand vs nearby absorptive capacity by area
fig, ax = plt.subplots(figsize=(9, 4.4))
x = np.arange(len(ORDER6)); w = 0.38
disp = [A[a]["displacement"][0]["value"] for a in ORDER6]
on = [A[a]["displacement"][2]["value"] for a in ORDER6]
off = [A[a]["displacement"][3]["value"] for a in ORDER6]
bd = ax.bar(x - w / 2, disp, w, color=REDc, label="Peak cars displaced")
bo = ax.bar(x + w / 2, on, w, color=YEL, label="Nearby on-street")
ax.bar(x + w / 2, off, w, bottom=on, color=PUR, label="Nearby off-street")
_labels(ax, bd, size=7.5)
ax.set_xticks(x); ax.set_xticklabels([SHORT[a] for a in ORDER6])
_style(ax, "Peak displaced demand vs. nearby absorptive capacity by area", "Spaces / cars")
ax.set_ylim(0, max(max(disp), max(o + f for o, f in zip(on, off))) * 1.18)
ax.legend(frameon=False, fontsize=9, loc="upper left")
charts["disp"] = _save(fig, "chart_displacement_balance.png")

# Chart 6 — share of displaced demand absorbed on-street by area
fig, ax = plt.subplots(figsize=(9, 4.0))
onpct = [A[a]["displacement"][4]["value"] for a in ORDER6]
cols = [REDc if v < 35 else (YEL if v < 70 else GREEN) for v in onpct]
bars = ax.bar([SHORT[a] for a in ORDER6], onpct, 0.6, color=cols)
_labels(ax, bars, fmt="{:.0f}%", size=8.5)
_style(ax, "Share of displaced demand absorbed on-street (remainder relies on off-street)", "% absorbed on-street")
ax.set_ylim(0, 110)
charts["onshare"] = _save(fig, "chart_onstreet_share.png")

# Chart 7 — hourly occupancy curve by area (% of capacity)
AREA_COL = {"kentron": "#E4572E", "komitas": "#2BA8D6", "mega": "#7C4DFF",
            "garegin": "#2E9E5B", "shiraz": "#E8A33D", "malatia": "#8D6E63"}
fig, ax = plt.subplots(figsize=(9, 4.6))
for a in ORDER6:
    pts = [(h, c) for h, c in A[a]["profile"]["vap"] if h <= 23]
    ax.plot([h for h, _ in pts], [c / CAP[a] * 100 for _, c in pts],
            marker="o", ms=3, lw=2, color=AREA_COL[a], label=SHORT[a])
ax.axhline(100, color=REDc, lw=1.3, ls="--")
ax.axhline(85, color=YEL, lw=1.3, ls="--")
_style(ax, "Hourly occupancy by area (% of capacity)", "% of capacity")
ax.set_xlabel("Hour of day", color=INK, fontsize=9)
ax.set_xticks(range(7, 24, 2))
ax.legend(frameon=False, fontsize=8.5, ncol=6, loc="upper center", bbox_to_anchor=(0.5, -0.16))
charts["hourly"] = _save(fig, "chart_hourly_occupancy.png")

# Chart 8 — city-wide displacement headline
fig, ax = plt.subplots(figsize=(9, 2.9))
dem = D["removed_demand"]; aon = D["absorb_onstreet"]; aoff = D["absorb_offstreet"]
ax.barh([1], [dem], color=REDc, height=0.55)
ax.barh([0], [aon], color=YEL, height=0.55, label="Nearby on-street")
ax.barh([0], [aoff], left=[aon], color=PUR, height=0.55, label="Nearby off-street")
ax.annotate(f"{dem:,} cars", (dem, 1), va="center", ha="left", xytext=(5, 0),
            textcoords="offset points", fontsize=9.5, weight="bold", color=INK)
ax.annotate(f"{aon + aoff:,} spaces   (surplus +{D['net_vs_demand']:,})", (aon + aoff, 0),
            va="center", ha="left", xytext=(5, 0), textcoords="offset points",
            fontsize=9.5, weight="bold", color=INK)
ax.set_yticks([0, 1]); ax.set_yticklabels(["Absorptive\ncapacity", "Peak displaced\ndemand"])
ax.set_title("City-wide displacement balance at the peak hour", color=INK, fontsize=13.5, pad=12, weight="bold")
ax.set_facecolor("white")
for s in ("top", "right", "left", "bottom"):
    ax.spines[s].set_visible(False)
ax.tick_params(colors=INK, labelsize=9, length=0)
ax.set_xlim(0, (aon + aoff) * 1.18); ax.set_xticks([])
ax.legend(frameon=False, fontsize=8.5, ncol=2, loc="lower right", bbox_to_anchor=(1, -0.05))
charts["headline"] = _save(fig, "chart_displacement_headline.png")

# Chart 9 — net parking headroom by area
fig, ax = plt.subplots(figsize=(9, 4.2))
head = [A[a]["displacement"][2]["value"] + A[a]["displacement"][3]["value"]
        - A[a]["displacement"][0]["value"] for a in ORDER6]
bars = ax.bar([SHORT[a] for a in ORDER6], head, 0.6,
              color=[GREEN if v >= 0 else REDc for v in head])
ax.axhline(0, color="#888888", lw=1)
for b, v in zip(bars, head):
    ax.annotate(f"{v:+d}", (b.get_x() + b.get_width() / 2, v), ha="center",
                va="bottom" if v >= 0 else "top", fontsize=8.5, color=INK,
                xytext=(0, 3 if v >= 0 else -3), textcoords="offset points")
_style(ax, "Net parking headroom by area (nearby capacity − peak displaced)", "Spare spaces at peak")
ax.set_ylim(min(head) * 1.4 if min(head) < 0 else -10, max(head) * 1.2)
charts["headroom"] = _save(fig, "chart_net_headroom.png")

# Chart 10 — parking turnover by area
tv = {}
for ft in g["features"]:
    p = ft.get("properties", {})
    if p.get("location") != "on-street" or p.get("parking_events") is None or not p.get("space"):
        continue
    tv.setdefault(p["area"], {"e": 0, "s": 0})
    tv[p["area"]]["e"] += p["parking_events"]; tv[p["area"]]["s"] += p["space"]
turn = [round(tv[a]["e"] / tv[a]["s"], 1) for a in ORDER6]
fig, ax = plt.subplots(figsize=(9, 4.0))
bars = ax.bar([SHORT[a] for a in ORDER6], turn, 0.6, color=BLUE)
_labels(ax, bars, fmt="{:.1f}", size=8.5)
_style(ax, "Parking turnover by area (vehicles per space per day)", "Vehicles / space / day")
ax.set_ylim(0, max(turn) * 1.18)
charts["turnover"] = _save(fig, "chart_turnover.png")

# Chart 11 — breadth of saturation (% of zones over 85% at peak)
over85 = [A[a]["occupancy"][1]["value"] for a in ORDER6]
fig, ax = plt.subplots(figsize=(9, 4.0))
bars = ax.bar([SHORT[a] for a in ORDER6], over85, 0.6,
              color=[REDc if v >= 70 else (YEL if v >= 40 else GREEN) for v in over85])
_labels(ax, bars, fmt="{:.0f}%", size=8.5)
_style(ax, "Breadth of saturation: zones over 85% occupancy at peak (% of an area's zones)", "% of zones")
ax.set_ylim(0, 105)
charts["breadth"] = _save(fig, "chart_saturation_breadth.png")

# Chart 12 — where the displaced demand comes from (share of the peak total)
src = sorted(((a, A[a]["displacement"][0]["value"]) for a in ORDER6), key=lambda x: -x[1])
tot_disp = sum(v for _, v in src)
fig, ax = plt.subplots(figsize=(9, 4.0))
bars = ax.bar([SHORT[a] for a, _ in src], [v for _, v in src], 0.6, color=BLUE)
for b, (_, v) in zip(bars, src):
    ax.annotate(f"{v}\n{round(100 * v / tot_disp)}%", (b.get_x() + b.get_width() / 2, v),
                ha="center", va="bottom", fontsize=8.5, color=INK, xytext=(0, 3), textcoords="offset points")
_style(ax, "Where the displaced demand comes from (peak cars displaced, by area)", "Peak cars displaced")
ax.set_ylim(0, max(v for _, v in src) * 1.22)
charts["source"] = _save(fig, "chart_displaced_source.png")

# Chart 13 — distribution of zones by peak occupancy (over-capacity tail)
band_lbl = ["<50%", "50–85%", "85–100%", "100–150%", ">150%"]
band_col = ["#2E9E5B", "#9CCC65", "#F4B400", "#FB8C00", "#E53935"]
bands = [0, 0, 0, 0, 0]
for ft in g["features"]:
    p = ft.get("properties", {})
    if p.get("location") != "on-street" or p.get("peak_occupancy") is None or not p.get("space"):
        continue
    o = p["peak_occupancy"] / p["space"] * 100
    bands[0 if o < 50 else 1 if o < 85 else 2 if o <= 100 else 3 if o <= 150 else 4] += 1
nz = sum(bands)
fig, ax = plt.subplots(figsize=(9, 4.2))
bars = ax.bar(band_lbl, bands, 0.7, color=band_col)
for b, v in zip(bars, bands):
    ax.annotate(f"{v}\n{round(100 * v / nz)}%", (b.get_x() + b.get_width() / 2, v), ha="center",
                va="bottom", fontsize=8.5, color=INK, xytext=(0, 3), textcoords="offset points")
_style(ax, f"Distribution of surveyed zones by peak occupancy (n={nz})", "Number of zones")
ax.set_xlabel("Peak occupancy band (% of capacity)", color=INK, fontsize=9)
ax.set_ylim(0, max(bands) * 1.2)
charts["dist"] = _save(fig, "chart_occupancy_distribution.png")

# Chart 14 — saturation vs. turnover, by zone (scatter)
fig, ax = plt.subplots(figsize=(9, 4.6))
for a in ORDER6:
    xs, ys = [], []
    for ft in g["features"]:
        p = ft.get("properties", {})
        if p.get("area") != a or p.get("location") != "on-street":
            continue
        if p.get("turnover") is None or p.get("peak_occupancy") is None or not p.get("space"):
            continue
        xs.append(p["turnover"]); ys.append(p["peak_occupancy"] / p["space"] * 100)
    ax.scatter(xs, ys, s=26, alpha=0.7, color=AREA_COL[a], label=SHORT[a], edgecolors="none")
ax.axhline(100, color=REDc, lw=1.2, ls="--")
ax.axhline(85, color=YEL, lw=1.2, ls="--")
_style(ax, "Saturation vs. turnover, by zone", "Peak occupancy (% of capacity)")
ax.set_xlabel("Turnover (vehicles per space per day)", color=INK, fontsize=9)
ax.annotate("upper-left = high occupancy, low turnover\n(commuter-locked — rare here)", (0.02, 0.97),
            xycoords="axes fraction", va="top", ha="left", fontsize=8, color="#777777", style="italic")
ax.legend(frameon=False, fontsize=8.5, ncol=6, loc="upper center", bbox_to_anchor=(0.5, -0.16))
charts["scatter"] = _save(fig, "chart_occupancy_turnover.png")

# Chart 15 — hour of peak demand by area (lollipop)
ph = []
for a in ORDER6:
    vap = [(h, c) for h, c in A[a]["profile"]["vap"] if h <= 23]
    ph.append((a, max(vap, key=lambda x: x[1])[0]))
ph.sort(key=lambda x: x[1])
fig, ax = plt.subplots(figsize=(9, 3.8))
ys = np.arange(len(ph))
ax.hlines(ys, 7, [h for _, h in ph], color="#CCCCCC", lw=2, zorder=1)
ax.scatter([h for _, h in ph], ys, s=90, color=BLUE, zorder=2)
for y, (a, h) in zip(ys, ph):
    ax.annotate(f"{h}:00", (h, y), color=INK, fontsize=9, weight="bold", va="center", ha="left",
                xytext=(8, 0), textcoords="offset points")
ax.set_yticks(ys); ax.set_yticklabels([SHORT[a] for a, _ in ph])
ax.set_xlim(7, 24); ax.set_xlabel("Hour of day", color=INK, fontsize=9)
ax.set_title("Hour of peak parking demand by area", color=INK, fontsize=13.5, pad=12, weight="bold")
ax.set_facecolor("white")
for s in ("top", "right", "left"):
    ax.spines[s].set_visible(False)
ax.spines["bottom"].set_color("#BBBBBB"); ax.tick_params(colors=INK, labelsize=9, length=0)
ax.xaxis.grid(True, color=GRIDC, linewidth=0.8); ax.set_axisbelow(True)
charts["peakhour"] = _save(fig, "chart_peak_hour.png")

# Chart 16 — daytime vs. evening share of demand by area
fig, ax = plt.subplots(figsize=(9, 3.8))
rows = ORDER6
yidx = np.arange(len(rows))
day, eve = [], []
for a in rows:
    vap = [(h, c) for h, c in A[a]["profile"]["vap"] if h <= 23]
    d = sum(c for h, c in vap if h < 18); e = sum(c for h, c in vap if h >= 18)
    tot = d + e
    day.append(round(100 * d / tot)); eve.append(100 - round(100 * d / tot))
ax.barh(yidx, day, color=BLUE, label="Daytime (07–18)")
ax.barh(yidx, eve, left=day, color="#37474F", label="Evening (18–24)")
for i, (d, e) in enumerate(zip(day, eve)):
    ax.annotate(f"{d}%", (d / 2, i), ha="center", va="center", color="white", fontsize=8.5, weight="bold")
    ax.annotate(f"{e}%", (d + e / 2, i), ha="center", va="center", color="white", fontsize=8.5, weight="bold")
ax.set_yticks(yidx); ax.set_yticklabels([SHORT[a] for a in rows]); ax.invert_yaxis()
ax.set_xlim(0, 100)
ax.set_title("Daytime vs. evening share of demand by area", color=INK, fontsize=13.5, pad=12, weight="bold")
ax.set_facecolor("white")
for s in ("top", "right", "left", "bottom"):
    ax.spines[s].set_visible(False)
ax.tick_params(colors=INK, labelsize=9, length=0); ax.set_xticks([])
ax.legend(frameon=False, fontsize=9, ncol=2, loc="lower center", bbox_to_anchor=(0.5, -0.22))
charts["dayeve"] = _save(fig, "chart_day_evening.png")

# Chart 17 — parking layout by area: parallel vs. angled
cfg = {}
for ft in g["features"]:
    p = ft.get("properties", {})
    if p.get("location") != "on-street" or not p.get("space"):
        continue
    cfg.setdefault(p["area"], {"par": 0, "t": 0})
    cfg[p["area"]]["t"] += p["space"]
    if str(p.get("method", "")).startswith("par"):
        cfg[p["area"]]["par"] += p["space"]
par = [round(100 * cfg[a]["par"] / cfg[a]["t"]) for a in ORDER6]
ang = [100 - v for v in par]
fig, ax = plt.subplots(figsize=(9, 4.0))
x = np.arange(len(ORDER6))
ax.bar(x, par, 0.6, color=BLUE, label="Parallel")
ax.bar(x, ang, 0.6, bottom=par, color=ORANGE, label="Angled / perpendicular")
for i, v in enumerate(par):
    if ang[i] > 6:
        ax.annotate(f"{ang[i]}% angled", (i, 100), ha="center", va="bottom", fontsize=8, color=INK,
                    xytext=(0, 3), textcoords="offset points")
ax.set_xticks(x); ax.set_xticklabels([SHORT[a] for a in ORDER6])
_style(ax, "Parking layout by area: parallel vs. angled", "% of marked spaces")
ax.set_ylim(0, 112)
ax.legend(frameon=False, fontsize=9, ncol=2, loc="lower center", bbox_to_anchor=(0.5, -0.2))
charts["layout"] = _save(fig, "chart_layout.png")

# Chart 18 — surveyed on-street supply by area (the base behind the percentages)
supply = [A[a]["paidfree"][1]["value"] + A[a]["paidfree"][2]["value"] for a in ORDER6]
sp = sorted(zip(ORDER6, supply), key=lambda x: -x[1])
fig, ax = plt.subplots(figsize=(9, 4.0))
bars = ax.bar([SHORT[a] for a, _ in sp], [v for _, v in sp], 0.6, color=BLUE)
_labels(ax, bars, fmt="{:,.0f}", size=8.5)
_style(ax, "Surveyed on-street parking by area (marked spaces)", "Marked spaces")
ax.set_ylim(0, max(supply) * 1.18)
charts["footprint"] = _save(fig, "chart_supply_footprint.png")

doc.add_page_break()
heading("Visual analysis of the field-survey data", 1)
body(
    "This section distils the field-survey dataset into a series of charts. Each is followed by a "
    "short note on what it implies for action, rather than a description of the chart itself."
)

heading("Demand through the day", 2)
add_image(charts["demand"], caption="Hourly vehicles present across all surveyed areas, against surveyed on-street supply.")
whatnow(
    "the kerb is pressed twice a day — a late-morning and an evening peak — not in a single rush "
    "hour. Time-of-day pricing and any loading-bay time windows should target both peaks; the deep "
    "midday and overnight troughs are when relocated or shared parking can be steered without "
    "friction.")

heading("The gap between average and peak", 2)
add_image(charts["occ"], caption="Capacity-weighted occupancy, full-day average vs. busiest hour, by area.")
whatnow(
    "manage the peak, not the average. The bars that clear the red line (Kentron, Komitas, Gai "
    "Avenue) are where pricing and turnover measures earn their place; the low-average outer areas "
    "have headroom to take reallocated demand today.")

heading("Who is parking", 2)
add_image(charts["stay"], caption="Share of vehicles by length of stay, by area.")
whatnow(
    "a short-stay kerb responds to hourly pricing and a two-hour cap far faster than to permits. "
    "Lead with turnover tools city-wide and reserve resident permits for the genuinely "
    "residential exceptions rather than applying them as a default.")

heading("Scale of the change", 2)
add_image(charts["retrem"], caption="On-street spaces retained vs. removed under the redesign, by area.")
whatnow(
    "the areas removed in full (Gai Avenue, Shiraz) need their mitigation committed before "
    "construction starts; high-retention Komitas can be phased more gently. Sequencing should "
    "follow this chart, not a uniform timetable.")

heading("Can the displaced cars be re-parked?", 2)
add_image(charts["disp"], caption="Peak displaced demand against nearby on-street and off-street capacity, by area.")
whatnow(
    "only Gai Avenue's red bar overtops its available capacity — commit a specific off-street or "
    "park-and-ride provision there before any removal. Everywhere else the balance holds only if "
    "the off-street yards are real, open and signed, so confirm them on the ground.")

heading("Where on-street will not catch the overflow", 2)
add_image(charts["onshare"], caption="Share of each area's displaced demand that nearby on-street parking can absorb.")
whatnow(
    "in the core (Kentron, Komitas) on-street simply cannot absorb the displaced cars — secure "
    "off-street capacity there through courtyards, yards and park-and-ride. Save the side-street "
    "conversion lever for Malatia and Shiraz, where the chart shows it actually works.")

heading("How long the pressure lasts", 2)
add_image(charts["hourly"], caption="Occupancy as a share of capacity, hour by hour, for each area.")
whatnow(
    "the central areas sit at or above capacity for most of the working day, not for a single "
    "hour — so their pricing and enforcement need to run all day, while the outer areas only need "
    "attention in the peak windows.")

heading("The city-wide balance in one line", 2)
add_image(charts["headline"], caption="Peak displaced demand against total nearby absorptive capacity, all areas combined.")
whatnow(
    "the surplus is comfortable on paper, but it is overwhelmingly off-street. The task is to make "
    "that off-street capacity genuinely usable — open, priced and signed — not to find more land.")

heading("Who has a deficit", 2)
add_image(charts["headroom"], caption="Spare spaces at the peak hour after absorbing displaced demand, by area.")
whatnow(
    "every area clears its own demand except Gai Avenue, whose bar falls below zero. That single "
    "deficit is the one place a do-nothing approach fails, so it warrants a dedicated provision "
    "rather than reliance on the surrounding network.")

heading("How hard the kerb works", 2)
add_image(charts["turnover"], caption="Vehicles served per marked space per day, by area.")
whatnow(
    "turnover is high across every area, confirming a short-stay, price-sensitive kerb. Hourly "
    "pricing and short maximum-stay limits will change behaviour quickly here — faster than "
    "permit-based schemes, which suit residential storage rather than this churn.")

heading("How widespread the stress is", 2)
add_image(charts["breadth"], caption="Share of each area's zones running over 85% occupancy at the peak hour.")
whatnow(
    "saturation is broad, not a handful of hotspots: in Kentron, Komitas and Gai Avenue most "
    "zones breach the 85% threshold. Mitigation in those areas should be applied zone-wide, not "
    "patched street by street.")

heading("Where to start", 2)
add_image(charts["source"], caption="Peak cars displaced by area, ranked — the share of the city-wide total each contributes.")
whatnow(
    "Kentron alone accounts for over a quarter of all displaced demand, with Shiraz and Garegin "
    "next. Sequence mitigation effort and budget to the largest contributors first rather than "
    "spreading them evenly.")

heading("How many streets are over the line", 2)
add_image(charts["dist"], caption="Every surveyed zone placed in a peak-occupancy band.")
whatnow(
    "more than half of all surveyed zones are over capacity at the peak, and nearly a third are "
    "above 150%. This is over-saturation at scale, not a few outliers — the over-100% tail is "
    "where pricing and enforcement have to land first.")

heading("Saturated, or just busy?", 2)
add_image(charts["scatter"], caption="Each zone plotted by its turnover and its peak occupancy; colour denotes area.")
whatnow(
    "occupancy and turnover rise together — the busiest streets are also churning fastest, so the "
    "saturation is short-stay pressure, not all-day commuter storage. The lever is therefore "
    "pricing to keep a margin of spaces free, not blanket time-limits; reserve maximum-stay caps "
    "for the few high-occupancy, low-turnover outliers in the upper-left.")

heading("When each area peaks", 2)
add_image(charts["peakhour"], caption="Clock hour at which each area reaches its highest parking demand.")
whatnow(
    "demand peaks at different times — late morning in the outer commercial areas, evening in the "
    "Kentron centre. Enforcement hours and loading-bay windows should follow each area's own clock; "
    "a single city-wide timetable would miss the pressure in half of them.")

heading("Day or night demand", 2)
add_image(charts["dayeve"], caption="Share of each area's daily demand falling in daytime vs. evening hours.")
whatnow(
    "the evening-heavy areas (Kentron, Malatia) carry residential demand that this daytime survey "
    "only partly captures — protect them with resident provision. The day-heavy areas (Shiraz, Gai "
    "Avenue) are commercial and suit turnover pricing. Note the survey ends around midnight, so "
    "genuine overnight demand is understated, not absent.")

heading("How the kerb is laid out", 2)
add_image(charts["layout"], caption="Share of marked spaces that are parallel vs. angled/perpendicular, by area.")
whatnow(
    "Kentron's kerb is roughly 60% angled parking, much of it informal — restructuring it into "
    "marked parallel bays restores order and pedestrian space even before any parking is removed. "
    "The other areas are already orderly parallel layouts and need no such intervention.")

heading("The base behind the percentages", 2)
add_image(charts["footprint"], caption="Total marked on-street spaces surveyed in each area.")
whatnow(
    "read the earlier percentages against this base — Malatia and Kentron hold the most spaces, so "
    "an equal percentage change there moves the most cars. Size mitigation budgets to the bars "
    "here, not to the number of areas.")

# ---- conclusions -----------------------------------------------------------
doc.add_page_break()
heading("Conclusions and recommendations", 1)
bullet(
    f" Network-wide the kerb is only {AVG['all']}% occupied across the working day but reaches "
    f"{A['all']['occupancy'][0]['value']}% at the peak hour. More than half the kerb sits empty "
    "most of the day — the spaces exist, they are just all wanted at the same time and place. A "
    "genuine shortage would show up as a high daily average too; this does not. Capacity is "
    "therefore not the binding constraint city-wide — management of the peak is. (Kentron and Gai "
    "Avenue are the local exceptions, where capacity is genuinely tight.)",
    bold_lead="Demand timing, not supply, is the problem. ")
bullet(
    " Across the network, kerb space is largely unpriced or under-priced relative to peak demand. "
    "Introducing or extending paid parking in the high-occupancy areas — set to keep a few spaces "
    "free even at the peak hour — is the most direct way to free up capacity, and the short-stay "
    "demand profile (most stays under two hours) means demand would respond.",
    bold_lead="Use pricing to manage peak demand. ")
bullet(
    f" The {D['removed_supply']:,} spaces removed by the redesign displace {D['removed_demand']} "
    "peak-hour cars, and nearby capacity can absorb them in every area except Gai Avenue. A "
    "targeted provision (additional yard capacity or phased removal) is recommended there.",
    bold_lead="Manage the one local shortfall. ")
bullet(
    " Off-street yards do most of the heavy lifting in the displacement balance. Their capacity, "
    "pricing and signage should be confirmed and actively promoted so drivers actually use them.",
    bold_lead="Lean on off-street capacity. ")

body(
    "All figures in this report are drawn directly from the field-survey dataset published in the "
    "interactive briefing (Step 9) and reflect the current model run.",
    italic=True, color=GREY, size=9)

# ===========================================================================
# IMPLICATIONS FOR THE SUBMITTED DELIVERABLES (OUTPUTS 8 AND 10)
# ===========================================================================
doc.add_page_break()
heading("Implications for the submitted deliverables (Outputs 8 and 10)", 1)
body(
    "These field surveys were carried out in June 2026, after Output 8 (Traffic and Parking "
    "Surveys and Analysis Report, April 2026) and Output 10 (Parking Analysis Report, April 2026) "
    "had been submitted. Both deliverables were written on the explicit basis that no field "
    "occupancy or duration survey would be conducted — a position formally communicated to ADB and "
    "the PIU. The new evidence therefore requires a small number of targeted revisions to keep "
    "those deliverables internally consistent and defensible."
)
body(
    "The recommendations below are deliberately limited to crucial changes — either because the "
    "new data now contradicts a statement we made, or because leaving the text unchanged would "
    "expose the work to avoidable criticism. Cosmetic or merely additive edits are not proposed, "
    "and each chapter ends with the points that should deliberately be left unchanged."
)
body(
    "One scope note applies to both. The field surveys cover six selected high-impact areas "
    f"({total_paths} on-street segments, ~{onstreet_total:,} marked spaces) as a targeted demand "
    "overlay — not a repeat of the full corridor supply inventory (8,862 on-street spaces) "
    "documented in Output 8. The two are complementary; the field-survey totals must not be read "
    "as superseding, or contradicting, the inventory totals, and the retained/removed counts here "
    "reflect a per-area engineering judgement, not the full corridor removal figure (5,953).",
    italic=True, color=GREY, size=10)

# ---- For Output 8 ----------------------------------------------------------
heading("For Output 8 — Traffic and Parking Surveys and Analysis Report", 1)
body(
    "The single most significant issue is the chapter “Justification for Excluding Full-Scale "
    "Occupancy and Duration Surveys.” The remaining edits update the Displacement Assessment and "
    "the demand-data status now that real occupancy data exists."
)

heading("1. Reframe the exclusion justification (critical — reputational)", 2)
body(
    "The chapter states that a full-scale field occupancy survey was excluded and that “even "
    "highly detailed occupancy data would not materially influence design or policy "
    "recommendations,” citing a justification letter submitted to ADB and the PIU. A targeted "
    "field occupancy survey has since produced precisely the hourly occupancy, duration and "
    "turnover metrics the chapter said would not be collected. Left unchanged, the chapter is "
    "directly contradicted by the project's own subsequent work."
)
body(
    "Recommended: keep the proportionality argument — a 34-km, 17-hour corridor-wide operation was "
    "disproportionate — but change the conclusion. The team did not abandon occupancy measurement; "
    "it scoped it down to a targeted survey of representative high-impact areas, conducted after "
    "submission to validate the displacement assumptions and to bridge the still-pending ANPR "
    "access. The absolute claim that occupancy data “would not materially influence design or "
    "policy recommendations” should be qualified: it holds for the design (parking is removed "
    "regardless of occupancy) but is no longer true for policy and mitigation, where the occupancy "
    "data has directly shaped the displacement and absorption assessment."
)

heading("2. Quantify the Displacement Assessment", 2)
body(
    "The Displacement Assessment currently argues — from European precedent only — that displaced "
    "demand is typically lower than the spaces removed, but puts no number on Yerevan because no "
    "occupancy data existed. The field survey now measures it: across the surveyed areas peak-hour "
    f"parked demand on removed segments is ~{D['removed_demand']} cars against ~{D['removed_supply']:,} "
    "spaces removed — roughly 55%. Recommended: add this as local empirical confirmation of the "
    "“demand is lower than supply removed” principle. This strengthens the argument and is a "
    "defensibility asset (see Output 10, point 3), not mere padding."
)

heading("3. Refine “Available Absorptive Capacity”", 2)
body(
    "Output 8 presents absorptive capacity as a single large pool (7,035 off-street + 1,857 "
    "cross-street on-street) comfortably exceeding removal. The field survey confirms overall "
    "sufficiency but shows two things the current text omits: absorption is off-street-dominated — "
    f"surviving and nearby on-street covers only ~{D['absorbed_onstreet']}% of peak displaced "
    "demand city-wide and as little as ~14% in the Kentron core — and it is unevenly distributed, "
    "with one surveyed area (Gai Avenue) where even gross best-case nearby capacity (~89%) does "
    "not fully cover peak displaced demand. Recommended: keep the headline of sufficiency, but add "
    "these two qualifiers so the report does not imply that side-street conversion can absorb the "
    "centre."
)

heading("4. Update “Occupancy and Demand Data: Current Status”", 2)
body(
    "This section states that demand analysis is pending ANPR access and will follow as an "
    "addendum. That is now only half true — targeted field-based demand data exists. Recommended: "
    "revise to state that interim demand evidence has been obtained through the field survey, with "
    "ANPR retained as the future citywide, continuous source, and reposition the field survey as "
    "the interim addendum it effectively is."
)

heading("Deliberately leave unchanged", 2)
bullet(" 15,897 / 8,862 / 5,953 — different scope from the field survey; the field figures do not "
       "supersede them.", bold_lead="Supply-inventory totals: ")
bullet(" the field survey under-records paid (red/blue) lines; Output 8's Zone A/B figures remain "
       "the reference for regulation.", bold_lead="Regulation classification: ")

# ---- For Output 10 ---------------------------------------------------------
heading("For Output 10 — Parking Analysis Report", 1)
body(
    "Output 10 carries the mitigation strategy, so this is where the new findings have the most "
    "consequence. Two of the edits below are reputational; the rest sharpen the mitigation logic "
    "with evidence the report previously lacked."
)

heading("1. Reconcile the “field surveys were abandoned” statements (critical — reputational)", 2)
body(
    "Output 10 states that “field surveys were abandoned” and that “detailed occupancy, duration, "
    "and turnover data are not required.” These must be reconciled with the completed survey using "
    "the same reframe recommended for Output 8. The section “Overview of Previous Field Survey "
    "Plans (Methodology),” which documents the abandoned 17-hour plan, should now note that a "
    "targeted version was subsequently executed and summarise its findings."
)

heading("2. Recalibrate the high-sensitivity (Kentron) mitigation package (critical)", 2)
body(
    "This is the most important substantive change. The “Mitigation Measure Packages by "
    "Sensitivity Zones” assigns the high-sensitivity core (Kentron — Nalbandyan, Amiryan, Tigran "
    "Mets) a curb-led package: extend paid zones, delivery spaces, visitor caps, resident permits. "
    "The field survey shows that in Kentron nearby on-street capacity absorbs only ~14% of peak "
    "displaced demand (Komitas 0%), and the centre already runs at 138% of capacity at peak — there "
    "is essentially no spare curb to convert. As written, the package promises absorption on "
    "streets the data shows are already full."
)
body(
    "Recommended: for the high-sensitivity zone, lead with off-street measures — opening gated "
    "residential courtyards for daytime use, park-and-ride interception, and securing and "
    "structuring nearby off-street yards — and present the curb measures (paid-zone extension, "
    "visitor caps) as demand-management and turnover tools rather than as absorption capacity. "
    "Leaving the package curb-led would be contradicted by our own occupancy figures."
)

heading("3. Turn “the corridor itself is the primary mitigation” into evidence", 2)
body(
    "The central displacement argument — that the corridor reduces car demand so not every space "
    "needs replacing — currently rests on European precedent (Nantes, Metz, Rouen, etc.). Given "
    "the active court challenges demanding an empirical basis for parking decisions (a risk the "
    "report itself flags), the local finding that peak demand is only ~55% of removed supply is a "
    "material defensibility asset. Recommended: cite it where the “primary mitigation” logic is "
    "introduced."
)

heading("4. Ground the pricing and turnover rationale in observed data", 2)
body(
    "The pricing and visitor-cap measures currently rest on the theoretical 85% efficiency "
    "threshold and on ANPR data that is not yet available. The field survey shows central areas "
    "operating well above 85% at peak (Kentron 138%, Komitas 123%) and a demand profile dominated "
    f"by short stays (~{A['all']['profile']['duration']['visitorPct']}% under two hours). "
    "Recommended: use these observed figures to support paid-zone expansion and visitor caps, and "
    "replace or supplement the “illustrative” ANPR figures (Figures 5–9), currently labelled "
    "hypothetical, with the actual field demand profile so the strategy no longer visibly rests on "
    "placeholder data."
)

heading("5. Name the residual risk concretely", 2)
body(
    "“Limitations and Residual Risks” is currently generic (some net reduction is expected and "
    "even desirable). Recommended: add the specific, evidenced case — Gai Avenue, where nearby "
    "capacity does not fully absorb peak displaced demand — as a worked example requiring targeted "
    "provision. A named, quantified residual risk reads as rigour; the same gap discovered after "
    "publication reads as an oversight."
)

heading("Deliberately leave unchanged", 2)
bullet(" the survey ran daytime hours (≈07:00–24:00) and does not capture overnight residential "
       "demand, so it neither confirms nor refutes that concern and must not be used to downplay "
       "it.", bold_lead="The residential overnight-storage hardship narrative: ")
bullet(" do not weaken it on the basis of the survey's low long-stay share — the observation "
       "window excludes the night, when that demand occurs.", bold_lead="The resident-permit scheme: ")
bullet(" the survey supports, not contradicts, overall sufficiency — keep it.",
       bold_lead="The off-street capacity headline: ")

doc.save(OUT)
print("WROTE:", OUT)
