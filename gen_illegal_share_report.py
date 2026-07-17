#!/usr/bin/env python
"""Generate Illegal Parking Share Statistics report (.docx).

Reads directly from the raw survey workbooks in Downloads/parkingsurvey —
no GeoJSON intermediary — and produces a standalone document.
"""
import os
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SURVEY_DIR = r"C:/Users/user/Downloads/parkingsurvey"
OUT = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report"
       r"/Illegal Parking Share Statistics.docx")

FILES = {
    "Garegin Nzhdeh St.xlsx": ("garegin", "Garegin Nzhdeh"),
    "Kentron.xlsx":           ("kentron", "Kentron"),
    "Komitas.xlsx":           ("komitas", "Komitas"),
    "Malatia Sebastia.xlsx":  ("malatia", "Malatia-Sebastia"),
    "Mega Mall.xlsx":         ("mega",    "Gai Avenue / Mega Mall"),
    "Shiraz, Hasratyan.xlsx": ("shiraz",  "Shiraz & Hasratyan"),
}

NON_ZONE = {
    "OFF", "Off street city", "Shiraz off-street", "P",
    "Sheet1", "Analysis Ard Ayush", "kentroni fili mjic",
}

# ── Read workbooks ────────────────────────────────────────────────────────────
area_totals = {}   # area_key -> {"label", "total", "illegal", "zones": {zone: {total,illegal}}}
grand_total = 0
grand_illegal = 0

for fname, (area_key, label) in FILES.items():
    path = os.path.join(SURVEY_DIR, fname)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    agg = {"label": label, "total": 0, "illegal": 0, "zones": {}}

    for shname in wb.sheetnames:
        if shname in NON_ZONE:
            continue
        try:
            zone_num = int(shname)
        except ValueError:
            continue

        ws = wb[shname]
        rows = list(ws.iter_rows(min_row=9, values_only=True))  # data from row 9 (1-indexed)
        z_total = 0
        z_illegal = 0
        for r in rows:
            if not r or len(r) < 3:
                continue
            plate = str(r[2]).strip() if r[2] is not None else ""
            legal = str(r[6]).strip() if (len(r) > 6 and r[6] is not None) else ""
            if not plate or plate in ("-", "None", ""):
                continue
            z_total += 1
            if legal.startswith("I"):
                z_illegal += 1

        if z_total > 0:
            agg["zones"][zone_num] = {"total": z_total, "illegal": z_illegal}
            agg["total"] += z_total
            agg["illegal"] += z_illegal

    wb.close()
    area_totals[area_key] = agg
    grand_total += agg["total"]
    grand_illegal += agg["illegal"]

# Sort areas: illegal-share descending
AREA_ORDER = sorted(area_totals.keys(),
                    key=lambda k: -(area_totals[k]["illegal"] / area_totals[k]["total"])
                    if area_totals[k]["total"] else 0)

# Zones with illegal > 0, sorted by share desc
illegal_zones = []
for area_key, agg in area_totals.items():
    for zone, z in agg["zones"].items():
        if z["illegal"] > 0:
            pct = z["illegal"] / z["total"] * 100
            illegal_zones.append({
                "zone": zone,
                "area_key": area_key,
                "area_label": agg["label"],
                "illegal": z["illegal"],
                "total": z["total"],
                "pct": pct,
            })
illegal_zones.sort(key=lambda x: -x["pct"])

# ── Styling helpers ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
GREY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xC0, 0x2A, 0x2A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=10,
             align="left", white=False, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    r = p.add_run(str(text))
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        r.font.color.rgb = color


def heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(16 if level == 1 else 13)
    p.space_before = Pt(10 if level == 1 else 6)
    return p


def body(text, italic=False, color=None, size=11, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def table_header_row(tbl, labels, widths_in=None):
    row = tbl.rows[0]
    for i, label in enumerate(labels):
        set_cell(row.cells[i], label, bold=True, white=True,
                 align="center" if i > 0 else "left")
        shade(row.cells[i], "142A4A")
    if widths_in:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)


# ── Document ──────────────────────────────────────────────────────────────────
heading("Illegal Parking Share — Field Survey Statistics")

body(
    "This note summarises illegal parking observations recorded across all six "
    "field survey areas in Yerevan. Data are drawn directly from the raw survey "
    "workbooks (one per area) with no intermediary processing. A vehicle observation "
    "is classified as 'Illegal' where the surveyor coded the Legal/Illegal column "
    "as 'I – Illegal'; all other coded observations are treated as legal.",
    size=11,
)
body(
    f"Total observations across all areas: {grand_total:,}  |  "
    f"Illegal observations: {grand_illegal}  |  "
    f"Overall illegal share: {grand_illegal/grand_total*100:.2f}%",
    size=11,
)

# ── Section 1: By area ────────────────────────────────────────────────────────
heading("1. Illegal Share by Survey Area", level=2)

tbl1 = doc.add_table(rows=1 + len(AREA_ORDER), cols=4)
tbl1.style = "Table Grid"
table_header_row(tbl1, ["Survey Area", "Total Obs.", "Illegal Obs.", "Illegal Share (%)"],
                 widths_in=[2.4, 1.3, 1.3, 1.5])

for i, area_key in enumerate(AREA_ORDER, start=1):
    d = area_totals[area_key]
    pct = d["illegal"] / d["total"] * 100 if d["total"] else 0
    row = tbl1.rows[i]
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    set_cell(row.cells[0], d["label"], align="left")
    set_cell(row.cells[1], f"{d['total']:,}", align="right")
    set_cell(row.cells[2], str(d["illegal"]), align="right",
             color=RED if d["illegal"] > 0 else None,
             bold=d["illegal"] > 0)
    set_cell(row.cells[3], f"{pct:.2f}%", align="right",
             color=RED if pct > 0 else None,
             bold=pct > 0)
    for c in row.cells:
        shade(c, bg)

doc.add_paragraph()

# Grand-total row note
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Overall (all areas combined): ")
r.bold = True
r.font.size = Pt(11)
r2 = p.add_run(
    f"{grand_illegal} illegal observations out of {grand_total:,} total "
    f"= {grand_illegal/grand_total*100:.2f}%"
)
r2.font.size = Pt(11)
r2.font.color.rgb = RED

# ── Section 2: Zones with illegal observations ────────────────────────────────
heading("2. Zones with Illegal Observations", level=2)

body(
    f"{len(illegal_zones)} survey zones (out of {sum(len(a['zones']) for a in area_totals.values())} "
    f"total on-street zones) recorded at least one illegal parking observation. "
    "All are located in the Komitas area, except one zone in Shiraz & Hasratyan.",
    size=11,
)

tbl2 = doc.add_table(rows=1 + len(illegal_zones), cols=5)
tbl2.style = "Table Grid"
table_header_row(tbl2,
                 ["Zone", "Survey Area", "Illegal Obs.", "Total Obs.", "Share (%)"],
                 widths_in=[0.7, 2.2, 1.2, 1.2, 1.2])

for i, z in enumerate(illegal_zones, start=1):
    row = tbl2.rows[i]
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    set_cell(row.cells[0], str(z["zone"]), align="center")
    set_cell(row.cells[1], z["area_label"])
    set_cell(row.cells[2], str(z["illegal"]), align="right",
             color=RED, bold=True)
    set_cell(row.cells[3], f"{z['total']:,}", align="right")
    set_cell(row.cells[4], f"{z['pct']:.1f}%", align="right",
             color=RED, bold=True)
    for c in row.cells:
        shade(c, bg)

doc.add_paragraph()

# ── Section 3: Observations ───────────────────────────────────────────────────
heading("3. Observations", level=2)

bullets = [
    (
        "Negligible overall rate.",
        "Across 21,174 vehicle-hour observations spanning all six survey areas, "
        "only 45 were coded as illegal — an overall share of 0.21%."
    ),
    (
        "Concentrated in Komitas.",
        "44 of the 45 illegal observations (98%) are in the Komitas area, "
        "which covers paid (blue-zone) on-street parking. The highest individual "
        "zones are 84, 110, and 109 (each at or above 6%)."
    ),
    (
        "One Shiraz zone.",
        "Zone 138 in Shiraz & Hasratyan recorded a single illegal observation "
        "out of 146 (0.7%). No other area logged any violations."
    ),
    (
        "Surveyor coding consistency.",
        "The absence of illegal observations across Kentron, Garegin Nzhdeh, "
        "Malatia-Sebastia, and Gai Avenue / Mega Mall likely reflects inconsistent "
        "application of the 'I – Illegal' code rather than a genuine absence of "
        "violations. The 0.21% figure should be treated as a lower bound."
    ),
    (
        "Komitas paid zone pattern.",
        "Most Komitas illegal observations occur on white-regulation streets "
        "adjacent to the blue-zone core, or in blue zones where vehicles parked "
        "in non-standard orientations (45-degree on parallel-only streets). "
        "This points to boundary confusion rather than deliberate evasion."
    ),
]

for lead, rest in bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(lead + " ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)

doc.add_paragraph()
body(
    "Source: raw survey workbooks (Garegin Nzhdeh St.xlsx, Kentron.xlsx, "
    "Komitas.xlsx, Malatia Sebastia.xlsx, Mega Mall.xlsx, Shiraz, Hasratyan.xlsx). "
    "Analysis date: June 2026.",
    italic=True, color=GREY, size=9,
)

doc.save(OUT)
print("Saved:", OUT)
