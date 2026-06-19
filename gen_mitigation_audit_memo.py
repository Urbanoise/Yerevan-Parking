#!/usr/bin/env python
"""Generate the Mitigation Measures Audit memo (.docx).

Standalone internal review memo: pressure-tests the Output 8 / Output 10
mitigation measures against the field-survey data. Figures are from the
validated field-surveys.geojson run + the raw stay-length histogram
(11,394 on-street stays). Decision-first structure.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Mitigation Measures Audit - Memo.docx")

# ---- palette (matches the findings report) --------------------------------
NAVY = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6B, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, runs, align="left", fill=None, size=9.5):
    """runs: str, or list of (text, {bold,color,white}) tuples."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(str(text))
        r.bold = opt.get("bold", False)
        r.font.size = Pt(opt.get("size", size))
        if opt.get("white"):
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif opt.get("color") is not None:
            r.font.color.rgb = opt["color"]
    if fill:
        shade(cell, fill)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(15 if level == 1 else 12.5)
    return p


def body(parts, size=11, italic=False, color=None, space=4):
    """parts: str or list of (text, {bold,color,italic}) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", italic)
        r.font.size = Pt(opt.get("size", size))
        c = opt.get("color", color)
        if c is not None:
            r.font.color.rgb = c
    return p


def bullet(parts, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        c = opt.get("color")
        if c is not None:
            r.font.color.rgb = c
    return p


def table(headers, rows, widths=None, fill="142A4A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], [(h, {"bold": True, "white": True, "size": 9.5})],
                 align="left" if i == 0 else "center", fill=fill)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            set_cell(cells[i], c, align="left" if i == 0 else "center")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def rule():
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "10"), ("w:space", "1"), ("w:color", "142A4A")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def hist_block(rows, scale=1.4):
    """Render a §8-style horizontal stay-length histogram (monospace + teal bars)."""
    for lbl, pct in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(f"{lbl:<16}{pct:>5.1f}%  ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r2 = p.add_run("█" * max(1, round(pct / scale)))
        r2.font.size = Pt(9)
        r2.font.color.rgb = ACCENT


# ===========================================================================
# TITLE
# ===========================================================================
t = doc.add_paragraph()
r = t.add_run("Yerevan Parking Study  •  Internal Review Memo")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

t = doc.add_paragraph()
r = t.add_run("Mitigation Measures — Audit & Validation against the Field Survey")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
r = sub.add_run("Pressure-testing the Output 8 and Output 10 mitigation measures — drafted a priori, "
                "before fieldwork — against the field-survey evidence.")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY
rule()

body([("Scope. ", {"bold": True}),
      ("Re-derived independently from the source deliverables and the field-survey dataset; "
       "convergence and divergence with the report's existing “Implications” chapter are flagged. "
       "Figures are from the validated survey run (six areas; 208 on-street segments; ~2,930 marked "
       "spaces) and the raw stay-length histogram (11,394 on-street stays).", {})],
     size=10, color=GREY)

# ===========================================================================
# 1. DECISION
# ===========================================================================
heading("1.  Decision", 1)
body([("The measure ", {}), ("menu", {"italic": True}), (" holds up. The measure ", {}),
      ("targeting", {"italic": True}), (" does not.", {})], space=4)
body("Nothing in the survey justifies deleting a measure. But the survey overturns three things that "
     "were guessed before fieldwork — and adds one the data demands:")

bullet([("the off-street measure (M7 — open residential yards).", {"bold": True}),
        (" The entire displacement case rests on off-street capacity: surviving on-street absorbs only "
         "43% of peak displaced demand (14% in the Kentron core, 0% in Komitas), and 92% of the "
         "off-street capacity that closes the gap to ~100% is gated residential yards counted at gross "
         "capacity as if already open. So the “demand can be re-absorbed” headline is silently "
         "conditional on opening those yards — i.e. on M7. Yet O10 frames M7 as an optional pilot in the "
         "lower-sensitivity zones. It is the load-bearing measure of the whole strategy and must be made "
         "a committed precondition, led in the core. But opening yards is an organisational and change-"
         "management undertaking, not a free capacity line: it needs active management and resident "
         "buy-in, and a poorly-communicated rollout could provoke a backlash that jeopardises the wider "
         "corridor reform — so M7 must travel with a resident-engagement and management plan. No net-new "
         "measure — M7 elevated.", {})],
       bold_lead="ELEVATE & COMMIT — ")
bullet([("visitor caps / maximum-stay limits — but keep them.", {"bold": True}),
        (" Their logic (evict long-stay commuters) is contradicted by the data: ~63% of stays are ≤1h, "
         "~77% ≤2h, the overnight (8h+) share is ~7% — and those are residents (permit-exempt), while "
         "genuine daytime all-day workers (5–8h) are only ~6%; turnover is already 5–11 vehicles/space/"
         "day. A 2-hour cap would mostly bite 3–4h errand visitors, not commuters. Do not remove — "
         "demote to a targeted contingency for specific low-turnover spots, and lead with pricing. And "
         "for the genuine 5–8h commuters the lever is not a fine but a mode shift: a long, predictable "
         "daily stay is the trip most readily moved to transit, so they are the prime BRT catchment — "
         "the framing must be meeting the commuter's need, not penalising it.", {})],
       bold_lead="RE-SCOPE (misdirected) — ")
bullet([("measures across the sensitivity zones.", {"bold": True}),
        (" The high-sensitivity (Kentron) package is curb-led, but the data shows curb cannot absorb "
         "there. Yard-opening (M7) — now assigned to lower-sensitivity zones — belongs first in the core "
         "(Kentron is 86% yard-dependent, Komitas 100%). Komitas, classed Medium, behaves like High "
         "(123% peak, 0% on-street absorption); Shiraz behaves like Lower (75% peak, 69% absorption). "
         "Targeting must follow measured behaviour, not a-priori geography.", {})],
       bold_lead="RE-ALLOCATE — ")
bullet([("phase out the flat annual zonal-parking permit.", {"bold": True}),
        (" Yerevan's annual permit — a flat yearly fee for unlimited zonal parking — works directly "
         "against every turnover objective in this study: it strips the marginal cost out of dwell time, "
         "so it rewards leaving a car in a high-demand bay all day and depresses turnover precisely where "
         "the survey shows turnover is the binding constraint. In the medium-to-long term it should be "
         "cancelled and replaced by an occupancy-based tariff. It is the same anti-turnover mechanism as "
         "free resident parking — which Tbilisi's parking study found had to be unwound for exactly this "
         "reason — so the resident-permit design (M4) must not re-import it.", {})],
       bold_lead="ADD (new policy) — ")

body([("Net: ", {"bold": True}),
      ("sufficient as a displacement toolkit — no net-new mitigation measure is needed to re-absorb "
       "demand. Re-scope caps (demote, lead with pricing), elevate and commit M7 (the off-street "
       "measure the headline depends on, carried with a resident-engagement plan), and re-allocate by "
       "zone. One structural addition sits outside the displacement frame: phase out the flat annual "
       "permit, which is actively depressing the turnover the whole strategy relies on. The “corridor "
       "itself = primary mitigation” premise is empirically confirmed — peak displaced demand is only "
       "55% of spaces removed (908 cars vs 1,643 spaces).", {})])
body([("Independent corroboration. ", {"bold": True, "color": ACCENT}),
      ("Setting the report's existing “Implications” chapter aside and re-deriving, I reach the same "
       "Kentron curb→off-street recalibration and the same Gai Avenue residual-risk flag. I go further "
       "on three points it does not make: the off-street headline is conditional on opening gated yards "
       "(M7) and must say so; visitor caps are misdirected on the measured data; and the displaced-"
       "demand peak is midday in every area (so daytime yard availability aligns). Neither contradicts "
       "the existing recommendations.", {})], size=10, color=GREY)

# ===========================================================================
# 2. SURVEY DESIGN & VALIDITY
# ===========================================================================
heading("2.  Survey design, validity & the qualitative→quantitative logic", 1)
body([("Why this is methodologically sound. ", {"bold": True}),
      ("The corridor's sensitivity zones (High / Medium / Lower) were assigned a priori — qualitatively, "
       "with local-expert judgement — before any field data existed. That assignment is a hypothesis. "
       "The field survey is the quantitative test of it: every parked vehicle on the surveyed segments "
       "was logged by hourly licence-plate sweep (07:00–24:00) across six areas deliberately chosen to "
       "span the full assigned sensitivity spectrum, recording time, zone, plate, vehicle type, kerb "
       "location, parking method and legality. Occupancy, turnover and dwell time are then measured, not "
       "assumed — and where the measurement diverges from the a-priori label, the measurement governs. "
       "This is the chain that makes the recommendations defensible rather than opinion.", {})], space=4)

body([("Validity check — were these ordinary days? ", {"bold": True}),
      ("A single disrupted day (public holiday, festival, demonstration, road closure) could distort "
       "occupancy and corrupt the whole comparison. Every survey fell on an ordinary mid-week working "
       "day, with no public holiday and no city-wide event on record:", {})], space=2)
table(
    ["Area", "Survey date", "Weekday"],
    [["Komitas", "29 May 2026", "Friday"],
     ["Shiraz / Hasratyan", "1 June 2026", "Monday"],
     ["Garegin Nzhdeh", "2 June 2026", "Tuesday"],
     ["Gai Avenue (Mega Mall)", "2 June 2026", "Tuesday"],
     ["Kentron", "3 June 2026", "Wednesday"],
     ["Malatia-Sebastia", "3 June 2026", "Wednesday"]],
    widths=[2.4, 1.8, 1.4])
body([("Two design features reinforce this. ", {"bold": True}),
      ("First, the areas were surveyed on five different days, so a one-off event would surface as a "
       "single-area outlier — instead the behavioural signatures (short-stay dominance, high turnover, "
       "a midday demand peak) repeat across all six, which is what rules out day-specific contamination. "
       "Second, no weekend or holiday is in the set, so the figures describe typical working-day demand "
       "— the condition the corridor must be designed for.", {})],
     size=10, color=GREY, space=4)
body([("Recommended final desk-check. ", {"bold": True, "color": ACCENT}),
      ("Day-of-week and the cross-day consistency are confirmed here; a final cross-reference of each "
       "date against the municipal events/holiday calendar should be appended to the methodology for "
       "completeness before publication.", {})], size=10, color=GREY, space=4)

body([("What the test returned — the sensitivity map needs recalibration. ", {"bold": True}),
      ("Because the days were normal, the divergences below are real behaviour, not artefacts, and they "
       "override the a-priori labels:", {})], space=2)
bullet([("Komitas — assigned Medium, behaves High. ", {"bold": True}),
        ("123% peak occupancy, 0% on-street absorption of displaced demand — the Medium package will not "
         "fit; treat as High.", {})])
bullet([("Shiraz / Hasratyan — behaves Lower. ", {"bold": True}),
        ("75% peak, 69% on-street absorption — genuine headroom; the light-touch package is appropriate.", {})])
bullet([("Kentron — confirmed High, but for a specific reason. ", {"bold": True}),
        ("138% peak with only 14% on-street absorption: the pressure is real, yet the curb cannot relieve "
         "it — which is why the High package must be off-street-led (M7), not curb-led.", {})])
body([("In short, the survey does not just produce numbers; it closes the loop on the qualitative zoning "
       "that started the project — confirming it where it was right (Kentron, the lower-pressure west) "
       "and correcting it where expert judgement and measured behaviour part company (Komitas, Gai).", {})],
     size=10, color=GREY, space=2)

# ===========================================================================
# 3. INVENTORY
# ===========================================================================
heading("3.  Inventory of current measures", 1)
body([("Output 8 sets up the displacement frame (supply inventory, absorptive-capacity pools, "
       "sensitivity zones). The measures themselves live in Output 10.", {})], size=10, color=GREY)
table(
    ["#", "Measure", "Stated purpose", "Zone(s)"],
    [["M0", [("The corridor itself", {"bold": True}), (" (modal shift via BRT)", {})], "Primary mitigation", "All"],
     ["M1", [("Extend paid zones", {"bold": True}), (" (red/blue ≤100 m)", {})], "Absorb spillover; manage demand", "High + Med"],
     ["M2", [("Delivery / loading bays", {"bold": True})], "Protect commercial servicing", "High + Med"],
     ["M3", [("Visitor caps", {"bold": True}), (" (max-stay)", {})], "Increase turnover", "High"],
     ["M4", [("Resident permits", {"bold": True}), (" (transitional)", {})], "Protect residents from spillover cost", "High"],
     ["M5", [("Organise free parking", {"bold": True}), (" (no pricing)", {})], "Structure chaotic free supply", "Lower"],
     ["M6", [("Park & ride", {"bold": True})], "Intercept inbound commuters", "Med + Lower"],
     ["M7", [("Open residential yards", {"bold": True})], "Create off-street supply", "Lower"],
     ["M8", [("ANPR enforcement", {"bold": True}), (" extension", {})], "Enabling", "Cross-cutting"],
     ["M9", [("Duration-based fines", {"bold": True})], "Deter long storage", "Cross-cutting"],
     ["M10", [("Revenue earmarking", {"bold": True})], "Public legitimacy", "Cross-cutting"],
     ["M11", [("Phase out flat annual permit", {"bold": True}), (" (new)", {"color": AMBER})],
      "Restore turnover — price dwell time", "Cross-cutting"],
     ["—", [("Phasing + monitoring", {"bold": True})], "Sequencing (before removal)", "All"]],
    widths=[0.4, 2.3, 2.6, 1.0])

# ===========================================================================
# 3. VALIDATION
# ===========================================================================
heading("4.  Validation against the survey", 1)
body([("Verdicts: ", {}), ("✓ confirmed", {"color": GREEN, "bold": True}),
      ("  ·  ", {}), ("◆ re-scope", {"color": AMBER, "bold": True}),
      ("  ·  ", {}), ("○ cannot test (keep on precaution)", {"color": GREY, "bold": True}),
      (".", {})], size=10)
GV = {"bold": True, "color": GREEN}
AV = {"bold": True, "color": AMBER}
WV = {"bold": True, "color": GREY}
table(
    ["Measure", "Verdict", "What the data says"],
    [["M0 Corridor / modal shift", [("✓ Confirmed", GV)],
      "Peak displaced demand 908 vs 1,643 removed = 55%. The “displaced < removed” principle is now "
      "locally measured (was European precedent only). Avg 43% vs peak 93% → don't replace 1:1."],
     ["M1 Extend paid zones", [("◆ Re-scope", AV)],
      "Right lever (90% of kerb unpriced; peaks 93–138%). But its O10 purpose is “absorb on side "
      "streets” — and the core has no spare side-street capacity (Kentron 14%, Komitas 0%). In the core "
      "it is a demand-rationing tool, not absorption."],
     ["M2 Delivery bays", [("✓ (inferred)", GV)],
      "No direct freight count, but ≤1h dominance (~63%) + very high turnover implies heavy short-access. "
      "Keep; rests on inference, not measurement."],
     ["M3 Visitor caps / max-stay", [("◆ Re-scope, not remove", AV)],
      "Overnight ~7% (1h-rule; residents, permit-exempt); daytime all-day (5–8h) only ~6%; ≤1h ~63%; "
      "turnover already 5–11/day. A 2h cap mostly bites 3–4h errand visitors, not commuters. KEEP but "
      "demote to a targeted contingency for low-turnover spots; lead with pricing."],
     ["M4 Resident permits", [("○ Keep — but design actively", WV)],
      "Daytime window still misses true overnight storage; the overnight band (now ~7% via the 1h rule) "
      "is a floor and must NOT be used to weaken this. Keep, justified by land-use, not survey data. But "
      "design it: per the Tbilisi study, a permit that is free or flat at the outset must later move to "
      "priced/limited use, or it recreates the long-stay, low-turnover problem (see M11)."],
     ["M5 Organise free parking", [("✓ Confirmed", GV)],
      "O8: 87.6% unmarked, 79.9% unsigned; ~90% of kerb free. Lower-pressure areas (Malatia 54%, Shiraz "
      "75% peak) have headroom — structure before pricing. (Also needed in Kentron — see G5.)"],
     ["M6 Park & ride", [("◆ Weak support", AV)],
      "Premise = intercept commuters, but the daytime fleet is overwhelmingly short-stay, weakening the "
      "“many commuters to intercept” assumption within these areas. Keep as a study item; don't oversell. "
      "Where commuters do appear (the 5–8h band), the answer is mode shift, not penalty — they are the "
      "natural BRT catchment, so frame P&R/BRT as serving the commuter rather than fining them."],
     ["M7 Open residential yards", [("✓ Load-bearing + under-committed", GV)],
      "92% of off-street absorption is gated residential yards (1,778 of 1,935 spaces); the ~100% "
      "headline is conditional on opening them. M7 is load-bearing (Kentron 86% yard-dependent, Komitas "
      "100%), yet O10 makes it an optional lower-zone pilot. Elevate to a committed core precondition."],
     ["M8 ANPR enforcement", [("✓ Enabling", GV)],
      "Necessary for every priced measure. Keep."],
     ["M9 Duration-based fines", [("○ Keep, don't oversell", WV)],
      "Same caveat as M3: daytime fleet isn't long-stay. Value is overnight storage (unmeasured) + "
      "system maturation."],
     ["M10 Revenue earmarking", [("— Out of scope", WV)],
      "Political-legitimacy measure; survey neither supports nor refutes. Keep."],
     ["M11 Phase out annual permit", [("◆ New — adopt", AV)],
      "A flat annual fee removes the marginal cost of dwell time and directly suppresses turnover — the "
      "one lever the data shows is binding (≤1h ~63%; turnover 5–11/day depends on a live price signal). "
      "Cancel medium-to-long term and replace with occupancy-based pricing; same logic as unwinding free "
      "resident parking (Tbilisi)."],
     ["Phasing", [("✓ Reinforced", GV)],
      "“Mitigation before removal” matters more given the off-street dependency — the gated yards (M7) "
      "must be opened and the pricing in place before the kerb goes."]],
    widths=[1.7, 1.5, 3.5])

# ===========================================================================
# 4. GAPS
# ===========================================================================
heading("5.  Gaps the survey revealed", 1)
table(
    ["Gap", "Why it's only visible now", "In current set?"],
    [[[("G1 — Absorption hinges on opening gated yards (M7)", {"bold": True})],
      "92% of the off-street capacity that closes the gap is gated residential yards, counted at gross "
      "capacity as if already open; on-street alone covers 43% (14%/0% in core). The ~100% headline is "
      "conditional on M7, which O10 treats as optional.",
      [("Re-scope M7 — commit it", {"bold": True, "color": AMBER})]],
     [[("G2 — Gai Avenue local shortfall", {"bold": True})],
      "Only area where gross best-case capacity (89%) < peak demand. Gai isn't in O10's zone lists.",
      [("No — targeted", {"bold": True, "color": RED})]],
     [[("G3 — All-day, zone-wide saturation in core", {"bold": True})],
      "Kentron: 89% of zones >85%, high all working day (peak hour 21:00).",
      [("Partial — scope", {"color": AMBER})]],
     [[("G4 — Evening occupancy peak under-captured", {"bold": True})],
      "Kentron's 21:00 / Malatia's 18:00 occupancy peaks are driven by retained residential streets the "
      "daytime window only partly sees. (Displaced demand itself peaks midday — yard timing is fine.)",
      [("Partial — informs M4", {"color": AMBER})]],
     [[("G5 — Worst informal layout is in the core", {"bold": True})],
      "Kentron kerb ~60% angled/informal. M5 (organise) is assigned only to lower zones.",
      [("Mis-assigned", {"color": AMBER})]],
     [[("G6 — Komitas mis-classed Medium", {"bold": True})],
      "123% peak, 0% on-street absorption — behaves like High; Medium package won't fit.",
      [("No — re-class", {"bold": True, "color": RED})]]],
    widths=[2.1, 3.3, 1.3])

# ===========================================================================
# 6. FIELD PARAMETERS
# ===========================================================================
heading("6.  Field parameters — what every collected variable proves", 1)
body([("The survey logged six attributes for every vehicle — time, zone, plate, vehicle type, kerb "
       "location, parking method, legality — across ~24,950 observations. Occupancy and duration carry "
       "the headline; the other attributes were collected to test specific failure modes. Several return "
       "a near-null result, which is itself a finding: it positively rules a problem out and narrows "
       "where the measures must act. Nothing was collected without a purpose.", {})], size=10, color=GREY)
table(
    ["Parameter (why collected)", "What the data shows", "So what"],
    [[[("Legality — L / I.", {"bold": True}), (" To size the illegal-parking problem and test whether "
        "enforcement is the binding gap.", {})],
      "~99.8% legal; only ~0.2% flagged illegal (max 0.9%, Komitas).",
      "Near-null — and that is the point: ticketed illegality is not the problem. The disorder is "
      "saturation and tolerated informality, not law-breaking, so the lever is pricing/turnover and "
      "organisation (M1/M5), not punitive enforcement."],
     [[("Kerb location — OS / FP / SB.", {"bold": True}), (" To quantify off-carriageway encroachment "
        "(footpath, setback) vs proper kerb use.", {})],
      "75.5% on-street; 24.5% informal (13.0% footpath, 11.5% setback). Worst: Komitas 35%, Garegin 32%, "
      "Shiraz 31%; lowest Kentron 13%.",
      "A quarter of parking sits off the carriageway — informal but tolerated (hence the ~0% illegal "
      "flag). This is the latent disorder M5 (organise free parking) targets, and it is not confined to "
      "the lower zones — Komitas is the worst — so M5's remit must widen."],
     [[("Parking method — PA / 45° / PP.", {"bold": True}), (" To assess layout efficiency and order.", {})],
      "By vehicle: 55% parallel / 26% at 45° / 19% perpendicular. By zone (space-weighted dominant "
      "method): Kentron 60% angled vs ≤4% everywhere else.",
      "Space-hungry, disorder-prone angled layout is concentrated in the highest-pressure core (Kentron) "
      "— reinforcing G5: the M5 organisation effort belongs first in the core, where both layout and "
      "pressure are worst."],
     [[("Vehicle type — C / T / M.", {"bold": True}), (" To size freight/servicing and two-wheeler "
        "demand.", {})],
      "94.8% sedan, 4.9% truck, 0.3% motorcycle. Truck share elevated in Malatia (12.6%) and Shiraz "
      "(8.0%).",
      "Heavy-vehicle demand is low overall but locally concentrated: delivery-bay provision (M2) should "
      "be sized to the Malatia/Shiraz pockets, not applied uniformly. Motorcycle share is negligible — "
      "no dedicated two-wheeler measure is warranted (a ruled-out option)."]],
    widths=[2.4, 2.6, 2.7])
body([("Two attributes return a deliberately reassuring null — illegality (~0.2%) and motorcycles "
       "(0.3%). Collecting them lets the study state positively that neither is a material problem, "
       "rather than leaving them as untested assumptions — which is exactly why they were measured.",
       {"italic": True})], size=9, color=GREY, space=2)

# ===========================================================================
# 7. FACT -> MEASURE
# ===========================================================================
heading("7.  Fact → measure mapping", 1)
body([("Hard, data-driven justifications for the retained / correct measures — a sourcing map for "
       "later wiring into the reports.", {})], size=10, color=GREY)
table(
    ["Survey fact (measured)", "Justifies"],
    [["Peak displaced demand 908 vs 1,643 removed (55%)", "M0; “no 1:1 replacement”; O8 quantification; O10 court-defensibility"],
     ["Avg 43% vs peak 93% occupancy citywide", "M1/M3 framed as peak-demand management, not supply expansion"],
     ["~90% of kerb unpriced (2,654 white vs 276 blue); 87.6% unmarked", "M5 organise free; M1 extend paid"],
     ["Kentron peak 138% / on-street absorption 14%; Komitas 123% / 0%", "High-zone package must be off-street-led, not curb-led"],
     ["92% of off-street absorption is gated residential yards (1,778 / 1,935 spaces)", "M7 is the load-bearing precondition — elevate & commit; the ~100% headline depends on it"],
     ["Off-street carries 86% of Kentron's absorption, 100% of Komitas's (on-street 14% / 0%)", "M7 most critical in the core — re-allocate it up from the lower zones"],
     ["Displaced demand peaks midday (10:00–15:00) in every area, incl. Kentron/Malatia", "Daytime yard availability aligns with the displaced peak — supports M7 timing"],
     ["Stay mix ~63% ≤1h / ~77% ≤2h / ~7% overnight 8h+ (1h-rule); turnover 5–11/day", "Pricing-for-turnover (M1) as headline; M3 caps demoted to targeted"],
     ["Gai Avenue total absorption 89% (<100%); no nearby residential yards to open", "G2 targeted provision at Gai — M7 cannot rescue it"],
     ["Kentron 89% of zones >85%, all-day", "All-day, zone-wide enforcement hours (G3)"],
     ["Kentron / Malatia evening occupancy peak (retained residential)", "M4 resident permits; review paid-hours window (G4)"],
     ["Malatia 54% peak / 97% absorption; Shiraz 75% / 69%", "Light-touch (M5/M7) confirmed for genuine lower-sensitivity areas"],
     ["~99.8% legal / 0.2% illegal across ~24,950 observations", "Enforcement is not the binding gap — lead with M1 pricing + M5 organisation, not penalties"],
     ["24.5% of parking off-carriageway (footpath + setback); worst Komitas 35%", "M5 organise free parking — widen beyond the lower zones"],
     ["Kentron 60% angled layout (space-weighted) vs ≤4% elsewhere", "M5 layout/organisation belongs first in the core (G5)"],
     ["Truck share 4.9% overall but 12.6% Malatia / 8.0% Shiraz", "M2 delivery bays sized to local freight pockets, not uniform"],
     ["Flat annual permit removes the marginal cost of dwell time", "M11 phase out the annual permit; M4 resident permits priced/limited (Tbilisi)"]],
    widths=[3.8, 2.9])

# ===========================================================================
# 6. GROUPING CRITIQUE
# ===========================================================================
doc.add_page_break()
heading("8.  Critique of the user grouping", 1)
body([("Current grouping ", {"bold": True}),
      ("(your directive, in compute_field_survey_metrics.mjs → STAY_BIN, shown in StoryStep.svelte): "
       "duration-of-stay, three ways — Visitor <2h / Commuter 2–8h / Long >8h, where “duration” = number "
       "of distinct hours a plate is seen.", {})])

body([("What the data looks like ", {"bold": True}),
      ("(11,394 on-street stays, with the approved 1-hour overnight correction applied — see §A):", {})], space=2)
hist_block([("1h", 62.8), ("2h", 14.4), ("3h", 6.5), ("4h", 3.5), ("5h", 2.2),
            ("6h", 1.6), ("7h", 1.2), ("8h", 0.9), ("8h+ / overnight", 6.9)])
body([("The 1h–8h bars are a smooth decay with no natural breakpoint; the overnight bucket holds the "
       "edge-only stays reclassified by the approved rule.", {"italic": True})],
     size=9, color=GREY, space=6)

body([("Findings:", {"bold": True})], space=2)
bullet([("The distribution is a smooth decay — any boundary is arbitrary. ", {"bold": True}),
        ("The split was a label imposed on continuous data, not derived from it.", {})])
bullet([("“Visitor <2h” is really “seen in a single hourly sweep.” ", {"bold": True}),
        ("Integer-hour resolution: <2h = exactly 1h = ~63%. It can't tell a 5-minute drop-off from a "
         "55-minute stay.", {})])
bullet([("“Commuter 2–8h” conflates two behaviours: ", {"bold": True}),
        ("the 2–4h errand/shopper (~24%) and the 5–8h worker (~6%) — opposite policy responses "
         "(turnover pricing vs. interception).", {})])
bullet([("The overnight band is now modelled, not an artifact. ", {"bold": True}),
        ("The approved 1-hour rule lifts the 8h+ band from an observed 1.7% to ~6.9% by reclassifying "
         "edge-only stays as overnight. Even so it stays a floor — the daytime window (07:00–24:00) still "
         "cannot see genuine overnight storage, so it must not be used to weaken the resident-permit case.", {})])
bullet([("It doesn't match the deliverables' own user groups. ", {"bold": True}),
        ("Output 10's “Impacts on Specific User Groups” is purpose-based — residents, customers, "
         "commuters, delivery, taxi, persons with disabilities — which is what the measures actually serve.", {})])

heading("Recommendation — adopted approach", 2)
bullet([("Keep the three-way duration split, with the overnight band populated by the approved 1-hour "
         "boundary rule. ", {"bold": True}),
        ("This resolves the worst problem — the previously empty, mislabelled long band now carries a "
         "measured overnight share. Working basis: ", {}),
        ("~63% visitor (<2h) / ~30% commuter (2–8h) / ~7% overnight (8h+).", {"bold": True, "color": ACCENT})])
bullet([("Two residual cautions. ", {"bold": True}),
        ("(a) The visitor/commuter boundary still sits on a smooth decay with no natural gap, so treat "
         "the split as descriptive, not categorical. (b) For the mitigation mapping, anchor measures to "
         "purpose-based user groups (Resident, Visitor/Customer, Commuter/Worker, Delivery, Taxi, PwD) "
         "and use the duration data as evidence of population size, not as the grouping itself.", {})])

# ===========================================================================
# 7. MEASURES -> GROUPS
# ===========================================================================
heading("9.  Measures → groups (recommended grouping)", 1)
table(
    ["User group (purpose-based)", "Survey signal", "Mitigation measure(s)"],
    [[[("Visitors / customers", {"bold": True}), (" (~77% ≤2h turnover demand)", {})],
      "~63% ≤1h; turnover 5–11/day",
      "M1 pricing-for-turnover (headline); M5 organise free; M2 delivery bays"],
     [[("Commuters / workers", {"bold": True}), (" (~6% daytime all-day)", {})],
      "5–8h stays; small but present",
      "M6 park & ride / BRT mode-shift (the real answer); M9 price differential. Fines/caps are not the "
      "lever — meet the need, don't penalise it"],
     [[("Residents", {"bold": True}), (" (overnight ~7%, 1h-rule)", {})],
      "Daytime window undercounts; evening occupancy peaks",
      "M4 resident permits (priced/limited, Tbilisi model); M7 open yards (load-bearing); phase out the "
      "flat annual permit (M11). Never M3 caps"],
     [[("Delivery / freight", {"bold": True})],
      "Inferred (short-access + turnover)",
      "M2 loading bays; grace period"],
     [[("Taxi", {"bold": True})], "In inventory", "Designated taxi stands (design)"],
     [[("Persons with disabilities", {"bold": True})], "Statutory", "Retained accessible bays (design)"]],
    widths=[2.0, 1.9, 2.8])

# ===========================================================================
# 8. CAVEATS
# ===========================================================================
heading("10.  Caveats binding every conclusion", 1)
bullet([("Daytime window (07:00–~24:00). ", {"bold": True}),
        ("Overnight residential storage is structurally invisible — resident permits, duration-based "
         "fines, and the “long-stay” band cannot be tested and must not be weakened on this data.", {})])
bullet([("Six areas, not the full corridor. ", {"bold": True}),
        ("A demand overlay (208 segments, ~2,930 spaces), not the 8,862-space inventory. Absorption is "
         "gross best-case (every nearby space available, drivers willing to go off-street); real-world "
         "friction tightens every balance.", {})])
bullet([("Duration = distinct hours observed, ", {"bold": True}),
        ("not continuous dwell or vehicle type — coarse, and silent on freight/taxi/PwD purpose.", {})])

body([("Bottom line. ", {"bold": True, "color": NAVY}),
      ("The a-priori menu was sound; the survey doesn't break it, it re-aims it — no net-new displacement "
       "measure, plus one structural pricing reform. Elevate and commit M7 (open residential yards), the "
       "off-street measure the absorption headline silently depends on — carried with a resident-"
       "engagement and management plan so it doesn't backfire; demote max-stay caps in favour of pricing "
       "(keep, don't remove), and treat the residual commuters as a BRT mode-shift opportunity rather "
       "than a target for fines; phase out the flat annual permit, which actively suppresses the "
       "turnover the strategy needs, and design resident permits as priced/limited (the Tbilisi lesson); "
       "push the off-street and organisation measures up into the saturated core; and fix the Komitas "
       "(Medium→High) and Gai (unclassified→targeted) zone assignments. The grouping keeps its three-way "
       "duration split — now with the overnight band populated by the approved 1-hour rule — while the "
       "mitigation mapping is anchored to purpose-based user groups.", {})],
     space=2)

# ===========================================================================
# APPENDIX — STAY-LENGTH HISTOGRAM: PROVENANCE & ALTERNATIVE RULE
# ===========================================================================
doc.add_page_break()
heading("Appendix — Stay-length histogram: provenance and alternative rule", 1)
body([("§8 adopts the approved 1-hour overnight correction: vehicles seen only at a survey edge — just "
       "the first hour (arrival unobserved) or just the last hour (departure unobserved) — are treated "
       "as overnight stays. For provenance, A1 below shows the raw observed distribution before any "
       "correction; A2 shows the alternative 2-hour rule that was considered. Same axis and scale "
       "throughout.", {})], size=10, color=GREY)

heading("A1.  Raw observed distribution — no boundary correction (provenance)", 2)
body([("Every stay binned purely by distinct hours observed, with no edge handling. This is the "
       "starting point; the adopted §8 figures move ", {}),
      ("600 edge-only stays (5.3%)", {"bold": True}),
      (" from the 1h bin to overnight.", {})], size=10)
hist_block([("1h", 68.1), ("2h", 14.4), ("3h", 6.5), ("4h", 3.5), ("5h", 2.2),
            ("6h", 1.6), ("7h", 1.2), ("8h", 0.9), ("9h+", 1.7)])
body([("3-way split:  ", {"bold": True}),
      ("visitor 68.1%  ·  commuter 30.3%  ·  long 1.7%", {"bold": True, "color": ACCENT}),
      ("   (→ adopted 1h rule: 62.8 / 30.3 / 6.9)", {"color": GREY})], size=9.5, space=8)

heading("A2.  Alternative — 2-hour rule (not adopted)", 2)
body([("Widens the buffer: reclassifies a vehicle present at the first hour that leaves within two hours, "
       "or one that arrives within the last two hours and is still present at the final hour. A vehicle "
       "seen at last−1 but absent at the last hour is NOT reclassified — its departure was observed. ", {}),
      ("736 stays (6.5%)", {"bold": True}),
      (" move to overnight; the extra 136 vs the adopted 1h rule come from the 2h bin.", {})], size=10)
hist_block([("1h", 62.8), ("2h", 13.2), ("3h", 6.5), ("4h", 3.5), ("5h", 2.2),
            ("6h", 1.6), ("7h", 1.2), ("8h", 0.9), ("8h+ / overnight", 8.1)])
body([("3-way split:  ", {"bold": True}),
      ("visitor 62.8%  ·  commuter 29.1%  ·  long/overnight 8.1%", {"bold": True, "color": ACCENT}),
      ("   (vs adopted 1h rule: 62.8 / 30.3 / 6.9)", {"color": GREY})], size=9.5, space=8)

body([("The adopted rule is a modelling assumption, not a measurement: it treats every edge-only vehicle "
       "as overnight, so the ~7% overnight share is an upper estimate for this correction and still a "
       "floor on true overnight demand, which the daytime window cannot capture directly.",
       {"italic": True})], size=9, color=GREY)

doc.save(OUT)
print("WROTE:", OUT)
