#!/usr/bin/env python
"""Generate the Output 10 — Proposed Changes (Redline Plan) memo (.docx).

Step-1 deliverable: a section-by-section change-list for
"Parking Analysis Report - 20260417.docx" (= Output 10, the measures report; its
internal cover label "Output 8: Parking Analysis Report" is a known mislabel).
For review/approval BEFORE any edit is applied. Full-rework scope. Figures from
the validated field-surveys.geojson run + the Mitigation Measures Audit memo.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Output 10 - Proposed Changes (Redline Plan).docx")

NAVY = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6B, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
INK = RGBColor(0x22, 0x22, 0x22)

TAG = {
    "REFRAME":     RGBColor(0x6A, 0x1B, 0x9A),
    "REPLACE":     RED,
    "INJECT":      RGBColor(0x00, 0x6D, 0x77),
    "ADD":         GREEN,
    "RECALIBRATE": AMBER,
    "FIGURE":      RGBColor(0x15, 0x65, 0xC0),
    "KEEP":        GREY,
}

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = INK


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, runs, align="left", fill=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(str(text))
        r.bold = opt.get("bold", False)
        r.font.size = Pt(opt.get("size", size))
        if opt.get("white"):
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif opt.get("color") is not None:
            r.font.color.rgb = opt["color"]
    if fill:
        shade(cell, fill)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(15 if level == 1 else 12)
    return p


def body(parts, size=10.5, color=None, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        r.font.size = Pt(opt.get("size", size))
        c = opt.get("color", color)
        if c is not None:
            r.font.color.rgb = c
    return p


def bullet(parts, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.font.size = Pt(opt.get("size", 10.5))
        c = opt.get("color")
        if c is not None:
            r.font.color.rgb = c
    return p


def rule():
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "8"), ("w:space", "1"), ("w:color", "142A4A")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def table(headers, rows, widths=None, fill="142A4A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], [(h, {"bold": True, "white": True, "size": 9})],
                 align="left" if i == 0 else "center", fill=fill)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            set_cell(cells[i], c, align="left")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def _line(label, parts, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.18)
    rl = p.add_run(f"{label}:  "); rl.bold = True
    rl.font.size = Pt(9.5); rl.font.color.rgb = color
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text); r.font.size = Pt(9.5)
        r.bold = opt.get("bold", False)
        if opt.get("color") is not None:
            r.font.color.rgb = opt["color"]


def change(section, refs, tag, current, proposed, basis):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(section); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = INK
    if refs:
        r2 = p.add_run("   " + refs); r2.font.size = Pt(8.5); r2.font.color.rgb = GREY
    tp = doc.add_paragraph(); tp.paragraph_format.space_after = Pt(1)
    for i, tg in enumerate(tag if isinstance(tag, list) else [tag]):
        if i:
            tp.add_run("  ")
        rr = tp.add_run(f" {tg} "); rr.bold = True; rr.font.size = Pt(8); rr.font.color.rgb = TAG.get(tg, GREY)
    _line("Current", current, GREY)
    _line("Proposed", proposed, RED)
    _line("Basis", basis, ACCENT)


# ===========================================================================
# TITLE
# ===========================================================================
t = doc.add_paragraph()
r = t.add_run("Yerevan Parking Study  •  Output 10 Revision")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GREY
t = doc.add_paragraph()
r = t.add_run("Output 10 — Parking Analysis Report: Proposed Changes")
r.bold = True; r.font.size = Pt(19); r.font.color.rgb = NAVY
sub = doc.add_paragraph()
r = sub.add_run("A section-by-section change-list for review and approval. No edits have been applied to the "
                "source report yet — on approval, every change will be made in the document in red so it can "
                "be tracked.")
r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = GREY
rule()

body([("Target file. ", {"bold": True}),
      ("“Parking Analysis Report - 20260417.docx”. Its cover label “Output 8: Parking Analysis Report” is a "
       "known mislabel — this is Output 10 (Parking Analysis Report). The companion surveys report (now "
       "“Output 8 - …08042026 (rev).docx”) is Output 8 and owns the survey/supply/displacement frame. "
       "Output 10 owns the mitigation measures and recommendations — which is where this plan does most of "
       "its work, applying the Mitigation Measures Audit memo.", {})], size=10, color=GREY)
body([("How to read. ", {"bold": True}),
      ("Each block: section + paragraph refs, a change-type tag, then ", {}),
      ("Current", {"bold": True, "color": GREY}), (" → ", {}),
      ("Proposed", {"bold": True, "color": RED}), (" → ", {}),
      ("Basis", {"bold": True, "color": ACCENT}),
      (". Tags: REFRAME · REPLACE · INJECT · ADD · RECALIBRATE · FIGURE · KEEP.", {})],
     size=9.5)

# ===========================================================================
# 1. SPINE
# ===========================================================================
heading("1.  What this revision does", 1)
body([("Output 10 carries the mitigation measures, written a priori (before fieldwork). The Mitigation "
       "Measures Audit memo pressure-tested them against the completed field survey and reached a clear "
       "verdict: the measure ", {}),
      ("menu", {"italic": True}), (" holds up, but the measure ", {}),
      ("targeting", {"italic": True}),
      (" does not. This revision applies that verdict — no measure is deleted; three are re-aimed, one is "
       "elevated to a committed precondition, and one further idea (phasing out the flat annual permit) is "
       "flagged for decision pending Parking City Service data — and, like Output 8, it "
       "converts the report's “field survey abandoned / demand via ANPR” framing to “survey conducted, "
       "demand measured.”", {})])
body([("The five headline moves: ", {"bold": True, "color": RED}),
      ("(1) elevate Open Residential Yards to a committed, core-led precondition with a resident-engagement "
       "plan; (2) re-scope Visitor Caps — demote to a targeted contingency, lead with pricing, and treat "
       "commuters as a BRT mode-shift opportunity not a fining target; (3) flag — not yet adopt — a "
       "candidate idea to phase out the flat annual permit, pending Parking City Service permit/tariff data; "
       "(4) widen Organise-Free-Parking into the core; (5) recalibrate the sensitivity zones (Komitas "
       "Medium→High; and treat Gai Avenue / Mega Mall as a site-specific carve-out — defer kerb removal "
       "until the mall's off-street parking is brought into play, rather than inventing a new zone class).", {})])

# ===========================================================================
# 2. SECTION-BY-SECTION
# ===========================================================================
heading("2.  Section-by-section change-list", 1)

change("Cover label", "title page",
       ["REPLACE"],
       "Cover reads “Output 8: Parking Analysis Report”.",
       [("Correct to “Output 10: Parking Analysis Report” (Output 8 is the surveys report). Matched style, "
         "tracked.", {})],
       "Known mislabel; confirmed with you for Output 8.")

change("Introduction — Purpose & Scope", "¶41–53",
       ["INJECT"],
       "Standalone management/mitigation report; demand context described qualitatively.",
       [("Note that a targeted field occupancy survey (six areas, ordinary weekdays) has been completed and "
         "now underpins the measure recommendations, cross-referencing the companion Output 8.", {})],
       "Field survey; memo §2.")

change("Yerevan Parking System Overview — Tariff / ANPR / Payment", "¶63–139",
       ["ADD", "KEEP"],
       "Institutional, tariff (2024 reform), ANPR enforcement and digital-payment context. Factual; "
       "largely sound.",
       [("Keep. Add a short description of the flat annual zonal-parking permit within the tariff/payment "
         "context — this sets up the candidate idea to phase it out (flagged, not committed; see §3). The "
         "report later relies on a turnover argument the annual permit undercuts.", {})],
       "Candidate idea, not a measure; needs PCS data; memo §1/§3.")

change("Parking Surveys — Evolution / Reliance on ANPR / Previous Field Survey Plans", "¶140–192",
       ["REFRAME", "REPLACE"],
       "Treats the field survey as abandoned/planned and demand metrics as ANPR-derived (“Occupancy and "
       "Turnover Metrics from ANPR Data”; “Overview of Previous Field Survey Plans”).",
       [("Recast as survey conducted: retitle “Previous Field Survey Plans” → “Field Occupancy Survey "
         "(conducted)”, add the methodology + weekday-validity table, and replace the ANPR-as-demand-source "
         "framing with the measured results (point to Output 8 for the full dataset). ANPR stays as "
         "existing operational context, not the demand engine.", {})],
       "Survey conducted; memo §2; consistent with Output 8.")

change("Parking Impacts — Differentiated Impacts by Land Use", "¶202–206",
       ["INJECT"],
       "Qualitative land-use narrative (Kentron commercial; Malatia/Nor Nork overnight residential; "
       "Arshakunyats wide ROW).",
       [("Inject measured occupancy/turnover/duration per area: Kentron 138% peak (commercial, "
         "short-stay-dominated); Malatia-Sebastia 54% peak with the highest long-stay/residential signature "
         "(confirming the overnight-storage reading); turnover 5–11/day citywide.", {})],
       "Per-area metrics; memo §6.")

change("Parking Impacts — Impacts on Specific User Groups", "¶207–211",
       ["REPLACE", "INJECT"],
       "Assumption-led: residents (long-duration), delivery/loading, persons with disabilities.",
       [("Anchor to measured, purpose-based groups: Visitors/customers (~63% ≤1h, ~77% ≤2h — the dominant "
         "group), Commuters/workers (~6% all-day → BRT mode-shift target), Residents (overnight ~7%, a "
         "floor the daytime survey under-counts — protect, don't weaken), Delivery/freight (truck pockets "
         "Malatia 12.6% / Shiraz 8.0%), Taxi and PwD (retained by design). Replace assumed shares with "
         "measured ones.", {})],
       [("Stay-mix + vehicle type. Truck share = trucks ÷ all classified vehicles in the area's hourly "
         "sweeps (Malatia 548/4,339 = 12.6%; Shiraz 198/2,469 = 8.0%; 4.9% citywide). Memo §6/§9.", {})])

change("Displacement & Mitigation — Displacement Potential by Corridor Zone", "¶223–228",
       ["INJECT", "RECALIBRATE"],
       "Qualitative High/Medium/Lower descriptions (Kentron high; Arshakunyats/Arabkir medium; outer "
       "lower).",
       [("Add measured peak occupancy + on-street absorption per surveyed area and reconcile labels: "
         "Kentron high but off-street-led (138% / 14%); Komitas reclassified Medium→High (123% / 0%); "
         "Shiraz lower (75% / 69%). Gai Avenue (Mega Mall) is the one area where, even counting every "
         "nearby yard at full capacity, there is still not enough room to re-absorb the cars displaced by "
         "kerb removal (123% peak, best-case absorption still below 100%) — so none of the standard "
         "High/Medium/Lower packages, which all assume the demand lands somewhere, fit it. Simplest "
         "handling: a site-specific carve-out — defer kerb removal here and lean on the Mega Mall's own "
         "off-street structure, rather than creating a new zone category. Add the measured headline "
         "908 vs 1,643 = 55%.", {})],
       "Displacement run; memo §2 & §5 (G2, G6).")

change("Displacement & Mitigation — Yerevan-Specific Framework", "¶238–243",
       ["ADD"],
       "City-specific feasibility factors (enforcement capacity, legal/political context, Soviet urban "
       "form / structural deficit).",
       [("Add the load-bearing off-street finding: 92% of the absorptive capacity that closes the gap to "
         "~100% is gated residential yards counted at gross capacity — so re-absorption is conditional on "
         "opening them (Kentron 86% yard-dependent, Komitas 100%). Frames the elevation of Open Residential "
         "Yards.", {})],
       "Absorption breakdown; memo §1, §4 (G1).")

change("Displacement & Mitigation — Mitigation Measures for Yerevan", "¶244–275  (core of the revision)",
       ["REFRAME", "INJECT", "ADD"],
       "Seven measures, a priori: Extend paid zones · Delivery spaces · Visitor caps · Resident permits "
       "(free, transitional) · Organise free parking · Park & ride · Open Residential yards.",
       [("Per-measure rewrite per the audit memo (detail in §3 table below): re-scope Extend-paid-zones as "
         "demand-rationing in the core; target Delivery bays to freight pockets; re-scope Visitor caps "
         "(demote, lead with pricing, BRT for commuters); design Resident permits as priced/limited from "
         "the outset — using Tbilisi as the cautionary case, not a model: Tbilisi today grants up to two "
         "free vehicles per apartment with no paid resident permit, and is now tightening that to one per "
         "apartment under its new strategy. Widen Organise-free-parking into the core; soften Park & ride "
         "claims; ELEVATE Open Residential Yards to a committed core precondition + resident-engagement "
         "plan; and FLAG (not yet adopt) a candidate idea to phase out the flat annual permit, conditional "
         "on Parking City Service permit-uptake and tariff data. Inject the legality null (~0.2% illegal → "
         "not an enforcement problem) and informal-location finding.", {})],
       "Memo §1, §3, §6, §7 (whole audit).")

change("Displacement & Mitigation — Phasing and Implementation", "¶276–279",
       ["INJECT"],
       "Three-phase, mitigation-before-removal sequencing.",
       [("Reinforce: because re-absorption depends on the gated yards, opening and managing them (and the "
         "pricing) must precede kerb removal — the off-street dependency makes ‘mitigation before removal’ "
         "load-bearing, not advisory.", {})],
       "Off-street dependency; memo §3 (Phasing).")

change("Displacement & Mitigation — Limitations and Residual Risks", "¶280–282",
       ["INJECT"],
       "Acknowledges some net supply reduction is expected/desirable.",
       [("Add the survey-specific caveats (daytime 07:00–24:00 window under-counts overnight; six areas, "
         "not the full corridor; absorption is gross best-case) and the Gai Avenue residual risk (the one "
         "area where best-case capacity still falls short).", {})],
       "Survey caveats; memo §8 (G2).")

change("Recommendations — Mitigation Measure Packages by Sensitivity Zones", "¶296–319 (Figs 10–11)",
       ["RECALIBRATE", "FIGURE"],
       "Three-zone packages (high/medium/lower) with measures grouped; sensitivity map.",
       [("Recalibrate assignments to measured behaviour: Komitas into High; Gai Avenue / Mega Mall handled "
         "as a site-specific carve-out (defer kerb removal pending the mall's off-street parking), not a "
         "new zone class; Open Residential Yards moved up into High. Note the candidate annual-permit idea "
         "against the cross-cutting package as a flagged option, pending operator data. Update Figure 10 "
         "(measures × zone) and Figure 11 (sensitivity map).", {})],
       "Memo §1 (re-allocate), §2; reuse Findings-report visuals.")

change("Conclusions — Key Conclusions & Further Studies", "¶320–342",
       ["INJECT", "REFRAME"],
       "Qualitative conclusions; “Further Studies Required” lists occupancy/demand work as still needed.",
       [("Add measured conclusions (peak 93%, 55% displacement, off-street-conditional re-absorption, zone "
         "recalibration). Prune the ‘further studies’ the survey has now delivered; keep the genuine "
         "residual gaps — overnight/residential demand, full-corridor extension, freight counts, and ANPR "
         "payment-compliance.", {})],
       "Whole survey; memo §1.")

# ===========================================================================
# 3. PER-MEASURE VERDICTS (the core)
# ===========================================================================
doc.add_page_break()
heading("3.  Per-measure changes (the §244–275 rewrite)", 1)
body([("Direct from the Mitigation Measures Audit memo. ‘Verdict’ shows the change to each measure's "
       "write-up.", {})], size=9.5, color=GREY)
GV = {"bold": True, "color": GREEN}
AV = {"bold": True, "color": AMBER}
table(
    ["Measure (as written)", "Verdict", "Change to the text"],
    [["Extend paid zones", [("◆ Re-scope", AV)],
      "Right lever (≈90% of kerb unpriced) but in the core it rations demand, it does not absorb "
      "(Kentron 14% / Komitas 0% side-street capacity). Reframe as demand management in the core."],
     ["Delivery spaces", [("✓ Keep — target", GV)],
      "Confirmed by short-access dominance; size bays to the freight pockets (truck share = trucks ÷ all "
      "classified vehicles surveyed per area: Malatia 12.6%, Shiraz 8.0%; 4.9% citywide) rather than "
      "uniformly."],
     ["Visitor caps", [("◆ Re-scope, not remove", AV)],
      "Data contradicts the premise (≈63% ≤1h; all-day only ~6%; turnover already 5–11/day). Demote to a "
      "targeted contingency for low-turnover spots; lead with pricing; commuters → BRT mode-shift, not "
      "fines."],
     ["Resident permits (free)", [("◆ Design actively", AV)],
      "Keep, but learn from Tbilisi as a cautionary case: it grants up to two free vehicles per apartment "
      "with no paid resident permit — exactly the long-stay, low-turnover problem — which is why its new "
      "strategy tightens this to one per apartment. Yerevan's permit should be priced/limited from the "
      "outset. Opening note: yards/permits need management + communication."],
     ["Organise free parking", [("✓ Confirmed — widen", GV)],
      "Strongly supported (88% unmarked; 24.5% off-carriageway). Widen beyond the lower zones — Komitas "
      "(35% informal) and Kentron (60% angled) need it most."],
     ["Park & ride", [("◆ Don't oversell", AV)],
      "Daytime fleet is short-stay, weakening the ‘many commuters to intercept’ premise. Keep as a study "
      "item; frame BRT/P&R as serving the commuter."],
     ["Open Residential yards", [("✓ Elevate (load-bearing)", GV)],
      "92% of off-street absorption; ~100% re-absorption is conditional on it. Move from optional "
      "outer-zone pilot to committed core precondition — with a resident-engagement and management plan so "
      "it doesn't backfire."],
     ["Phase out flat annual permit", [("◇ Idea — flag, not yet adopt", AV)],
      "IDEA, not yet a measure. The flat annual fee removes the marginal cost of dwell time and so "
      "suppresses turnover — the one lever the data shows is binding. Present as a candidate to cancel "
      "medium-to-long term and replace with occupancy-based pricing, but it must first be substantiated "
      "with Parking City Service permit-uptake and tariff data before adoption as a measure."]],
    widths=[1.7, 1.4, 4.0])

# ===========================================================================
# 4. MEASURED FIGURES (reference)
# ===========================================================================
heading("4.  Measured figures to inject (reference)", 1)
table(["Theme", "Value (surveyed six areas)"],
      [["Occupancy", "Peak (cap-weighted) 93% citywide; Kentron 138 · Komitas 123 · Gai 123 · Garegin 92 · Shiraz 75 · Malatia 54; 69% of zones >85% at peak"],
       ["Turnover / duration", "5–11 veh/space/day; ≈63% ≤1h, ≈77% ≤2h, ~7% overnight (8h+); avg 1.6–2.2 h"],
       ["Displacement", "908 displaced vs 1,643 removed = 55%; on-street absorbs 43% (Kentron 14 / Komitas 0 / Malatia 97); off-street 1,935 (92% gated yards) → ~100% conditional"],
       ["Parameters", "99.8% legal / 0.2% illegal; 24.5% off-carriageway (FP 13.0 / SB 11.5); Kentron 60% angled; 94.8% sedan, truck Malatia 12.6 / Shiraz 8.0"],
       ["Recalibration", "Komitas Medium→High; Shiraz→Lower (confirmed); Gai→site-specific carve-out (defer kerb removal until mall off-street opens); Kentron High but off-street-led"]],
      widths=[1.5, 5.7])

# ===========================================================================
# 5. OPEN DECISIONS
# ===========================================================================
heading("5.  Open decisions and points resolved in review", 1)
bullet([("Per your review, this is presented as a flagged idea, not a measure, pending Parking City Service "
         "permit-uptake and tariff data. Open decision: once that data is in hand, (a) adopt it as a full "
         "measure, or (b) leave it as a standing recommendation in the conclusions.", {})],
       bold_lead="Annual-permit idea. ")
bullet([("Resolved (confirmed): correct the cover from “Output 8” to “Output 10”, as on the surveys "
         "report.", {"color": GREEN})], bold_lead="Cover label. ")
bullet([("Resolved: reuse the Findings-Report charts (occupancy, displacement, layout) for Figures 10–11 "
         "where relevant, rather than generating Output-10-specific versions.", {"color": GREEN})],
       bold_lead="Figures. ")
bullet([("Resolved (yes): apply to a copy “…20260417 (rev).docx” in red text, original untouched.",
         {"color": GREEN})], bold_lead="Working copy. ")
body([("On approval, I will apply each block above directly in the report in red, section by section.",
       {"bold": True, "color": NAVY})], space=2)

doc.save(OUT)
print("WROTE:", OUT)
