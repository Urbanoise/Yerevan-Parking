#!/usr/bin/env python
"""Generate the Output 8 — Proposed Changes (Redline Plan) memo (.docx).

Step-1 deliverable: a section-by-section change-list for
"Parking Surveys and Analysis Report 08042026.docx" (= Output 8; its internal
"Part of Output 10" label is a known mislabel). For review/approval BEFORE any
edit is applied to the source report. Full-rework scope. Figures are from the
validated field-surveys.geojson run + the audit memo.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Output 8 - Proposed Changes (Redline Plan).docx")

NAVY = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6B, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
INK = RGBColor(0x22, 0x22, 0x22)

# change-type tag palette
TAG = {
    "REFRAME":     RGBColor(0x6A, 0x1B, 0x9A),
    "REPLACE":     RED,
    "INJECT":      RGBColor(0x00, 0x6D, 0x77),
    "ADD":         GREEN,
    "RECALIBRATE": AMBER,
    "FIGURE":      RGBColor(0x15, 0x65, 0xC0),
    "KEEP":        GREY,
}

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = INK


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, runs, align="left", fill=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(str(text))
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        r.font.size = Pt(opt.get("size", size))
        if opt.get("white"):
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif opt.get("color") is not None:
            r.font.color.rgb = opt["color"]
    if fill:
        shade(cell, fill)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(15 if level == 1 else 12)
    return p


def body(parts, size=10.5, italic=False, color=None, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", italic)
        r.font.size = Pt(opt.get("size", size))
        c = opt.get("color", color)
        if c is not None:
            r.font.color.rgb = c
    return p


def bullet(parts, bold_lead=None, lead_color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        if lead_color is not None:
            r.font.color.rgb = lead_color
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.bold = opt.get("bold", False)
        r.font.size = Pt(opt.get("size", 10.5))
        c = opt.get("color")
        if c is not None:
            r.font.color.rgb = c
    return p


def rule():
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "8"), ("w:space", "1"), ("w:color", "142A4A")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def table(headers, rows, widths=None, fill="142A4A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], [(h, {"bold": True, "white": True, "size": 9})],
                 align="left" if i == 0 else "center", fill=fill)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            set_cell(cells[i], c, align="left")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def change(section, refs, tag, current, proposed, basis):
    """One change block: section title + ¶refs + tag, then Current / Proposed / Basis."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(section)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    if refs:
        r2 = p.add_run("   " + refs)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GREY
    # tag chip
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(1)
    for i, tg in enumerate(tag if isinstance(tag, list) else [tag]):
        if i:
            tp.add_run("  ")
        rr = tp.add_run(f" {tg} ")
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = TAG.get(tg, GREY)
    _line("Current", current, GREY)
    _line("Proposed", proposed, RED)
    _line("Basis", basis, ACCENT)


def _line(label, parts, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.18)
    rl = p.add_run(f"{label}:  ")
    rl.bold = True
    rl.font.size = Pt(9.5)
    rl.font.color.rgb = color
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opt in parts:
        r = p.add_run(text)
        r.font.size = Pt(opt.get("size", 9.5))
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        c = opt.get("color")
        if c is not None:
            r.font.color.rgb = c


# ===========================================================================
# TITLE
# ===========================================================================
t = doc.add_paragraph()
r = t.add_run("Yerevan Parking Study  •  Output 8 Revision")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

t = doc.add_paragraph()
r = t.add_run("Output 8 — Parking Surveys & Analysis Report: Proposed Changes")
r.bold = True
r.font.size = Pt(19)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
r = sub.add_run("A section-by-section change-list for review and approval. No edits have been applied to "
                "the source report yet — on approval, every change will be made in the document with red "
                "text so it can be tracked.")
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = GREY
rule()

body([("Target file. ", {"bold": True}),
      ("“Parking Surveys and Analysis Report 08042026.docx”. Its internal cover label “Part of Output 10” "
       "is a known mislabel — this is Output 8 (Traffic & Parking Surveys and Analysis Report). The "
       "companion “Parking Analysis Report - 20260417.docx” is Output 10 (it carries the mitigation "
       "measures M1–M10 and is mislabelled “Output 8” on its own cover). Measure-level edits (M7 elevate, "
       "the new annual-permit measure, M3 re-scope) therefore belong mainly to Output 10; this plan keeps "
       "Output 8 to what it owns — the survey, supply, sensitivity zones and the displacement frame.", {})],
     size=10, color=GREY)

body([("How to read. ", {"bold": True}),
      ("Each block names the section and paragraph refs, a change-type tag, then ", {}),
      ("Current", {"bold": True, "color": GREY}), (" → ", {}),
      ("Proposed", {"bold": True, "color": RED}), (" → ", {}),
      ("Basis", {"bold": True, "color": ACCENT}),
      (". Tags: ", {}),
      ("REFRAME", {"bold": True, "color": TAG["REFRAME"]}), (" recast framing · ", {}),
      ("REPLACE", {"bold": True, "color": TAG["REPLACE"]}), (" swap obsolete content · ", {}),
      ("INJECT", {"bold": True, "color": TAG["INJECT"]}), (" add measured numbers into existing prose · ", {}),
      ("ADD", {"bold": True, "color": TAG["ADD"]}), (" new content · ", {}),
      ("RECALIBRATE", {"bold": True, "color": TAG["RECALIBRATE"]}), (" zone reassignment · ", {}),
      ("FIGURE", {"bold": True, "color": TAG["FIGURE"]}), (" update/add a visual.", {})],
     size=9.5, color=INK)

# ===========================================================================
# THE SPINE
# ===========================================================================
heading("1.  The one change that drives all the others", 1)
body([("Output 8 is dated 8 April 2026 — it predates the field survey (29 May – 3 June 2026). It "
       "therefore (a) contains a whole chapter ", {}),
      ("“Justification for Excluding Full-Scale Occupancy and Duration Surveys”", {"italic": True}),
      (", (b) defers all demand data to a ", {}),
      ("pending", {"italic": True}),
      (" municipal ANPR feed (“data-sharing approval … is pending”), and (c) derives displacement from "
       "European precedent rather than local measurement. A targeted field occupancy survey has since "
       "been completed on six representative areas, producing exactly the occupancy, duration, turnover "
       "and locally-measured displacement the report says are unavailable.", {})])
body([("The spine of the revision: ", {"bold": True, "color": RED}),
      ("move Output 8 from “surveys excluded, demand inferred / pending” to “a targeted field survey was "
       "conducted, and here is what it measured.” The proportionality logic (we did not survey all 34 km) "
       "is kept and reframed — it now explains why six representative areas were surveyed rather than the "
       "whole corridor, instead of why no occupancy survey was done at all.", {})])
body([("Scope boundary to hold throughout. ", {"bold": True}),
      ("The field survey is a demand overlay on six areas (208 segments, ~2,930 marked spaces, ~24,950 "
       "observations, 11,394 stays) — not the full 8,862-space inventory. Its displacement figures "
       "(908 displaced vs 1,643 removed = 55%) are for the surveyed areas and are presented as a locally-"
       "measured calibration of “displaced < removed,” not a corridor-wide total. The full-corridor "
       "supply inventory (15,897 spaces; 5,953 removed) stays as the structural backbone.", {})],
     size=10, color=GREY)

# ===========================================================================
# 2. SECTION-BY-SECTION
# ===========================================================================
heading("2.  Section-by-section change-list", 1)

change("Executive Summary", "¶26–29",
       ["REPLACE", "INJECT"],
       "Supply-inventory framing only (corridors, removal scale); no demand findings — there were none.",
       [("Add a measured-findings paragraph: the targeted field survey (6 areas, ordinary weekdays) "
         "found citywide peak occupancy 93% (cap-weighted), 69% of zones over 85% "
         "at peak, short-stay dominance (≈63% ≤1h, ≈77% ≤2h), turnover 5–11/space/day, and peak "
         "displaced demand of only 55% of spaces removed. State up front that displacement is now "
         "locally measured, not just inferred from European precedent.", {})],
       "Field-survey run; audit memo §1.")

change("Introduction — Purpose & Scope; Relationship to ToR", "¶31–40",
       ["REFRAME"],
       "Scope = desk-based supply inventory + “targeted demand analysis using ANPR where available”; "
       "ToR full-survey requirement recorded as refined away.",
       [("Reframe scope to include the completed targeted field occupancy survey as a third, realised "
         "evidence stream. In the ToR sub-section, note that the occupancy/duration requirement — "
         "initially descoped — was subsequently met in representative form by the field survey, so the "
         "deliverable now satisfies the intent of ToR ¶59–63 rather than departing from it.", {})],
       "Reframe; strengthens ToR compliance.")

change("Survey Methodology — Evolution / Adopted Methodology / Data Sources", "¶41–59",
       ["REFRAME", "ADD"],
       "Three-component framework: desk inventory + ANPR + 15% field verification. ANPR named as the "
       "demand-side source. ‘Evolution’ explains why the full field survey was dropped.",
       [("Add a fourth component / new sub-section “Field Occupancy Survey (targeted)”: licence-plate "
         "hourly sweeps 07:00–24:00 across six areas spanning the sensitivity spectrum; per-vehicle "
         "logging of time, zone, plate, vehicle type, kerb location, parking method, legality; sample "
         "size and the area/date table (weekday-validity). Recast ‘Evolution’ so the field survey is the "
         "next step beyond the desk inventory. Per comment #1: since no action follows this deliverable, "
         "strip the ANPR-as-future-demand-engine framing — remove Component 2 (ANPR Data Integration) as "
         "the demand source; the field survey is the demand evidence. Retain at most a one-line factual "
         "note that the municipal ANPR system exists, with no promise of future analysis.", {})],
       "Survey methodology; memo §2 (qualitative→quantitative logic); comment #1.")

change("NEW sub-section — Survey validity (ordinary days)", "after ¶44",
       ["ADD"],
       "Absent.",
       [("Insert a short validity note + table: all six surveys fell on ordinary mid-week working days "
         "(Komitas Fri 29 May; Shiraz Mon 1 Jun; Garegin & Gai Tue 2 Jun; Kentron & Malatia Wed 3 Jun); "
         "no weekend/holiday; surveyed on five separate days, so a one-off event would show as a single-"
         "area outlier — the patterns repeat, which rules out day-specific distortion. Recommend a final "
         "desk-check against the municipal events/holiday calendar.", {})],
       "Survey dates; memo §2.")

change("Justification for Excluding Full-Scale Occupancy and Duration Surveys", "¶60–67  (whole chapter)",
       ["REFRAME", "REPLACE"],
       "Stands as a live justification for NOT collecting occupancy/duration data (analytical rationale, "
       "resource proportionality, ANPR-instead).",
       [("This is the most out-of-date chapter. Retitle to “Survey Scope: From Full-Corridor to Targeted "
         "Occupancy Survey.” Keep the resource-proportionality argument (a 34-km census was "
         "disproportionate) but recast its conclusion: rather than excluding occupancy work entirely, a "
         "proportionate, representative six-area occupancy survey was conducted. Remove/replace language "
         "that asserts occupancy data is unnecessary or unavailable. Cross-reference the new results "
         "chapter.", {})],
       "Survey now exists; memo §2; avoids an internal contradiction.")

change("Parking Supply Inventory Results — Configuration & Location", "¶94–127",
       ["INJECT", "KEEP"],
       "Supply by configuration (parallel 75.3%, 90° 24.7%) and by location (on-street 71.4%) from the "
       "desk inventory. Corridor street-level detail.",
       [("Keep the supply inventory intact (it is the corridor-wide backbone). Add a one-paragraph bridge "
         "distinguishing marked configuration (how spaces are striped) from observed parking behaviour "
         "from the field survey (how vehicles actually parked: 55% parallel / 26% at 45° / 19% "
         "perpendicular by vehicle; Kentron 60% angled by space) — and flag that ~24.5% of observed "
         "parking sat off-carriageway (footpath/setback), the latent informality the supply map cannot "
         "see.", {})],
       "Field parameters (location, method); memo §6.")

change("Parking Impacts — Scale of Removal; Impact by Sensitivity Zone", "¶156–164 (Table 9)",
       ["INJECT", "RECALIBRATE", "FIGURE"],
       "Removal scale from supply (5,953). Impact by Sensitivity Zone (Table 9) assigned qualitatively, "
       "a priori.",
       [("Keep the corridor-wide removal scale. Rebuild Table 9 to carry measured columns for the "
         "surveyed areas — peak occupancy, % zones >85%, turnover, on-street absorption — and reconcile "
         "the a-priori labels with behaviour: Komitas reclassified Medium→High (123% peak, 0% on-street "
         "absorption); Shiraz confirmed Lower (75%, 69%); Kentron confirmed High but for an off-street-"
         "led reason (138%, 14%); Gai/Mega flagged as an unclassified gap (123% peak, absorption <100%). "
         "Update the sensitivity-zone figure/map accordingly.", {})],
       "Per-area occupancy/absorption; memo §2 & §5 (G2, G6).")

change("Displacement Assessment — Understanding / Absorptive Capacity", "¶165–175",
       ["INJECT", "ADD"],
       "Displacement explained via European precedent (displaced < removed); absorptive capacity = "
       "7,035 off-street + 1,857 on-street, framed as gross availability.",
       [("Add the locally-measured result: across the surveyed areas, peak displaced demand was 908 vs "
         "1,643 removed = 55% — the European principle now measured in Yerevan, not borrowed. Add the "
         "absorption breakdown (on-street covers only 43%, and 92% of the off-street capacity that closes "
         "the gap is gated residential yards counted at gross capacity), with per-area absorption (Kentron "
         "14%, Komitas 0%, Garegin 30%, Shiraz 69%, Malatia 97%, Gai 68%). State plainly that the "
         "“re-absorbed” conclusion is conditional on opening those yards — the load-bearing dependency "
         "(handed to Output 10 as M7).", {})],
       "Displacement run; memo §1, §3 (M7), §4 (G1).")

change("Occupancy and Demand Data: Current Status", "¶176–184  (whole chapter)",
       ["REPLACE"],
       "Says demand analysis depends on ANPR; “data-sharing approval from the Mayor’s office is pending”; "
       "results to follow “as an addendum.”",
       [("Replace wholesale with “Occupancy and Demand: Field Survey Results.” Present the measured "
         "occupancy (vs the 85% threshold the chapter already names), duration/stay-mix (visitor/"
         "commuter/resident), turnover, and the legality finding (≈0.2% illegal). Per comment #1, delete "
         "the ANPR dependency and the “pending / Mayor’s approval / addendum” framing outright — no action "
         "follows this deliverable, so the field survey stands as the demand evidence, not a placeholder "
         "for data still to come.", {})],
       "Field-survey results directly answer this chapter’s open questions; comment #1.")

change("NEW sub-section — Field parameters: what every collected variable shows", "new, near results",
       ["ADD"],
       "Absent.",
       [("Add the parameter analysis (mirrors audit memo §6) with the ‘why collected / what it proves’ "
         "logic, including the reassuring nulls: legality ≈99.8% legal / 0.2% illegal (rules out law-"
         "breaking as the problem); kerb location 24.5% informal (footpath/setback), worst in Komitas; "
         "parking method (Kentron 60% angled); vehicle mix 94.8% sedan with truck pockets in Malatia "
         "(12.6%) and Shiraz (8.0%), motorcycles negligible (0.3%).", {})],
       "Field parameters; memo §6. Satisfies the ‘analyse every collected parameter’ requirement.")

change("Conclusions", "¶185–191",
       ["INJECT", "REFRAME"],
       "Five supply-only conclusions (15,897 spaces; 8,862 on-street; 11.3% paid; 5,953 removed; "
       "absorptive capacity).",
       [("Keep the supply conclusions; add measured conclusions: peak occupancy 93% with 69% of zones "
         "over 85%; short-stay/high-turnover kerb; displaced demand 55% of removal; absorption ~100% but "
         "conditional on opening gated yards; sensitivity map recalibrated (Komitas→High, Gai gap). Point "
         "demand-side recommendations to Output 10.", {})],
       "Whole survey; memo §1.")

change("Figures, Tables & Live Link", "List of Figures/Tables; ¶191 link",
       ["FIGURE", "KEEP"],
       "Figures 1–8 supply/absorptive; Table 9 sensitivity; live link to the Vercel app.",
       [("Add field-survey figures (hourly demand vs supply, peak vs average occupancy by area, "
         "displacement balance, duration mix, layout) — reuse the charts already built for the Findings "
         "Report rather than re-author. Update List of Figures/Tables and confirm the live link still "
         "resolves.", {})],
       "Findings-report chart set already generated from the same geojson.")

# ===========================================================================
# 3. MEASURED FIGURES REFERENCE
# ===========================================================================
doc.add_page_break()
heading("3.  Measured figures to inject (reference)", 1)
body([("All from the validated field-surveys.geojson run (six areas) unless marked full-corridor. "
       "Surveyed subset ≠ full inventory — keep the two denominators distinct in the text.", {})],
     size=9.5, color=GREY)

heading("Survey scope & validity", 2)
table(["Item", "Value"],
      [["Areas / dates (all weekdays)",
        "Komitas Fri 29 May · Shiraz Mon 1 Jun · Garegin & Gai Tue 2 Jun · Kentron & Malatia Wed 3 Jun"],
       ["Sample", "208 segments · ~2,930 marked spaces · ~24,950 observations · 11,394 on-street stays"],
       ["Window", "07:00–24:00 (hourly licence-plate sweeps)"]],
      widths=[2.2, 5.0])

heading("Occupancy, turnover, duration", 2)
table(["Metric", "Value"],
      [["Peak occupancy (cap-weighted)", "Citywide 93% · Kentron 138 · Komitas 123 · Gai 123 · Garegin 92 · Shiraz 75 · Malatia 54"],
       ["Zones over 85% at peak", "69% citywide"],
       ["Turnover", "5–11 vehicles/space/day"],
       ["Stay mix", "≈63% ≤1h · ≈77% ≤2h · ≈7% overnight (8h+, 1h-rule) · avg 1.6–2.2 h"]],
      widths=[2.2, 5.0])

heading("Displacement (surveyed areas)", 2)
table(["Metric", "Value"],
      [["Removed / displaced (peak)", "1,643 spaces removed · 908 cars displaced = 55%"],
       ["On-street absorption", "43% citywide · Kentron 14 · Komitas 0 · Garegin 30 · Shiraz 69 · Malatia 97 · Gai 68"],
       ["Off-street absorption", "1,935 spaces, of which 92% gated residential yards (→ Output 10 M7)"],
       ["Total absorbed", "~100%, conditional on opening the gated yards"]],
      widths=[2.2, 5.0])

heading("Field parameters", 2)
table(["Parameter", "Value (≈24,950 obs)"],
      [["Legality", "99.8% legal · 0.2% illegal (max 0.9%, Komitas)"],
       ["Kerb location", "75.5% on-street · 24.5% informal (13.0% footpath, 11.5% setback); worst Komitas 35%"],
       ["Parking method", "By vehicle 55% parallel / 26% 45° / 19% perpendicular · by space Kentron 60% angled, ≤4% elsewhere"],
       ["Vehicle type", "94.8% sedan · 4.9% truck (Malatia 12.6, Shiraz 8.0) · 0.3% motorcycle"]],
      widths=[2.2, 5.0])

heading("Full-corridor backbone (unchanged, keep)", 2)
table(["Item", "Value"],
      [["Inventory", "15,897 spaces / 994 features · 8,862 on-street / 728 segments · 7,035 off-street / 266 facilities"],
       ["Regulation", "11.3% paid (396 Zone A + 633 Zone B) · 87.9% free/unregulated"],
       ["Removal", "5,953 on-street removed (C1 1,836 · C2 3,310 · C3 807)"]],
      widths=[2.2, 5.0])

# ===========================================================================
# 4. OPEN DECISIONS
# ===========================================================================
heading("4.  Open decisions before drafting in red", 1)
body([("Resolved from review comments. ", {"bold": True, "color": GREEN}),
      ("(#0) The ~24,950-observation count is dropped from the Executive Summary and kept only as the "
       "denominator under the field-parameter percentages. (#1) The ANPR-as-future-demand-source thread "
       "is deleted outright — Component 2 and the “Occupancy: Current Status” pending/addendum chapter — "
       "since no action follows this deliverable; the field survey is the demand evidence.", {})],
     size=10, color=GREY)
bullet([("Confirm the six surveyed areas map cleanly onto Output 8’s sensitivity zones and corridor "
         "segments (Komitas, Kentron, Shiraz/Hasratyan, Garegin Nzhdeh, Malatia-Sebastia, Gai/Mega) so "
         "the recalibration lands on the right Table 9 rows.", {})],
       bold_lead="Zone mapping. ")
bullet([("Output 8 keeps the displacement frame; measure-level moves (M7 elevate + comms plan, the new "
         "annual-permit measure, M3 re-scope) go to Output 10. Confirm you want Output 8 to point to "
         "Output 10 for these rather than restate them.", {})],
       bold_lead="O8/O10 division. ")
bullet([("The ‘Justification for Excluding…’ chapter: retire-and-reframe (recommended) vs delete "
         "entirely. Reframing preserves the audit trail of the scope decision.", {})],
       bold_lead="Obsolete chapter. ")
bullet([("Reuse the Findings-Report charts as Output 8 figures, or generate Output-8-specific versions?", {})],
       bold_lead="Figures. ")
bullet([("Apply edits to a copy (e.g. “…08042026 (rev).docx”) with red text, leaving the original "
         "untouched — confirm the filename convention.", {})],
       bold_lead="Working copy. ")

body([("On approval, I will apply each block above directly in the report with red text, section by "
       "section, leaving everything else untouched.", {"bold": True, "color": NAVY})], space=2)

doc.save(OUT)
print("WROTE:", OUT)
