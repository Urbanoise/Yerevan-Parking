#!/usr/bin/env python
"""Backing note: residential-yard opening — international evidence for the
dedicated feasibility/implementation study that Output 10 defers to.

This is NOT part of Output 8 or Output 10. O10 deliberately stops at naming a few
international cases and the candidate organisational models, and defers the
implementation mechanism to a separate study. This note seeds that study: it
captures the fuller implementation detail (organisational models, preconditions,
de-risking machinery, risk register, indicative phasing) WITH explicit source
caveats, so the optimism of the secondary synthesis is not carried into the
project file unchecked.

Two evidence bases are reconciled here:
  • NotebookLM "Yerevan notebook" synthesis (Residential_Yards_Parking_Integration_
    International_Cases.md) — broad, implementation-oriented, uniformly optimistic;
    several figures modelled/aspirational or vendor-sourced.
  • An adversarially cross-checked precedent-research pass run for this project —
    more sobering: no documented post-Soviet case of an association running paid
    public parking in its own gated courtyard; the default post-Soviet pattern is
    gating to EXCLUDE (Kyiv shelved a courtyard-charging plan; Moscow/Tbilisi
    subsidise barriers); Beijing's headline shared-space counts are weighted to
    commercial/government lots and partly the reverse direction (residents renting
    overnight in commercial lots), not day-into-residential at scale.

Where the two disagree, this note flags it rather than averaging it.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
       r"Residential Yards - Implementation Backing Note (for feasibility study).docx")

NAVY = RGBColor(0x14, 0x2A, 0x4A)
ACCENT = RGBColor(0x00, 0x6D, 0x77)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)


def _shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_)
    tcPr.append(sh)


def title(text, sub=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = NAVY
    if sub:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(sub)
        r2.italic = True
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = ACCENT


def heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT
    r.font.size = Pt(13 if level == 1 else 11.5)
    p.space_before = Pt(10)
    return p


def body(text, bold_lead=None, color=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = color or NAVY
    r = p.add_run(text)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, "142A4A")
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            rr = cells[j].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t


def callout(text, lead="Caveat. "):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(lead)
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY
    return p


# ============================================================ HEADER
title("Residential-Yard Opening — Implementation Backing Note",
      "International evidence for the dedicated feasibility/implementation study referenced in "
      "Output 10. Internal project-file note — not part of Output 8 or Output 10.")

t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
for k, v in [
    ("Status", "Backing note / study seed — illustrative, not a recommendation"),
    ("Relation to deliverables", "Output 10 names a few cases and the candidate organisational models, then "
        "defers the mechanism here. This note must NOT be read back into O10 as a committed plan."),
    ("Evidence basis", "NotebookLM 'Yerevan notebook' synthesis + an adversarially cross-checked "
        "precedent-research pass; conflicts between them are flagged, not averaged."),
    ("Bottom line", "A marginal-supply tool worth piloting, not a structural fix — and genuinely "
        "untested in a post-Soviet courtyard context."),
]:
    cells = t.rows[[ "Status","Relation to deliverables","Evidence basis","Bottom line"].index(k)].cells
    cells[0].text = ""
    rr = cells[0].paragraphs[0].add_run(k)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = NAVY
    cells[1].text = ""
    rr2 = cells[1].paragraphs[0].add_run(v)
    rr2.font.size = Pt(9)
    cells[0].width = Inches(1.7)
    cells[1].width = Inches(5.0)

# ============================================================ 1. REALITY CHECK
heading("1. Reality check — read this before the optimistic material")
body("The international literature shows the planning logic is settled (daytime residential "
     "vacancy coincides with commercial peaks) and the technology is off-the-shelf (app booking, "
     "ANPR gating). The binding constraint everywhere is the governance / consent / land-title "
     "layer — and in the post-Soviet case it is compounded by unresolved courtyard ownership and a "
     "resident default of gating yards to keep outsiders out.")
bullet("There is no well-documented case of a post-Soviet residents' association running a paid "
       "public-parking operation in its own gated courtyard. The concept is genuinely novel here.",
       bold_lead="Novelty. ")
bullet("Kyiv floated charging for courtyard parking and shelved it under public backlash, tripped "
       "up by ownership ambiguity. Moscow and Tbilisi subsidise/install courtyard barriers to "
       "exclude outsiders. The demonstrated post-Soviet impulse is exclusion, not sharing.",
       bold_lead="Counter-trend. ")
bullet("Schemes succeed where residents voluntarily rent their own surplus for their own revenue; "
       "they fail where the programme is perceived as taking from residents (top-down charging) or "
       "as self-dealing by an association board.", bold_lead="Success condition. ")
callout("Treat courtyard-sharing as a marginal-supply and demand-management tool, not as the "
        "structural fix for the BRT-displacement deficit. Even showcase cases report aspirational "
        "utilisation targets and owner earnings often too small to motivate participation.")

# ============================================================ 2. INTERNATIONAL CASES
heading("2. International cases (transferable elements)")
table(
    ["Case", "Model type", "Most transferable element", "Evidence quality"],
    [
        ["Stockholm Parkering (SE)", "City-owned intermediary", "City-backed company bridging private lots and "
            "demand; satisfies developer parking needs without new build", "Documented (municipal)"],
        ["Vancouver West End (CA)", "Municipal price alignment", "On-street permit price set at off-street market "
            "rate so cars move into under-used parkades", "Modelled outcomes; policy documented"],
        ["Hangzhou 邻里停 (CN)", "Platform + revenue split", "Revenue split owner / owners' committee / "
            "management — association becomes a paid partner", "Gov-sourced; uptake modest"],
        ["Beijing staggered (CN)", "Policy + enforcement tech", "Time-staggered sharing principle; ANPR / app",
            "Headline counts weighted to commercial/gov lots — see caveat"],
        ["JustPark / Akippa / Share with Oscar", "P2P platforms (UK/JP/AU)", "Trust machinery: insurance, "
            "address privacy, tax allowance, easements", "Vendor-sourced revenue figures"],
        ["King County Metro (US)", "Shared-parking lease agreements", "Leasing residential/church spaces for "
            "Park & Ride; statutory impoundment for overstay", "Documented (transit agency)"],
        ["Stockholm / ParqEx / WhereiPark", "Whole-building monetisation", "Board-controlled opening of surplus "
            "building spaces with tiered commission + ANPR", "Vendor case studies (proof-of-concept)"],
    ],
    widths=[1.7, 1.5, 2.6, 1.4],
)
callout("Beijing's widely-cited ~28,000 shared spaces / 330 lots is largely COMMERCIAL and GOVERNMENT "
        "lots, and partly the reverse direction (residents renting overnight in commercial lots) — not "
        "day-into-residential-courtyard at scale. Do not present it as a like-for-like Yerevan analogue. "
        "Social-credit-linked enforcement is noted in the source but is not an appropriate model to cite "
        "approvingly.")

# ============================================================ 3. ORGANISATIONAL MODELS
heading("3. Candidate organisational models (the choice the study must make)")
table(
    ["Model", "Who leads", "Mechanism", "Trade-off"],
    [
        ["Public / state-agency-led", "Municipal parking authority or PCS; or a dedicated city-owned "
            "company (Stockholm Parkering model)", "City brokers, leases and enforces; association consents "
            "and is paid", "Clear accountability; carries top-down political risk (Kyiv)"],
        ["Public-private partnership", "Licensed operator or parking platform under a municipal framework",
            "Operator runs booking/access/payment; revenue shared with association", "Fast to stand up; "
            "depends on a credible operator and a clean revenue split"],
        ["Association-led, city-enabled", "The hamatirutyun itself, under a municipal template",
            "Association opens and manages its courtyard; light regulation + standard agreement", "Highest "
            "consent legitimacy; weakest guarantee of dependable public capacity"],
        ["Combination", "City framework + PPP operator + association consent", "Layered: enabling regulation, "
            "operator, paid association, revenue-back district", "Most robust on the evidence; most moving parts"],
    ],
    widths=[1.7, 2.0, 2.0, 1.5],
)

# ============================================================ 4. PRECONDITIONS
heading("4. Load-bearing preconditions the study must resolve")
bullet("Who legally owns each courtyard (municipality / hamatirutyun / unregistered). This determines "
       "who may authorise the arrangement and retain revenue — and is the failure point in every "
       "post-Soviet case.", bold_lead="Land title. ")
bullet("Whether the hamatirutyun is recognised as a contracting entity able to execute a shared-parking "
       "agreement on behalf of owners, and the consent procedure / vote threshold under Armenian "
       "condominium law (the owners' meeting is the supreme management body).", bold_lead="Association capacity. ")
bullet("Tax treatment of association parking income (in the US, non-member parking income is taxable / "
       "UBIT; verify the Armenian position — do not assume tax-free), and whether commercial/public use "
       "voids residential insurance cover.", bold_lead="Tax & insurance. ")
bullet("A legal instrument that survives a board change or property sale (covenant running with the land; "
       "the Dutch 'qualitative obligation'; the UK 'deed of easement').", bold_lead="Durability. ")
bullet("Complete the geo-referenced on-street (red/blue-line) inventory first, so yard capacity is mapped "
       "to where it most efficiently replaces removed kerb; define an open data standard so yard "
       "availability can feed the city app.", bold_lead="Inventory & data. ")

# ============================================================ 5. DE-RISKING MACHINERY
heading("5. De-risking machinery observed elsewhere")
table(
    ["Risk", "Mechanism used elsewhere"],
    [
        ["Resident security anxiety", "Address privacy until booking confirmed (Share with Oscar); visible ANPR"],
        ["Overstay blocking residents", "Time-boxed digital passes; statutory impoundment (US RCW 46.55); "
            "ground locks (Beijing) — heavy-handed, weigh carefully"],
        ["Vehicle-damage liability", "Dedicated parking-sharing insurance co-developed with insurers (Akippa + Sompo)"],
        ["Owner inertia / 'untouchable asset'", "Revenue-back to the courtyard (participatory budgeting, Vancouver; "
            "parking benefit district, Shoup)"],
        ["Weak financial incentive", "Hangzhou owner/association/management split; UK tax-free allowance on rental income"],
        ["Platform fragmentation", "Mandate open-API integration into one city dashboard"],
    ],
    widths=[2.3, 4.4],
)

# ============================================================ 6. INDICATIVE PHASING
heading("6. Indicative phasing (illustrative — for the study to test, not a commitment)")
body("Sequencing should follow the international experience that single-owner institutional and "
     "commercial lots open far more readily than residential courtyards and should be activated first.")
bullet("Anchor on willing institutional/commercial lots; pilot 3–5 willing building associations in a "
       "high-sensitivity zone using existing PCS ANPR; price daytime yard access at parity with adjacent "
       "blue-line rates; evaluate uptake, conflict rate and revenue at 6 months.", bold_lead="Phase 1 (pilot). ")
bullet("Integrate yard availability into the PCS app; stand up an association registration framework with "
       "standardised agreements; map further courtyard supply.", bold_lead="Phase 2 (expansion). ")
bullet("Move from resident-exemption pricing toward demand-responsive pricing across on- and off-street "
       "supply; fold revenue into a neighbourhood-governed benefit district.", bold_lead="Phase 3 (maturity). ")

# ============================================================ 7. SOURCE CAVEATS
heading("7. Source caveats")
bullet("The NotebookLM synthesis is uniformly optimistic and omits the post-Soviet counter-trend; this "
       "note restores it.")
bullet("Revenue and scale figures from platforms (JustPark, Spacer, ParqEx, WhereiPark, Akippa) are "
       "vendor-sourced proof-of-concept, not independently audited.")
bullet("Vancouver (88%→60%) and Beijing (20–40% supply reduction; 7.3 Mt CO₂) figures are MODELLED "
       "scenarios, not realised outcomes — do not quote as Yerevan facts.")
bullet("No ITDP/World Bank/GIZ study specific to Yerevan or Tbilisi courtyard-sharing was located; the "
       "topic is under-documented for this context, which is itself a reason to commission the study.")

doc.save(OUT)
print("WROTE:", OUT)
