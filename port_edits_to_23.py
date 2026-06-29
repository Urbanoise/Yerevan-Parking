#!/usr/bin/env python
"""Port the user's red-edited blocks from the 22062026 (rev) backups onto the
content-identical 23062026 base docs, producing 23062026 (rev) deliverables.

The 23062026 files are a content-identical re-save of 22062026 (same text, same
figures — verified by paragraph diff and media CRC compare). So we transplant the
ACTUAL inserted XML elements (paragraphs + the off-street table) from the user's
backup revs — preserving their exact wording and red formatting — into the real
23 base, rather than re-typing. The single 7.2 sentence (a run appended to an
existing paragraph) is re-applied as a red run.
"""
import copy
import shutil
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

RPT = r"C:/Users/user/Yerevan-Parking/Field Surveys/Field Surveys Report/"
SCR = (r"C:/Users/user/AppData/Local/Temp/claude/"
       r"C--Users-user-Yerevan-Parking/7a175dc9-dc3c-4c9c-aba3-e664b7fc504e/scratchpad/")
RED = RGBColor(0xC0, 0x00, 0x00)

SEVEN_TWO = (
    " The administration of residential-yard opening in particular — who should lead it and "
    "under what organisational and legal arrangement — warrants a dedicated study before any "
    "roll-out."
)


def find_p(doc, sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    raise LookupError(sub)


def transplant_after(usr_doc, dst_doc, anchor_sub, n_elems):
    """Deep-copy the n elements following `anchor_sub` in usr_doc, insert them in
    order after the same anchor in dst_doc."""
    ua = find_p(usr_doc, anchor_sub)
    elems, el = [], ua._p.getnext()
    for _ in range(n_elems):
        elems.append(copy.deepcopy(el))
        el = el.getnext()
    ref = find_p(dst_doc, anchor_sub)._p
    for e in elems:
        ref.addnext(e)
        ref = e


def port_o8():
    src = RPT + "Output 8 - Parking Surveys and Analysis Report 23062026.docx"
    rev = RPT + "Output 8 - Parking Surveys and Analysis Report 23062026 (rev).docx"
    shutil.copyfile(src, rev)
    dst = Document(rev)
    usr = Document(SCR + "USER_o8_rev.docx")
    # anchor -> [intro para, off-street table, (emptied) caption para] = 3 elements
    transplant_after(usr, dst, "Measured absorption (surveyed areas)", 3)
    dst.save(rev)
    print("WROTE:", rev)


def port_o10():
    src = RPT + "Output 10 - Parking Analysis Report - 23062026.docx"
    rev = RPT + "Output 10 - Parking Analysis Report - 23062026 (rev).docx"
    shutil.copyfile(src, rev)
    dst = Document(rev)
    usr = Document(SCR + "USER_o10_rev.docx")
    # Comment 1: occupancy paragraph (1 element after the off-street dependency anchor)
    transplant_after(usr, dst, "re-absorption of displaced demand is conditional on off-street", 1)
    # Comment 2: the 3 residential-yard paragraphs after the verdict
    transplant_after(usr, dst, "Elevate to a committed precondition", 3)
    # 7.2: append the brief red sentence to the existing further-studies paragraph
    body = find_p(dst, "willingness-to-pay")
    if SEVEN_TWO.strip() not in body.text:
        r = body.add_run(SEVEN_TWO)
        r.font.color.rgb = RED
    dst.save(rev)
    print("WROTE:", rev)


if __name__ == "__main__":
    port_o8()
    port_o10()
